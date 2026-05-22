from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage
from collections.abc import Generator
from typing import Any

import os
import re
import json
import copy
import requests
import tempfile
from docx import Document
from docx.oxml.ns import qn

class LbyToolsTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        #获取参数如文件url
        payload_json = json.loads(tool_parameters["payload_json"])

        # orig_paras = payload_json["orig_paras"]
        iter2_list = payload_json["iter2_list"]
        orgin_docx_url = payload_json["orgin_docx_url"]
        abbrev_map = payload_json.get("abbrev_map") or []

        tmp_path = download_to_temp(orgin_docx_url)
        para_id_to_text = build_rewrite_map(iter2_list)

        # out_bytes = apply_rewrites_to_docx(tmp_path, para_id_to_text)
        out_bytes = apply_rewrites_to_docx(tmp_path, para_id_to_text, abbrev_map=abbrev_map)
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = "final.docx"

        yield self.create_blob_message(
                blob=out_bytes,          # 直接传入 bytes，不需要 base64 编码
                meta={
                    "mime_type": mime_type,   # 告诉 Dify 这是什么类型的文件
                    "filename": filename,
                },
            )

# XML 1.0 允许的字符范围，去掉不允许的控制字符/NULL等
_ILLEGAL_XML_RE = re.compile(
    r"[\x00-\x08\x0B\x0C\x0E-\x1F]"  # C0控制字符(除\t \n \r)
)

ITALIC_REGEXES = [
    r"\b[pP](?:\s*-\s*value)?\s*(?:<=|>=|<|>|=|≤|≥)\s*(?:0?\.\d+|\.\d+|\d+(?:\.\d+)?)\b",
    r"\bin\s+vivo\b",
    r"\bin\s+vitro\b",
]

def sanitize_for_docx_xml(s: str) -> str:
    if s is None:
        return ""
    # 1) 去掉非法控制字符
    s = _ILLEGAL_XML_RE.sub("", s)
    # 2) 保险起见：去掉 Unicode surrogate（有些模型/流程可能产生）
    s = s.encode("utf-8", "surrogatepass").decode("utf-8", "ignore")
    return s


def build_rewrite_map(iter2_list):
    """
    iter2_list: list[dict], each has 'para_id' and 'rewritten_text'
    Return: dict {para_id: rewritten_text}
    """
    mp = {}
    for item in iter2_list or []:
        pid = item.get("para_id")
        txt = (item.get("rewritten_text") or "").strip()
        if pid and txt:
            mp[pid] = txt
    return mp

def normalize_parenthetical_abbrev(text: str, seen_pairs: set[tuple[str, str]]) -> str:
    """
    仅处理形如: full (ABBR)
    - 首次出现：保留 full (ABBR)
    - 再次出现：替换为 ABBR
    """
    if not text:
        return text

    # full 尽量允许: 字母/数字/空格/连字符/斜杠 等；ABBR 允许大写+数字+连字符
    pattern = re.compile(r'(?P<full>[A-Za-z][A-Za-z0-9 \-\/]+?)\s*\((?P<abbr>[A-Z0-9][A-Z0-9\-]{1,20})\)')

    def repl(m: re.Match):
        full = m.group("full").strip()
        abbr = m.group("abbr").strip()
        key = (full, abbr)
        if key in seen_pairs:
            return abbr
        seen_pairs.add(key)
        return f"{full} ({abbr})"

    return pattern.sub(repl, text)

def normalize_with_abbrev_map(text: str, abbrev_map: list[dict], seen_abbr: set[str]) -> str:
    """
    规则（全文级）：
    - 第一次允许出现 full (abbr) 或 full
    - 只要该 abbr 已经“出现过一次”（不论以哪种形式），后续：
        full (abbr) -> abbr
        full -> abbr
    """
    if not text:
        return text

    items = sorted(
        [x for x in (abbrev_map or []) if x.get("abbr") and x.get("full")],
        key=lambda d: len(d["full"]),
        reverse=True,
    )

    for item in items:
        full = item["full"].strip()
        abbr = item["abbr"].strip()

        # full (abbr) 形式：允许空格变化
        pat_full_paren = re.compile(
            rf"\b{re.escape(full)}\s*\(\s*{re.escape(abbr)}\s*\)",
            flags=re.IGNORECASE
        )
        # 仅 full
        pat_full_only = re.compile(rf"\b{re.escape(full)}\b", flags=re.IGNORECASE)

        if abbr in seen_abbr:
            # 后续出现：括号形式与full都替换成abbr
            text = pat_full_paren.sub(abbr, text)
            text = pat_full_only.sub(abbr, text)
        else:
            # 首次出现：如果本段包含 full(abbr) 或 full，则标记为已出现
            if pat_full_paren.search(text) or pat_full_only.search(text):
                seen_abbr.add(abbr)

    return text

def _run_is_inside_hyperlink_field(run) -> bool:
    """
    识别 Word 域字段形式的超链接：
    fldChar begin -> instrText包含'HYPERLINK' -> fldChar separate -> (显示文本runs) -> fldChar end
    我们只要 run 落在该 field 范围内，就跳过，避免破坏链接结构/位置。
    """
    el = run._element

    # 先向前回溯，看看是否在 begin..end 的 field 区间里，并且 instrText 是 HYPERLINK
    node = el
    saw_begin = False
    saw_hyperlink_instr = False

    # 回溯有限步，避免极端文档死循环
    for _ in range(200):
        node = node.getprevious()
        if node is None:
            break

        # fldChar begin
        fld = node.find(qn("w:fldChar"))
        if fld is not None and fld.get(qn("w:fldCharType")) == "begin":
            saw_begin = True
            break

        # instrText ... HYPERLINK ...
        instr = node.find(qn("w:instrText"))
        if instr is not None:
            txt = (instr.text or "")
            if "HYPERLINK" in txt:
                saw_hyperlink_instr = True

    if not (saw_begin and saw_hyperlink_instr):
        return False

    # 再向后看看当前 run 后面是否能遇到 fldChar end（说明确实在一个field里）
    node = el
    for _ in range(400):
        node = node.getnext()
        if node is None:
            break
        fld = node.find(qn("w:fldChar"))
        if fld is not None and fld.get(qn("w:fldCharType")) == "end":
            return True

    return False

def iter_runs_skip_hyperlink(paragraph):
    """
    跳过两类超链接：
    1) <w:hyperlink> 容器
    2) 域字段(field)形式的 HYPERLINK（fldChar/instrText/separate/end）
    """
    for run in paragraph.runs:
        # 1) 容器型 hyperlink
        el = run._element
        parent = el.getparent()
        in_hlink = False
        while parent is not None:
            if parent.tag == qn("w:hyperlink"):
                in_hlink = True
                break
            parent = parent.getparent()
        if in_hlink:
            continue

        # 2) 域字段型 hyperlink
        if _run_is_inside_hyperlink_field(run):
            continue

        yield run


def replace_text_preserve_hyperlinks(paragraph, new_text: str):
    """
    将 new_text 仅填充到“非超链接run”里，保留 hyperlink 结构与位置
    """
    new_text = sanitize_for_docx_xml(new_text)

    runs = list(iter_runs_skip_hyperlink(paragraph))
    if not runs:
        # 段落全部是超链接等复杂结构，退化策略：不替换（避免破坏链接）
        return

    # 计算原“非超链接文本”总长度，用于按比例回填
    old_lengths = [len(r.text or "") for r in runs]
    total_old = sum(old_lengths)

    if total_old == 0:
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""
        return

    consumed = 0
    for i, (r, l) in enumerate(zip(runs, old_lengths)):
        if i == len(runs) - 1:
            part = new_text[consumed:]
        else:
            take = round(len(new_text) * (l / total_old))
            part = new_text[consumed:consumed + take]

        r.text = sanitize_for_docx_xml(part)
        consumed += len(part)

def apply_italics_in_paragraph(paragraph, regexes=ITALIC_REGEXES):
    pattern = re.compile("|".join(f"(?:{r})" for r in regexes), flags=re.IGNORECASE)

    for run in list(iter_runs_skip_hyperlink(paragraph)):
        txt = run.text or ""
        if not txt:
            continue

        matches = list(pattern.finditer(txt))
        if not matches:
            continue

        base_rPr = run._element.rPr
        run.text = ""  # 清空原run文本，原run保留位置作为锚点

        prev = run
        last = 0

        def _insert_piece(piece_text: str, italic: bool):
            nonlocal prev
            if piece_text == "":
                return
            newr = paragraph.add_run(piece_text)
            prev._element.addnext(newr._element)
            prev = newr

            if base_rPr is not None:
                new_rpr = newr._element.get_or_add_rPr()
                for child in list(new_rpr):
                    new_rpr.remove(child)
                new_rpr.append(copy.deepcopy(base_rPr))

            if italic:
                newr.italic = True

        for m in matches:
            _insert_piece(txt[last:m.start()], italic=False)
            _insert_piece(txt[m.start():m.end()], italic=True)
            last = m.end()

        _insert_piece(txt[last:], italic=False)

def reset_run_color_to_default(run):
    """
    将run的字体颜色恢复为默认（移除<w:color>设置）
    """
    r = run._element
    rPr = r.rPr
    if rPr is None:
        return

    color = rPr.find(qn('w:color'))
    if color is not None:
        rPr.remove(color)

def reset_paragraph_run_colors_to_default(paragraph):
    for run in paragraph.runs:
        # 只要存在颜色设置就移除；没有则跳过
        rPr = run._element.rPr
        if rPr is None:
            continue
        if rPr.find(qn('w:color')) is not None:
            reset_run_color_to_default(run)

def redistribute_text_to_docx_runs(new_text, runs):
    """
    runs: list[docx.text.run.Run]
    原 run 样式保留，仅替换 run.text
    """
    if not runs:
        return

    new_text = sanitize_for_docx_xml(new_text)

    old_lengths = [len(r.text or "") for r in runs]
    total_old = sum(old_lengths)

    if total_old == 0:
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""
        return

    consumed = 0
    for i, (r, l) in enumerate(zip(runs, old_lengths)):
        if i == len(runs) - 1:
            part = new_text[consumed:]
        else:
            take = round(len(new_text) * (l / total_old))
            part = new_text[consumed:consumed + take]

        r.text = sanitize_for_docx_xml(part)
        consumed += len(part)

def apply_rewrites_to_docx(input_docx_path, para_id_to_text,abbrev_map):
    """
    读取 input_docx_path -> 回填 -> 保存到临时docx -> 读取bytes(blob) -> 删除临时文件(输出+输入) -> 返回bytes
    """
    doc = Document(input_docx_path)

    # for idx, p in enumerate(doc.paragraphs, start=1):
    #     pid = f"p_{idx:05d}"
    #     new_text = para_id_to_text.get(pid)
    #     if not new_text:
    #         continue

    #     redistribute_text_to_docx_runs(new_text, p.runs)
    #     reset_paragraph_run_colors_to_default(p)
    seen_abbr  = set()

    for idx, p in enumerate(doc.paragraphs, start=1):
        pid = f"p_{idx:05d}"
        new_text = para_id_to_text.get(pid)
        if not new_text:
            continue

        # 仅处理 full (ABBR) 范式：首次保留，后续只保留 ABBR
        # new_text = normalize_parenthetical_abbrev(new_text, seen_pairs)
        new_text = normalize_with_abbrev_map(new_text, abbrev_map or [], seen_abbr)

        # 替换文本：保留超链接结构
        replace_text_preserve_hyperlinks(p, new_text)

        # 清颜色（你的原逻辑保留）
        reset_paragraph_run_colors_to_default(p)

        # 斜体短语
        apply_italics_in_paragraph(p)

    out_tmp_path = None
    try:
        fd, out_tmp_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)  # 避免 Windows 文件占用

        doc.save(out_tmp_path)

        with open(out_tmp_path, "rb") as f:
            blob = f.read()

        return blob

    finally:
        # 删除输出临时文件
        if out_tmp_path and os.path.exists(out_tmp_path):
            try:
                os.remove(out_tmp_path)
            except Exception:
                pass

        # 删除输入文件（input_docx_path）
        if input_docx_path and os.path.exists(input_docx_path):
            try:
                os.remove(input_docx_path)
            except Exception:
                pass


def download_to_temp(url: str, suffix=".docx", headers=None, timeout=60) -> str:
    headers = headers or {}
    r = requests.get(url, stream=True, headers=headers, timeout=timeout)
    r.raise_for_status()
    fd, path = tempfile.mkstemp(suffix=suffix)
    os.close(fd)
    with open(path, "wb") as f:
        for chunk in r.iter_content(chunk_size=1024 * 1024):
            if chunk:
                f.write(chunk)
    return path
