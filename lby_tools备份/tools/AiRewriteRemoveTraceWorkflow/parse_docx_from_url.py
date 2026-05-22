from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage
from collections.abc import Generator
from typing import Any

import os
import re
import tempfile
import requests
from docx import Document

class LbyToolsTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        #获取参数如文件url
        docx_url = tool_parameters["docx_url"]
        paras_json, plain_paras, ref_start = extract_main_paragraphs(docx_url)
        yield self.create_json_message({
            "paras_json": paras_json,
            "plain_paras": plain_paras,
            "ref_start": ref_start,
        })


def norm_title(s: str) -> str:
    return re.sub(r"\s+", "", (s or "").strip().lower())

def download_to_temp(url: str, suffix=".docx", headers=None, timeout=60) -> str:
    headers = headers or {}
    r = requests.get(url, stream=True, headers=headers, timeout=timeout)
    r.raise_for_status()
    fd, path = tempfile.mkstemp(suffix=suffix)
    os.close(fd)
    with open(path, "wb") as f:
        for chunk in r.iter_content(chunk_size=1024 * 1024):
            if chunk:
                f.write(chunk)
    return path

def extract_main_paragraphs(docx_url: str):
    try:
        tmp_path = download_to_temp(docx_url)
        
        doc = Document(tmp_path)

        paras_json = []
        ref_start = None

        # 注意：doc.paragraphs 不含表格内段落；表格天然“保留不改”
        for i, p in enumerate(doc.paragraphs, start=1):
            para_id = f"p_{i:05d}"
            text = "".join(r.text or "" for r in p.runs)

            title_norm = norm_title(text)
            if ref_start is None and title_norm in ("references", "参考文献"):
                ref_start = para_id

            rewrite_flag = True
            if ref_start is not None:
                # references 起点及之后全部不改
                if i >= int(ref_start.split("_")[1]):
                    rewrite_flag = False

            runs = []
            for r in p.runs:
                runs.append({
                    "text": r.text or "",
                    "bold": bool(r.bold),
                    "italic": bool(r.italic),
                    "underline": bool(r.underline),
                    "font": r.font.name,
                    "size": float(r.font.size.pt) if r.font.size else None,
                    "style": r.style.name if r.style else None,
                })

            paras_json.append({
                "index": i,
                "para_id": para_id,
                "style": p.style.name if p.style else None,
                "is_heading": (p.style.name.startswith("Heading") if p.style else False),
                "rewrite_flag": rewrite_flag and bool(text.strip()),
                "runs": runs,
                "text": text
            })

        plain_paras = [
            {"para_id": x["para_id"], "text": x["text"], "rewrite_flag": x["rewrite_flag"]}
            for x in paras_json
        ]
        return paras_json, plain_paras, ref_start
    finally:
        try:
            os.remove(tmp_path)
        except OSError:
            pass
