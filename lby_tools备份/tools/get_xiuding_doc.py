from collections.abc import Generator
from typing import Any

from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage
from io import BytesIO

import requests
from docx import Document
from docx.oxml.ns import qn


class LbyToolsTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        #获取参数如文件url
        docx_url = tool_parameters["old_file_url"]
        

        # out_bytes = read_docx_with_revisions_from_url(docx_url, include_deleted=False)
        text = read_docx_with_revisions_from_url(docx_url, include_deleted=False)

        yield self.create_text_message(
            # blob=out_bytes,          # 直接传入 bytes，不需要 base64 编码
            # meta={
            #     "mime_type": "text/plain; charset=utf-8",
            #     # "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",   # 告诉 Dify 这是什么类型的文件
            # },
            text
        )



W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def _iter_text_in_element(el, include_deleted=False):
    """
    Extract visible text from an oxml element, including tracked insertions (<w:ins>).
    By default, deletions (<w:del>) are ignored.
    """
    out = []

    def walk(node):
        tag = node.tag

        # If deletion and not included -> skip whole subtree
        if (not include_deleted) and tag == qn("w:del"):
            return

        # Normal text
        if tag == qn("w:t"):
            if node.text:
                out.append(node.text)

        # Deleted text is stored in <w:delText>
        if include_deleted and tag == qn("w:delText"):
            if node.text:
                out.append(node.text)

        # Tab / line break
        if tag == qn("w:tab"):
            out.append("\t")
        if tag in (qn("w:br"), qn("w:cr")):
            out.append("\n")

        for child in node:
            walk(child)

    walk(el)
    return "".join(out)

def read_docx_with_revisions_from_url(url: str, include_deleted=False) -> str:
    r = requests.get(url, timeout=60)
    r.raise_for_status()
    doc = Document(BytesIO(r.content))

    parts = []

    # paragraphs (body)
    for p in doc.paragraphs:
        txt = _iter_text_in_element(p._p, include_deleted=include_deleted).strip()
        if txt:
            parts.append(txt)

    # tables
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                txt = _iter_text_in_element(cell._tc, include_deleted=include_deleted).strip()
                if txt:
                    parts.append(txt)


    text = "\n".join(parts)

    # text -> bytes（UTF-8）
    return text.encode("utf-8")

    # # 如果你需要 BytesIO 字节流对象
    # bio = BytesIO(b)

    # # 需要读出 bytes
    # return bio.getvalue()


