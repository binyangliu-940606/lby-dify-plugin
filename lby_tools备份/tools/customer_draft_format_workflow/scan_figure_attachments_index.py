# import os
# import re
# import json

# from collections.abc import Generator
# from typing import Dict, Any, List

# from dify_plugin import Tool
# from dify_plugin.entities.tool import ToolInvokeMessage


# class LbyToolsTool(Tool):
#     def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
#         #获取参数如文件url
#         payload_json = json.loads(tool_parameters["payload_json"])
#         figure_raw_dir = payload_json["figure_raw_dir"]
#         figure_files_meta = payload_json["figure_files_meta"]

#         yield self.create_json_message(
#             scan_figure_attachments_index(figure_raw_dir, figure_files_meta) 
#         )


# def _parse_fig_from_name(name: str) -> Dict[str, Any]:
#     """
#     从文件名尽力解析 Figure 编号与 S 标记（确定性规则，不推理）
#     """
#     base = os.path.splitext(os.path.basename(name))[0]
#     low = base.lower()

#     # 提取 S?数字
#     m = re.search(r"(?:figure|fig|图)\s*([sS]?\s*\d+)", base, flags=re.I)
#     num = None
#     sup = False
#     if m:
#         raw = m.group(1).replace(" ", "")
#         if raw.upper().startswith("S"):
#             sup = True
#             num = raw[1:]
#         else:
#             num = raw

#     return {"num": num, "supplementary": sup, "base": base, "lower": low}

# def scan_figure_attachments_index(figure_raw_dir: str, figure_files_meta: List[Dict[str, Any]]) -> Dict[str, Any]:
#     idx = []
#     for f in figure_files_meta or []:
#         name = f["name"]
#         path = f["path"]
#         ext = f["ext"]
#         p = _parse_fig_from_name(name)
#         idx.append({
#             "name": name,
#             "path": path,
#             "ext": ext,
#             "num": p["num"],
#             "supplementary": p["supplementary"],
#             "lower": p["lower"]
#         })
#     return {"figure_attachment_index": idx}




import os
import re
import json
import zipfile
import hashlib
from urllib.parse import urlparse

import requests

from collections.abc import Generator
from typing import Dict, Any, List, Optional
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx import Document
from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage


class LbyToolsTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        payload_json = json.loads(tool_parameters["payload_json"])
        figure_raw_dir = payload_json["figure_raw_dir"]
        figure_files_meta = payload_json.get("figure_files_meta") or []

        manuscript_url = payload_json.get("orig_path")  # 新增：原稿url（可选）

        # 1) 扫描已有附件索引（原逻辑）
        scanned_idx = scan_figure_attachments_index(figure_raw_dir, figure_files_meta)["figure_attachment_index"]

        # 2) 从原稿提取图片并生成索引（新逻辑）
        extracted_idx = []
        if manuscript_url:
            extracted_idx = extract_figures_from_manuscript_url(
                manuscript_url=manuscript_url,
                save_dir=figure_raw_dir
            )

        # 3) 合并去重
        merged = merge_dedup_figure_index(scanned_idx, extracted_idx)

        yield self.create_json_message({"figure_attachment_index": merged})


def _parse_fig_from_name(name: str) -> Dict[str, Any]:
    base = os.path.splitext(os.path.basename(name))[0]
    low = base.lower()

    m = re.search(r"(?:figure|fig|图)\s*([sS]?\s*\d+)", base, flags=re.I)
    num = None
    sup = False
    if m:
        raw = m.group(1).replace(" ", "")
        if raw.upper().startswith("S"):
            sup = True
            num = raw[1:]
        else:
            num = raw
    return {"num": num, "supplementary": sup, "base": base, "lower": low}


def _parse_caption_to_fig_token(text: str):
    if not text:
        return None

    t = re.sub(r"\s+", " ", text.strip())

    # 不用 \b，避免中文导致边界不成立
    # m = re.match(r"^(?:figure|fig|图)\s*(s?\s*\d+)", t, flags=re.I)
    m = re.match(r"^(?:figure|fig|图)[.\s]*(\s*s?\d+)", t, flags=re.I)
    if not m:
        return None

    raw = re.sub(r"\s+", "", m.group(1))  # "S 1" -> "S1"
    if raw.upper().startswith("S"):
        return f"Figure_S{raw[1:]}"
    return f"Figure_{raw}"


def _ensure_figures_dir(figure_raw_dir: str) -> str:
    figures_dir = os.path.join(figure_raw_dir, "figure_raw")
    if not os.path.isdir(figures_dir):
        os.makedirs(figures_dir, exist_ok=True)
    return figures_dir


def _paragraph_has_image(paragraph) -> bool:
    # paragraph._p 是 lxml element，判断是否含有 a:blip（图片引用）
    return bool(paragraph._p.xpath(".//a:blip"))


def _extract_rId_list_from_paragraph(paragraph) -> List[str]:
    # 一个段落里可能有多张图，取所有 embed 的 rId
    rIds = []
    blips = paragraph._p.xpath(".//a:blip")
    for b in blips:
        rid = b.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if rid:
            rIds.append(rid)
    return rIds


def extract_figures_from_manuscript_url(manuscript_url: str, save_dir: str) -> List[Dict[str, Any]]:
    # 你现在传的是本地 orig_path，这里直接当本地路径用
    local_path = manuscript_url
    ext = os.path.splitext(local_path)[1].lower()
    if ext == ".docx":
        return _extract_from_docx(local_path, save_dir)
    return []


def _deduce_ext_from_partname(partname: str) -> str:
    e = os.path.splitext(partname)[1].lstrip(".").lower()
    return e or "bin"


# def _extract_from_docx(docx_path: str, figure_raw_dir: str) -> List[Dict[str, Any]]:
#     """
#     按“图片所在段落的下一段 caption”命名并保存到 figure_raw_dir/Figures
#     """
#     figures_dir = _ensure_figures_dir(figure_raw_dir)
#     doc = Document(docx_path)

#     idx: List[Dict[str, Any]] = []

#     paragraphs = list(doc.paragraphs)
#     for i, p in enumerate(paragraphs):
#         if not _paragraph_has_image(p):
#             continue

#         # caption：下一段文本（按你的需求）
#         caption = paragraphs[i + 1].text.strip() if i + 1 < len(paragraphs) else ""
#         token = _parse_caption_to_fig_token(caption)

#         rids = _extract_rId_list_from_paragraph(p)
#         for k, rid in enumerate(rids):
#             rel = doc.part.rels.get(rid)
#             if not rel or rel.reltype != RT.IMAGE:
#                 continue

#             img_part = rel.target_part
#             data = img_part.blob
#             ext = _deduce_ext_from_partname(img_part.partname)

#             # 文件名策略：
#             # 1) token 存在：Figure_1 / Figure_S1
#             # 2) token 不存在：Image_<段落序号>_<k>
#             if token:
#                 base_name = token
#             else:
#                 base_name = f"Image_{i+1}_{k+1}"

#             # 同一个 Figure 可能多张子图：追加 _2/_3
#             # 注意：如果 token 存在且 rids >1，追加序号避免覆盖
#             suffix = f"_{k+1}" if (token and len(rids) > 1) else ""
#             save_name = f"{base_name}{suffix}.{ext}"
#             save_path = os.path.join(figures_dir, save_name)

#             # 若重名，追加内容hash
#             if os.path.exists(save_path):
#                 h = hashlib.md5(data).hexdigest()[:10]
#                 save_name = f"{base_name}{suffix}_{h}.{ext}"
#                 save_path = os.path.join(figures_dir, save_name)

#             with open(save_path, "wb") as f:
#                 f.write(data)

#             pinfo = _parse_fig_from_name(save_name)
#             idx.append({
#                 "name": save_name,
#                 "path": save_path,
#                 "ext": ext,
#                 "num": pinfo["num"],
#                 "supplementary": pinfo["supplementary"],
#                 "lower": pinfo["lower"],
#             })

#     return idx
def _extract_from_docx(docx_path: str, figure_raw_dir: str) -> List[Dict[str, Any]]:
    """
    按“图片紧接着的 caption”命名并保存到 figure_raw_dir/Figures
    """

    figures_dir = _ensure_figures_dir(figure_raw_dir)
    doc = Document(docx_path)

    idx: List[Dict[str, Any]] = []

    paragraphs = list(doc.paragraphs)
    for i, p in enumerate(paragraphs):
        if not _paragraph_has_image(p):
            continue

        # ================== 修改部分 ==================
        caption = ""
        text_after_image = ""
        has_seen_image = False
        
        # 1. 遍历当前段落的 run，剥离出图片前后的文字
        for run in p.runs:
            # 通过底层 XML 判断该 run 是否包含图片 (w:drawing 或 v:shape/imagedata)
            xml_str = run._element.xml
            if '<w:drawing' in xml_str or '<v:shape' in xml_str or '<v:imagedata' in xml_str:
                has_seen_image = True
                text_after_image = "" # 遇到新图片则清空，确保抓取的是段落内【最后一张图片】后的文字
            else:
                if has_seen_image:
                    text_after_image += run.text

        text_after_image = text_after_image.strip()

        # 2. 如果当前段落图片后面有文字（即文字B），则作为 caption
        if text_after_image:
            caption = text_after_image
        else:
            # 3. 如果当前段落图片后没有文字，则去寻找下一个非空段落
            for j in range(i + 1, len(paragraphs)):
                next_text = paragraphs[j].text.strip()
                if next_text:
                    caption = next_text
                    break
        # ==============================================

        token = _parse_caption_to_fig_token(caption)

        rids = _extract_rId_list_from_paragraph(p)
        for k, rid in enumerate(rids):
            rel = doc.part.rels.get(rid)
            if not rel or rel.reltype != RT.IMAGE:
                continue

            img_part = rel.target_part
            data = img_part.blob
            ext = _deduce_ext_from_partname(img_part.partname)

            if token:
                base_name = token
            else:
                base_name = f"Image_{i+1}_{k+1}"

            suffix = f"_{k+1}" if (token and len(rids) > 1) else ""
            save_name = f"{base_name}{suffix}.{ext}"
            save_path = os.path.join(figures_dir, save_name)

            if os.path.exists(save_path):
                h = hashlib.md5(data).hexdigest()[:10]
                save_name = f"{base_name}{suffix}_{h}.{ext}"
                save_path = os.path.join(figures_dir, save_name)

            with open(save_path, "wb") as f:
                f.write(data)

            pinfo = _parse_fig_from_name(save_name)
            idx.append({
                "name": save_name,
                "path": save_path,
                "ext": ext,
                "num": pinfo["num"],
                "supplementary": pinfo["supplementary"],
                "lower": pinfo["lower"],
            })

    return idx


def scan_figure_attachments_index(figure_raw_dir: str, figure_files_meta: List[Dict[str, Any]]) -> Dict[str, Any]:
    idx = []
    for f in figure_files_meta or []:
        name = f["name"]
        path = f["path"]
        ext = f["ext"]
        p = _parse_fig_from_name(name)
        idx.append({
            "name": name,
            "path": path,
            "ext": ext,
            "num": p["num"],
            "supplementary": p["supplementary"],
            "lower": p["lower"]
        })
    return {"figure_attachment_index": idx}


def merge_dedup_figure_index(a: List[Dict[str, Any]], b: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    seen = set()

    def key(x: Dict[str, Any]) -> str:
        # 以 figure 号优先去重（更符合“图1”这种命名），没有再用 lower+ext
        num = (x.get("num") or "").strip()
        sup = bool(x.get("supplementary"))
        ext = (x.get("ext") or "").strip().lower()
        lower = (x.get("lower") or "").strip()

        if num:
            return f"f:{'S' if sup else ''}{num}.{ext}"
        if lower and ext:
            return f"n:{lower}.{ext}"
        return f"p:{(x.get('path') or '').strip()}"

    for item in (a or []) + (b or []):
        k = key(item)
        if k in seen:
            continue
        seen.add(k)
        out.append(item)
    return out
