import re
import json

from typing import Dict, Any, List
from collections.abc import Generator
from docx import Document
from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage


class LbyToolsTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        #获取参数如文件url
        payload_json = json.loads(tool_parameters["payload_json"])
        orig_path = payload_json["orig_path"]
        normalized_fig_keys = payload_json["normalized_fig_keys"]

        yield self.create_json_message({
            "return_data": extract_fig_legends_from_orig(orig_path, normalized_fig_keys),
        })

FIG_TITLE_PAT = re.compile(r"(?i)^\s*(Figure|Fig\.?|图)\s*(S?\s*\d+)\s*[:：.\-]?\s*(.*)$")

def _norm_fig_key(num_str: str) -> str:
    s = num_str.replace(" ", "")
    if s.upper().startswith("S"):
        return f"Figure S{s[1:]}"
    return f"Figure {s}"

def extract_fig_legends_from_orig(orig_path: str, normalized_fig_keys: List[str]) -> Dict[str, Any]:
    """
    从原稿提取图题与图注（不做生成）。
    图题段落：以 Figure/Fig/图 + 编号 开头。
    图注：取后续若干段，直到遇到下一张图题或空行/结束。
    子图注：按 A: 或 (A) 结构简单切分（不做推理）。
    """
    want = set(normalized_fig_keys or [])
    doc = Document(orig_path)

    legends: Dict[str, Any] = {}
    missing = []

    i = 0
    while i < len(doc.paragraphs):
        txt = (doc.paragraphs[i].text or "").strip()
        m = FIG_TITLE_PAT.match(txt)
        if not m:
            i += 1
            continue

        num = m.group(2)
        title_rest = (m.group(3) or "").strip()
        fig_key = _norm_fig_key(num)

        # 收集 caption：从下一段开始直到遇到下一个图题或连续空行
        cap_lines = []
        j = i + 1
        empty_count = 0
        while j < len(doc.paragraphs):
            t2 = (doc.paragraphs[j].text or "").strip()
            if FIG_TITLE_PAT.match(t2):
                break
            if t2 == "":
                empty_count += 1
                if empty_count >= 2:
                    break
            else:
                empty_count = 0
                cap_lines.append(t2)
            j += 1

        caption = "\n".join(cap_lines).strip()

        # 子图注拆分：只按固定模式 A: 或 (A)
        caption_by_part = {}
        if caption:
            # A: / B: / C:
            part_pat1 = re.compile(r"(?m)^\s*([A-Z])\s*:\s*")
            # (A) / (B)
            part_pat2 = re.compile(r"(?m)^\s*\(([A-Z])\)\s*")

            # 统一找所有标记位置
            marks = []
            for mm in part_pat1.finditer(caption):
                marks.append((mm.start(), mm.group(1), "A:"))
            for mm in part_pat2.finditer(caption):
                marks.append((mm.start(), mm.group(1), "(A)"))
            marks.sort(key=lambda x: x[0])

            if len(marks) >= 2:
                for idx, (pos, letter, _) in enumerate(marks):
                    end = marks[idx + 1][0] if idx + 1 < len(marks) else len(caption)
                    chunk = caption[pos:end].strip()
                    # 去掉头部标记
                    chunk = re.sub(r"^\s*(?:[A-Z]\s*:|\([A-Z]\))\s*", "", chunk).strip()
                    caption_by_part[letter] = chunk

        if fig_key in want:
            legends[fig_key] = {
                "title": title_rest,
                "caption": caption,
                "caption_by_part": caption_by_part
            }

        i = j

    for k in want:
        if k not in legends:
            missing.append(k)

    return {"fig_legends": legends, "fig_legend_missing": missing}