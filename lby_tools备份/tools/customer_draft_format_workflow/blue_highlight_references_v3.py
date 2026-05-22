import json
import re
import os


# from collections.abc import Generator
from typing import List, Dict, Any, Tuple
from docx import Document
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph

# from dify_plugin import Tool
# from dify_plugin.entities.tool import ToolInvokeMessage



# class LbyToolsTool(Tool):
#     def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
#         #获取参数如文件url
#         payload_json = json.loads(tool_parameters["payload_json"])
#         docx_in = payload_json["docx_in"]
#         normalized_fig_keys = payload_json["normalized_fig_keys"]
#         normalized_table_keys = payload_json["normalized_table_keys"]
#         normalized_additional_keys = payload_json["normalized_additional_keys"]

        
#         base_dir = os.path.dirname(docx_in)
#         base_name = os.path.splitext(os.path.basename(docx_in))[0]
#         docx_out = os.path.join(base_dir, f"{base_name}.blue.docx")

#         yield self.create_json_message({
#             "return_data": blue_highlight_references_v3(
#                 docx_in,
#                 docx_out,
#                 normalized_fig_keys,
#                 normalized_table_keys,
#                 normalized_additional_keys,
#             ),
#         })


BLUE = RGBColor(29, 65, 213)

# PMID 格式
PMID_GROUP_PAT = re.compile(r"\(PMID:\s*\d+(?:;\s*PMID:\s*\d+)*\)")
PMID_SINGLE_PAT = re.compile(r"\bPMID\s*[:：]\s*\d{5,10}\b")

# 图例/表例标题行（用于识别 legend 开始）
FIG_LEGEND_TITLE_PAT = re.compile(r"(?i)^\s*Figure\s+S?\d+\s*\.")
TAB_LEGEND_TITLE_PAT = re.compile(r"(?i)^\s*Table\s+S?\d+\s*\.")

# 用于判断 legend 结束的“下一块区域”标记（可按需扩展）
SECTION_STOP_PAT = re.compile(r"(?i)^\s*(Figure\s+S?\d+\s*\.|Table\s+S?\d+\s*\.|References|参考文献)\b")


def _build_run_char_map(p: Paragraph) -> List[Tuple[int, int, int]]:
    cur = 0
    mapping = []
    for i, r in enumerate(p.runs):
        t = r.text or ""
        mapping.append((i, cur, cur + len(t)))
        cur += len(t)
    return mapping


def _split_run_at(p: Paragraph, run_idx: int, split_pos_in_run: int):
    run = p.runs[run_idx]
    txt = run.text or ""
    left = txt[:split_pos_in_run]
    right = txt[split_pos_in_run:]

    run.text = left
    new_run = p.add_run(right)

    # 复制样式
    new_run.bold = run.bold
    new_run.italic = run.italic
    new_run.underline = run.underline
    if run.font is not None:
        new_run.font.name = run.font.name
        new_run.font.size = run.font.size
        if run.font.color and run.font.color.rgb:
            new_run.font.color.rgb = run.font.color.rgb

    run._r.addnext(new_run._r)


def _merge_ranges(ranges: List[Tuple[int, int]]) -> List[Tuple[int, int]]:
    if not ranges:
        return []
    ranges = sorted(ranges, key=lambda x: (x[0], x[1]))
    merged = [ranges[0]]
    for s, e in ranges[1:]:
        ps, pe = merged[-1]
        if s <= pe:
            merged[-1] = (ps, max(pe, e))
        else:
            merged.append((s, e))
    return merged


def _colorize_paragraph_ranges(p: Paragraph, ranges: List[Tuple[int, int]]):
    ranges = _merge_ranges([(s, e) for s, e in ranges if e > s])
    if not ranges:
        return

    for s, e in ranges:
        mapping = _build_run_char_map(p)
        if not mapping:
            continue
        total_len = mapping[-1][2]

        s = max(0, min(s, total_len))
        e = max(0, min(e, total_len))
        if e <= s:
            continue

        def locate(pos: int):
            for run_idx, a, b in mapping:
                if a <= pos < b:
                    return run_idx, pos - a, a, b
            if pos == total_len:
                last = mapping[-1]
                return last[0], last[2] - last[1], last[1], last[2]
            return None

        L = locate(s)
        R = locate(e)
        if L is None or R is None:
            continue

        l_run, l_in, l_a, l_b = L
        r_run, r_in, r_a, r_b = R

        # 拆边界：先右后左
        if r_in != 0 and r_in != (r_b - r_a):
            _split_run_at(p, r_run, r_in)
        if l_in != 0 and l_in != (l_b - l_a):
            _split_run_at(p, l_run, l_in)
            l_run += 1

        mapping2 = _build_run_char_map(p)
        cover = []
        for run_idx, a, b in mapping2:
            if a >= s and b <= e and a != b:
                cover.append(run_idx)

        for ri in cover:
            p.runs[ri].font.color.rgb = BLUE


# def _find_all_occurrences(text: str, needle: str) -> List[Tuple[int, int]]:
#     out = []
#     if not needle:
#         return out
#     start = 0
#     while True:
#         idx = text.find(needle, start)
#         if idx == -1:
#             break
#         out.append((idx, idx + len(needle)))
#         start = idx + len(needle)
#     return out

def _find_all_occurrences(text: str, needle: str) -> List[Tuple[int, int]]:
    """
    默认按 needle 精确匹配返回范围；
    若 needle 命中位置前一字符是 '(' 或 '（'，且括号内从该位置到闭括号的长度<=15，
    则直接返回括号内该段范围（不做正则解析）。
    """
    out: List[Tuple[int, int]] = []
    if not needle:
        return out

    start = 0
    while True:
        idx = text.find(needle, start)
        if idx == -1:
            break

        # 命中前是左括号：尝试取括号内范围
        if idx > 0 and text[idx - 1] in ("(", "（"):
            r1 = text.find(")", idx)
            r2 = text.find("）", idx)
            rights = [i for i in (r1, r2) if i != -1]
            if rights:
                end_bracket = min(rights)
                inside_len = end_bracket - idx  # 从 Figure... 到闭括号前

                if inside_len <= 30:
                    out.append((idx, end_bracket))  # 不包含闭括号
                    start = end_bracket
                    continue

        # 否则按原逻辑
        out.append((idx, idx + len(needle)))
        start = idx + len(needle)

    return out


def _compute_legend_skip_paragraphs(paragraph_texts: List[str]) -> set:
    """
    计算哪些段落属于 legend 区（需要跳过标蓝）：
    - figure legend：从 "Figure X." 标题行开始，直到遇到：
      空行（一次即可） 或 遇到 SECTION_STOP_PAT（下一块开始）
    - table legend：同理
    返回需要跳过的段落索引集合
    """
    skip = set()
    i = 0
    n = len(paragraph_texts)
    while i < n:
        t = (paragraph_texts[i] or "").strip()
        is_fig = bool(FIG_LEGEND_TITLE_PAT.match(t))
        is_tab = bool(TAB_LEGEND_TITLE_PAT.match(t))

        if is_fig or is_tab:
            # 标题行也跳过
            skip.add(i)

            # # 向后跳过图注/表注段落
            # j = i + 1
            # while j < n:
            #     tj = (paragraph_texts[j] or "").strip()
            #     if tj == "":
            #         skip.add(j)
            #         break
            #     if SECTION_STOP_PAT.match(tj):
            #         # 遇到下一块（可能是下一张Figure标题行或References等），停止
            #         break
            #     # 仍属于 legend 区
            #     skip.add(j)
            #     j += 1

            # i = j
            # continue

        i += 1

    return skip


def blue_highlight_references_v3(
    docx_in: str,
    docx_out: str,
    normalized_fig_keys: List[str],
    normalized_table_keys: List[str],
    normalized_additional_keys: List[str],
) -> Dict[str, Any]:
    """
    只标蓝“正文引用处”：
    - 跳过 Figure/Table 的 legend 区段落（标题行 + 后续图注/表注段落）
    - 其余段落中匹配 Figure/Table/Additional file/PMID 并标蓝
    """
    doc = Document(docx_in)
    paragraph_texts = [p.text or "" for p in doc.paragraphs]
    skip_set = _compute_legend_skip_paragraphs(paragraph_texts)

    needles = []
    needles.extend(normalized_fig_keys or [])
    needles.extend(normalized_table_keys or [])
    needles.extend(normalized_additional_keys or [])
    needles = sorted(set(needles), key=len, reverse=True)

    hit_ranges = 0

    for p_i, p in enumerate(doc.paragraphs):
        if p_i in skip_set:
            continue

        text = p.text or ""
        if not text:
            continue

        ranges: List[Tuple[int, int]] = []

        # Figure/Table/Additional file
        for nd in needles:
            if nd in text:
                ranges.extend(_find_all_occurrences(text, nd))

        # PMID
        for m in PMID_GROUP_PAT.finditer(text):
            ranges.append((m.start(), m.end()))
        for m in PMID_SINGLE_PAT.finditer(text):
            ranges.append((m.start(), m.end()))

        if ranges:
            _colorize_paragraph_ranges(p, ranges)
            hit_ranges += len(ranges)

    doc.save(docx_out)
    return {"ok": True, "hit_ranges": hit_ranges, "skipped_paragraphs": len(skip_set),"final_docx_blue_path":docx_out}