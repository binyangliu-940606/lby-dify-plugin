import re
import requests
import json

from collections.abc import Generator
from io import BytesIO
from typing import List, Dict, Tuple, Any
from docx import Document
from docx.shared import Pt, RGBColor

from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage



class LbyToolsTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        #获取参数如文件url
        url = tool_parameters["old_file_url"]
        checked_table_list = json.loads(tool_parameters["checked_tables"])
        checked_figure_list = json.loads(tool_parameters["checked_figures"])
        verify_table_list = json.loads(tool_parameters["verify_table_list"])
        verify_figure_list = json.loads(tool_parameters["verify_figure_list"])
        file_name = tool_parameters["file_name"]
        
        out_bytes = verify_and_annotate_docx(
            url=url,
            checked_table_list=checked_table_list,
            checked_figure_list=checked_figure_list,
            verify_table_list=verify_table_list,
            verify_figure_list=verify_figure_list,
        )

        yield self.create_blob_message(
            blob=out_bytes,          # 直接传入 bytes，不需要 base64 编码
            meta={
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",   # 告诉 Dify 这是什么类型的文件
                "filename": file_name,
            },
        )

# ========================== 1) 文档读取 / 保存 ==========================

def load_docx_from_url(url: str) -> Document:
    """从 URL 下载 docx，并用 python-docx 打开成 Document 对象"""
    r = requests.get(url, timeout=60, allow_redirects=True)
    r.raise_for_status()
    return Document(BytesIO(r.content))

def docx_to_bytes(doc: Document) -> bytes:
    """将 python-docx 的 Document 保存为 bytes（用于上传、回传给调用方等）"""
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ========================== 2) 规范化 & 解析辅助函数 ==========================

def _clean_spaces(s: str) -> str:
    """把多余空白压缩成单空格，去掉首尾空白"""
    return re.sub(r"\s+", " ", s).strip()

def normalize_label(s: str) -> str:
    """
    将图表引用做“统一格式化”，减少同义写法导致的匹配不稳定：
    - Fig / Fig. / Figures -> Figure
    - Tables -> Table
    - Supplemental -> Supplementary
    - 保留原来是否带 Supplementary 前缀（不强行加/删）
    """
    s = _clean_spaces(s)

    # 统一 Fig 系列写法
    s = re.sub(r"\bFig(?:\.)?\b", "Figure", s, flags=re.I)
    s = re.sub(r"\bFigures\b", "Figure", s, flags=re.I)

    # 统一 Table 复数
    s = re.sub(r"\bTables\b", "Table", s, flags=re.I)

    # 统一 Supplemental -> Supplementary
    s = re.sub(r"\bSupplemental\b", "Supplementary", s, flags=re.I)
    s = re.sub(r"\bsupplementary\b", "Supplementary", s)

    # 开头大小写统一（Figure/Table）
    s = re.sub(r"^figure\b", "Figure", s, flags=re.I)
    s = re.sub(r"^table\b", "Table", s, flags=re.I)
    return s

def is_s(ref: str) -> bool:
    """
    判断是否是 S 编号（Supplementary 系列）：
    - Figure S1 / Figure S1A
    - Table S1
    S 编号在你的规则里“不要求顺序”，只校验存在性
    """
    return bool(re.search(r"\b[FT]igure\s+S\d+", ref)) or bool(re.search(r"\bTable\s+S\d+", ref))

def parse_main_index(ref: str) -> int | None:
    """
    解析主编号（仅对非 S 图表有效）：
    - Figure 12A -> 12
    - Figure 3 -> 3
    - Figure S2 -> None
    - Table S1 -> None
    """
    m = re.search(r"\b(?:Figure|Table)\s+(\d+)\b", ref)
    if m:
        return int(m.group(1))
    return None

def figure_table_kind(ref: str) -> str | None:
    """判断引用属于 Figure 还是 Table（用于做顺序校验）"""
    if re.search(r"\bFigure\b", ref):
        return "Figure"
    if re.search(r"\bTable\b", ref):
        return "Table"
    return None

def _extract_letter_suffix(ref: str) -> str | None:
    """
    提取小图字母后缀（不是必需，但保留以便后续扩展）：
    - Figure 2A -> 'A'
    - Figure 2 -> None
    - Figure S1B -> 'B'
    """
    m = re.search(r"\b(?:Figure|Table)\s+(?:S\d+|\d+)([A-Z])\b", ref)
    return m.group(1) if m else None

def _strip_non_ref_name(name: str) -> str:
    """
    用于“用于验证的 figure_list”里可能出现的键名清洗：
    - "Figure 2.pdf" -> "Figure 2"
    - "Figure 1.docx" -> "Figure 1"
    """
    name = _clean_spaces(name)
    name = re.sub(r"\.pdf$", "", name, flags=re.I)
    name = re.sub(r"\.docx?$", "", name, flags=re.I)
    return name


# ========================== 3) 将“用于验证的列表”展开成期望引用 ==========================

def build_expected_tables(verify_table_list: List[Dict[str, Any]]) -> Tuple[List[str], List[str]]:
    """
    将用于验证的 table_list（结构化输入）拆成两类期望列表：

    返回:
      (expected_s_tables, expected_numbered_tables_in_order)

    规则：
    - {"excel_name":"Table S1"} 属于 S 表：只要求“存在”，不要求顺序
    - {"all_tables":[...]} 里面的每个 Table 1/2/3...：
        既要求存在，也要求按 1->2->3... 的顺序出现
    """
    expected_s = []
    expected_num = []

    for item in verify_table_list:
        if "excel_name" in item and item["excel_name"]:
            expected_s.append(normalize_label(item["excel_name"]))

        if "all_tables" in item and isinstance(item["all_tables"], list):
            expected_num.extend([normalize_label(x) for x in item["all_tables"] if x])

    return expected_s, expected_num

def build_expected_figures(verify_figure_list: List[Dict[str, List[str]]]) -> Tuple[List[str], List[str]]:
    """
    将用于验证的 figure_list（结构化输入）拆成两类期望列表：

    返回:
      (expected_s_figs, expected_numbered_figs_in_order)

    输入形态样例：
      {"Figure 2": ["A","B","C"]}  -> 期望包含 Figure 2A, Figure 2B, Figure 2C
      {"Figure 4A": []}           -> 值为空时，只要求 key 本身存在（会清洗成 Figure 4A / Figure 4）

    规则：
    - S 编号（Figure S1 / S1A...）只校验“存在”，不校验顺序
    - 非 S 编号（Figure 1/2/3... 以及带字母 Figure 1A...）：
        校验存在性 + 顺序（按主编号 1->2->3...）
    """
    expected_s = []
    expected_num = []

    for d in verify_figure_list:
        if not isinstance(d, dict):
            continue

        for k, v in d.items():
            # 先把 key 的文件名后缀（.pdf 等）去掉
            k0 = normalize_label(_strip_non_ref_name(k))

            # 值为空：只要求主键存在（比如 Figure 2 或 Figure 11A）
            if not v:
                if is_s(k0):
                    expected_s.append(k0)
                else:
                    expected_num.append(k0)
            else:
                # 值非空：展开成小图（Figure 2A / Figure 2B ...）
                for letter in v:
                    if not letter:
                        continue
                    ref = normalize_label(f"{k0}{str(letter).strip()}")
                    if is_s(ref):
                        expected_s.append(ref)
                    else:
                        expected_num.append(ref)

    return expected_s, expected_num


# ========================== 4) 核心校验逻辑：缺失 + 顺序 ==========================

def _find_missing(expected: List[str], checked: List[str]) -> List[str]:
    """
    检查 expected 里的每个元素是否在 checked 中出现。
    注意：这里是严格字符串匹配（且不去重）。
    """
    missing = []
    for x in expected:
        if x not in checked:
            missing.append(x)
    return missing

def _order_issues_numbered(checked: List[str], kind: str) -> List[str]:
    """
    对“待检验列表 checked”做顺序粗检：
    - 只看非 S 的 Figure/Table
    - 主编号不能出现倒序：例如 ... Figure 2 ... Figure 1 ... => 报错
    该检查不依赖“用于验证列表”，用于发现明显倒序情况。
    """
    msgs = []
    last = None

    for i, ref in enumerate(checked):
        if figure_table_kind(ref) != kind:
            continue
        if is_s(ref):
            continue

        idx = parse_main_index(ref)
        if idx is None:
            continue

        if last is None:
            last = idx
            continue

        if idx < last:
            msgs.append(
                #f"{kind} 顺序错误（待检验列表）：第{i}项 {ref} 的主编号 {idx} 出现在 {last} 之后（应保持 1->2->3...）"
                f"{kind} 顺序错误（待检验列表）：{ref} 的主编号 {idx} 出现在 {last} 之后（应保持 1->2->3...）"
            )
        else:
            last = idx

    return msgs

def validate(
    checked_table_list: List[str],
    checked_figure_list: List[str],
    verify_table_list: List[Dict[str, Any]],
    verify_figure_list: List[Dict[str, List[str]]],
) -> List[str]:

    checked_tables = [normalize_label(x) for x in checked_table_list if x]
    checked_figures = [normalize_label(x) for x in checked_figure_list if x]

    exp_s_tables, exp_num_tables = build_expected_tables(verify_table_list)
    exp_s_figs, exp_num_figs = build_expected_figures(verify_figure_list)

    msgs = []

    # ---------- 4.1 缺失校验 ----------
    miss_tables_s = _find_missing(exp_s_tables, checked_tables)
    miss_tables_num = _find_missing(exp_num_tables, checked_tables)
    miss_figs_s = _find_missing(exp_s_figs, checked_figures)
    miss_figs_num = _find_missing(exp_num_figs, checked_figures)

    if miss_tables_s:
        msgs.append("缺失的 S 表（只校验存在性）: " + ", ".join(miss_tables_s))
    if miss_tables_num:
        msgs.append("缺失的普通表（需存在且有顺序）: " + ", ".join(miss_tables_num))
    if miss_figs_s:
        msgs.append("缺失的 S 图（只校验存在性）: " + ", ".join(miss_figs_s))
    if miss_figs_num:
        msgs.append("缺失的普通图（需存在且有顺序）: " + ", ".join(miss_figs_num))

    # ---------- 4.2 顺序校验（数字顺序） ----------
    msgs.extend(_order_issues_numbered(checked_figures, "Figure"))
    msgs.extend(_order_issues_numbered(checked_tables, "Table"))

    # msgs.extend(_order_issues_expected_sequence(exp_num_tables, checked_tables))
    # msgs.extend(_order_issues_expected_sequence(exp_num_figs, checked_figures))

    # ---------- 4.3 多余校验（新增） ----------
    allowed_tables = build_allowed_tables_set(verify_table_list)
    allowed_figures = build_allowed_figures_set(verify_figure_list)

    extra_tables = _find_extras(checked_tables, allowed_tables)
    extra_figures = _find_extras(checked_figures, allowed_figures)

    if extra_tables:
        msgs.append("多余的 Table（DOC 中出现但 ZIP 未定义）: " + ", ".join(extra_tables))
    if extra_figures:
        msgs.append("多余的 Figure（DOC 中出现但 ZIP 未定义）: " + ", ".join(extra_figures))

    # ---------- 4.4 字母顺序校验（新增） ----------
    msgs.extend(_letter_order_issues(checked_figures, "Figure"))
    msgs.extend(_letter_order_issues(checked_tables, "Table"))  # 若你的 Table 不会有 A/B，可保留也无妨

    return msgs

# ========================== 5) 将校验结果写到文档底部 ==========================

def append_validation_result_to_doc(doc: Document, url: str, messages: List[str]) -> None:
    """
    将校验报告写入 docx 的末尾（底部追加段落）
    """
    doc.add_paragraph("")  # 空行分隔
    p = doc.add_paragraph("=== 校验报告 Validation Report ===")
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 红色

    # doc.add_paragraph(f"文档 URL: {url}")

    if not messages:
        doc.add_paragraph("结果：通过（未发现缺失或顺序问题）")
        return

    doc.add_paragraph("")
    p = doc.add_paragraph(f"发现问题数：{len(messages)}")
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 红色

    i = 1
    for m in messages:
        p = doc.add_paragraph()  
        run = p.add_run(str(i) +". " +m)
        i += 1
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 红色


def build_allowed_tables_set(verify_table_list: List[Dict[str, Any]]) -> set[str]:
    """
    从 verify_table_list 推导出“允许出现的 Table 引用全集”（含 S 表 + 普通表）。
    """
    exp_s, exp_num = build_expected_tables(verify_table_list)
    return set(exp_s) | set(exp_num)


def build_allowed_figures_set(verify_figure_list: List[Dict[str, List[str]]]) -> set[str]:
    """
    从 verify_figure_list 推导出“允许出现的 Figure 引用全集”。

    注意：
    - 若某个 key 对应的 value 为空（如 {"Figure 4A": []}），允许集合里只包含 "Figure 4A"
    - 若 value 非空（如 {"Figure 2": ["A","B"]}），允许集合里包含 "Figure 2A","Figure 2B"...（不自动包含 "Figure 2"）
      （你现在的规则就是这样；如果你也希望 Figure 2 本身也允许，可再加一行 allow.add(k0)）
    """
    exp_s, exp_num = build_expected_figures(verify_figure_list)
    return set(exp_s) | set(exp_num)

def _find_extras(checked: List[str], allowed: set[str]) -> List[str]:
    """
    找出 checked 中“多余”的元素：即不在 allowed 集合内的引用。
    去重输出，但保持出现的先后顺序。
    """
    extras = []
    seen = set()
    for x in checked:
        if x not in allowed and x not in seen:
            extras.append(x)
            seen.add(x)
    return extras

def _letter_order_issues(checked: List[str], kind: str) -> List[str]:
    """
    校验同一主编号下的小图字母顺序：A->B->C...
    例：Figure 1C 出现在 Figure 1B 之前 => 报错
    """
    msgs = []

    # last_letter_ord[(kind, main_index)] = ord_value
    last_letter_ord: Dict[int, int] = {}

    for ref in checked:
        if figure_table_kind(ref) != kind:
            continue
        if is_s(ref):
            continue

        main_idx = parse_main_index(ref)
        if main_idx is None:
            continue

        letter = _extract_letter_suffix(ref)
        if not letter:
            continue

        cur = ord(letter.upper())

        if main_idx not in last_letter_ord:
            last_letter_ord[main_idx] = cur
            continue

        prev = last_letter_ord[main_idx]
        if cur < prev:
            prev_letter = chr(prev)
            msgs.append(
                f"{kind} 字母顺序错误：同一主编号 {main_idx} 下，{ref}（{letter}）出现在 {prev_letter} 之后（应 A->B->C...）"
            )
        else:
            last_letter_ord[main_idx] = cur

    return msgs

# ========================== 6) 总入口：校验 + 写入 + 输出 bytes ==========================

def verify_and_annotate_docx(
    url: str,
    checked_table_list: List[str],
    checked_figure_list: List[str],
    verify_table_list: List[Dict[str, Any]],
    verify_figure_list: List[Dict[str, List[str]]],
) -> bytes:
    """
    入参（五个参数）：
      1) url：docx 文档地址
      2) checked_table_list：待检验的 table 引用列表（你抽取出来的结果）
      3) checked_figure_list：待检验的 figure 引用列表（你抽取出来的结果）
      4) verify_table_list：用于验证的 table 基准列表（excel_name / all_tables）
      5) verify_figure_list：用于验证的 figure 基准列表（大图 -> 小图字母列表）

    出参：
      - 将“校验报告”写入到原 docx 底部后，返回该 docx 的 bytes 字节流
    """
    doc = load_docx_from_url(url)
    messages = validate(
        checked_table_list=checked_table_list,
        checked_figure_list=checked_figure_list,
        verify_table_list=verify_table_list,
        verify_figure_list=verify_figure_list,
    )
    append_validation_result_to_doc(doc, url, messages)
    return docx_to_bytes(doc)

