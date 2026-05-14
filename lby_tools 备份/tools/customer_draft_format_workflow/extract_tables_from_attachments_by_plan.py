import os
import pandas as pd
import json

from collections.abc import Generator
from docx import Document
from typing import Dict, Any, List, Optional

from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage


class LbyToolsTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        #获取参数如文件url
        payload_json = json.loads(tool_parameters["payload_json"])
        table_extract_plan = payload_json["table_extract_plan"]
        table_raw_dir = payload_json["table_raw_dir"]

        yield self.create_json_message(
            extract_tables_from_attachments_by_plan(table_extract_plan, table_raw_dir)
        )


def _extract_docx_nearest_table(docx_path: str, keyword: str) -> Dict[str, Any]:
    """
    docx 抽表（可运行版，尽量稳定）：
    - 先在段落中找到包含 keyword 的位置（用于“计划执行”的基本校验）
    - 但 python-docx 无混合遍历 w:p/w:tbl，无法100%保证“最近表”
    - 这里采取折中：如果文档只有一张表就取第一张；多表时仍取第一张（建议后续升级为OOXML混合遍历）
    """
    doc = Document(docx_path)

    # 简单确认 keyword 是否出现（不出现也可能仍有表，仍尝试）
    found = False
    if keyword:
        for p in doc.paragraphs[:300]:
            if keyword in (p.text or ""):
                found = True
                break

    if not doc.tables:
        return {}

    # 折中策略：取第一张表
    tbl = doc.tables[0]
    grid = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
    return {"grid": grid, "found_keyword": found}


def _extract_xlsx_by_sheet_name(path: str, sheet_name: str) -> Dict[str, Any]:
    df = pd.read_excel(path, sheet_name=sheet_name, header=None)
    grid = df.fillna("").astype(str).values.tolist()
    return {"grid": grid, "sheet_name": sheet_name}


def _extract_xlsx_by_cell_search_then_block(path: str, sheet_name: Optional[str], keyword: str) -> Dict[str, Any]:
    """
    在 xlsx 中搜 keyword，找到后抽取一个“块”作为表：
    - 从命中行下一行开始
    - 直到连续两行全空为止
    """
    xls = pd.ExcelFile(path)
    sheets = [sheet_name] if sheet_name else xls.sheet_names

    for sh in sheets:
        df = pd.read_excel(path, sheet_name=sh, header=None)

        found = None
        # 搜索 keyword
        if keyword:
            for r in range(df.shape[0]):
                for c in range(df.shape[1]):
                    v = df.iat[r, c]
                    if isinstance(v, str) and keyword in v:
                        found = (r, c)
                        break
                if found:
                    break

        if not found:
            continue

        start_r = found[0] + 1
        end_r = start_r
        empty_streak = 0

        for r in range(start_r, df.shape[0]):
            row_vals = df.iloc[r].fillna("").astype(str).tolist()
            if all(v.strip() == "" for v in row_vals):
                empty_streak += 1
                if empty_streak >= 2:
                    break
            else:
                empty_streak = 0
                end_r = r

        block = df.iloc[start_r:end_r + 1].fillna("").astype(str)
        grid = block.values.tolist()

        # 清理全空行
        grid = [row for row in grid if not all((x or "").strip() == "" for x in row)]
        if grid:
            return {"grid": grid, "sheet_name": sh, "keyword_cell": {"row": found[0], "col": found[1]}}

    return {}


def _extract_csv_all(path: str) -> Dict[str, Any]:
    """
    csv 类型差异较大，这里先采用保守方案：整表读入作为 grid
    （后续如需按 keyword 截块，可仿照 xlsx 的 block 抽取逻辑）
    """
    df = pd.read_csv(path, header=None, encoding="utf-8", engine="python")
    grid = df.fillna("").astype(str).values.tolist()
    return {"grid": grid}


def extract_tables_from_attachments_by_plan(table_extract_plan: List[Dict[str, Any]], table_raw_dir: str) -> Dict[str, Any]:
    """
    适配 Node8 输出结构：
    table_extract_plan = [
      {"table_id":"Table 1", "table_info": {"file_path": "...", "file_type":"xlsx", "sheet_name":"...", "rule":"...", "keyword":"..."}},
      ...
    ]
    """
    tables_from_attach: Dict[str, Any] = {}
    missing: List[Dict[str, Any]] = []

#  "table_extract_plan": [
#     {
#       "table_id": "Table S6",
#       "table_info": {
#         "file_path": "C:\\Users\\ADMINI~1\\AppData\\Local\\Temp\\paper_work_4nnc0b7c\\table_raw\\Tables\\Table S6 Tumor_GO_KEGG.xlsx",
#         "file_type": "xlsx",
#         "sheet_name": null,
#         "rule": "cell_search_then_block",
#         "keyword": "Table S6"
#       }
#     },

    for item in table_extract_plan or []:
        table_id = item.get("table_id")
        info = item.get("table_info") or {}

        file_path = info.get("file_path")
        file_type = info.get("file_type")
        rule = info.get("rule")
        keyword = info.get("keyword") or table_id
        sheet_name = info.get("sheet_name")  # 可能为空/None

        if not table_id:
            continue

        if not file_path or not os.path.exists(file_path):
            missing.append({"table": table_id, "reason": "file_not_found", "file_path": file_path})
            continue

        try:
            if file_type == "docx":
                if rule != "docx_nearest_table":
                    # docx 目前只支持该 rule（按你的schema enum）
                    missing.append({"table": table_id, "reason": "unsupported_rule_for_docx", "rule": rule})
                    continue

                r = _extract_docx_nearest_table(file_path, keyword)
                if r.get("grid"):
                    tables_from_attach[table_id] = {
                        "grid": r["grid"],
                        "source": {"file": file_path, "type": "docx", "rule": rule, "keyword": keyword, "found_keyword": r.get("found_keyword")}
                    }
                else:
                    missing.append({"table": table_id, "reason": "no_table_found_in_docx", "file_path": file_path})

            elif file_type == "xlsx":
                if rule == "sheet_name_match":
                    if not sheet_name:
                        missing.append({"table": table_id, "reason": "sheet_name_required", "file_path": file_path})
                        continue
                    r = _extract_xlsx_by_sheet_name(file_path, sheet_name)
                    if r.get("grid"):
                        tables_from_attach[table_id] = {
                            "grid": r["grid"],
                            "source": {"file": file_path, "type": "xlsx", "rule": rule, "sheet": r.get("sheet_name")}
                        }
                    else:
                        missing.append({"table": table_id, "reason": "empty_sheet", "sheet": sheet_name})

                elif rule == "cell_search_then_block":
                    r = _extract_xlsx_by_cell_search_then_block(file_path, sheet_name, keyword)
                    if r.get("grid"):
                        tables_from_attach[table_id] = {
                            "grid": r["grid"],
                            "source": {"file": file_path, "type": "xlsx", "rule": rule, "sheet": r.get("sheet_name"), "keyword": keyword, "hit": r.get("keyword_cell")}
                        }
                    else:
                        missing.append({"table": table_id, "reason": "keyword_not_found_in_xlsx", "keyword": keyword, "file_path": file_path})

                else:
                    missing.append({"table": table_id, "reason": "unsupported_rule_for_xlsx", "rule": rule})

            elif file_type == "csv":
                # schema enum 允许 csv，但 rule 里没专门的 csv 规则，这里做兼容：
                # - 若 rule 是 cell_search_then_block：暂时也先整表读入（你后续要按keyword截块我再帮你增强）
                r = _extract_csv_all(file_path)
                if r.get("grid"):
                    tables_from_attach[table_id] = {
                        "grid": r["grid"],
                        "source": {"file": file_path, "type": "csv", "rule": rule, "keyword": keyword}
                    }
                else:
                    missing.append({"table": table_id, "reason": "empty_csv", "file_path": file_path})

            else:
                missing.append({"table": table_id, "reason": "unsupported_file_type", "file_type": file_type})

        except Exception as e:
            missing.append({"table": table_id, "reason": str(e), "file_path": file_path, "file_type": file_type, "rule": rule})

    return {"tables_from_attach": tables_from_attach, "table_missing_report": missing}