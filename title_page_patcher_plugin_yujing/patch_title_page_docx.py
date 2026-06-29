from __future__ import annotations

import json
import mimetypes
import os
import re
import shutil
import tempfile
from collections.abc import Generator
from copy import deepcopy
from pathlib import Path
from typing import Any
from urllib.parse import urljoin

import requests
from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph


EMAIL_RE = re.compile(r"([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,})")
BACK_MATTER_HEADINGS = [
    "ethics statement",
    "ethical statement",
    "clinical trial registration",
    "consent to publish",
    "funding",
    "acknowledgment",
    "acknowledgement",
    "conflict of interest",
    "data availability",
    "author contributions",
]
CONTACT_LABEL_RE = re.compile(r"(投稿邮箱作者|投稿邮箱|姓名|邮箱|author|name|email)\s*[:：]\s*", re.IGNORECASE)


class LbyToolsTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage, None, None]:
        source_file_url = tool_parameters.get("source_file_url")
        submission_contact = str(tool_parameters.get("submission_contact") or "").strip()
        patch_json_raw = tool_parameters.get("patch_json")
        output_filename = str(tool_parameters.get("output_filename") or "Title_page_patched.docx").strip() or "Title_page_patched.docx"

        if not source_file_url:
            raise ValueError("source_file_url is required")

        with tempfile.TemporaryDirectory(prefix="title_page_patch_") as temp_dir:
            temp_root = Path(temp_dir)
            source_path = temp_root / "source.docx"
            output_path = temp_root / output_filename

            self._download_source(source_file_url, source_path)
            shutil.copyfile(source_path, output_path)

            qc_report: dict[str, Any] = {
                "status": "SUCCESS",
                "input_output": {
                    "output_filename": output_filename,
                },
                "submission_contact_parse": {},
                "patch_json_parse": {"ok": True, "warnings": []},
                "actions": [],
                "warnings": [],
                "needs_human_check": [],
            }

            parsed_contact = self._parse_submission_contact(submission_contact)
            qc_report["submission_contact_parse"] = parsed_contact
            if not parsed_contact["parse_contact_success"]:
                qc_report["needs_human_check"].append("submission_contact_parse_failed")

            patch_data = self._parse_patch_json(patch_json_raw, qc_report)
            document = Document(str(output_path))

            self._apply_email_patch(document, patch_data, qc_report)
            self._apply_co_corresponding_patch(document, patch_data, parsed_contact, qc_report)
            self._apply_back_matter_patch(document, patch_data, qc_report)

            if qc_report["needs_human_check"]:
                qc_report["status"] = "NEED_HUMAN_CHECK"

            document.save(str(output_path))

            file_meta = {
                "filename": output_filename,
                "size_bytes": output_path.stat().st_size,
                "mime_type": mimetypes.guess_type(output_filename)[0]
                or "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            }

            yield self.create_blob_message(
                blob=output_path.read_bytes(),
                meta={"mime_type": file_meta["mime_type"], "file_name": output_filename},
            )
            yield self.create_variable_message("files", [file_meta])
            yield self.create_variable_message("qc_report", json.dumps(qc_report, ensure_ascii=False))
            yield self.create_json_message({"files": [file_meta], "qc_report": qc_report})
            yield self.create_text_message(f"Patched title page saved as {output_filename}")

    def _parse_patch_json(self, raw_patch: Any, qc_report: dict[str, Any]) -> dict[str, Any]:
        if isinstance(raw_patch, dict):
            return raw_patch
        if raw_patch is None:
            qc_report["patch_json_parse"]["ok"] = False
            qc_report["patch_json_parse"]["warnings"].append("patch_json_missing")
            qc_report["needs_human_check"].append("patch_json_missing")
            return {}
        try:
            return json.loads(str(raw_patch))
        except Exception as exc:  # noqa: BLE001
            qc_report["patch_json_parse"]["ok"] = False
            qc_report["patch_json_parse"]["warnings"].append(f"patch_json_invalid: {exc}")
            qc_report["needs_human_check"].append("patch_json_invalid")
            return {}

    def _download_source(self, source_input: Any, target_path: Path) -> None:
        if isinstance(source_input, str):
            if os.path.exists(source_input):
                shutil.copyfile(source_input, target_path)
                return
            if source_input.startswith(("http://", "https://")):
                response = requests.get(source_input, timeout=120)
                response.raise_for_status()
                target_path.write_bytes(response.content)
                return

        if isinstance(source_input, dict):
            local_path = source_input.get("local_path") or source_input.get("path")
            if local_path and os.path.exists(str(local_path)):
                shutil.copyfile(str(local_path), target_path)
                return

            url = self._resolve_file_url(source_input)
            if not url:
                raise ValueError("source_file_url did not contain a downloadable URL.")
            response = requests.get(url, timeout=120)
            response.raise_for_status()
            target_path.write_bytes(response.content)
            return

        raise ValueError(f"Unsupported source_file_url payload: {type(source_input)!r}")

    def _resolve_file_url(self, file_input: dict[str, Any]) -> str:
        candidates = [
            file_input.get("url"),
            file_input.get("download_url"),
            file_input.get("remote_url"),
            file_input.get("file_url"),
        ]
        url = next((str(item).strip() for item in candidates if item), "")
        if not url:
            return ""
        if url.startswith(("http://", "https://")):
            return url

        base_candidates = [
            os.getenv("FILES_URL", "").strip(),
            os.getenv("CONSOLE_API_URL", "").strip(),
            os.getenv("DIFY_API_URL", "").strip(),
        ]
        base_url = next((item for item in base_candidates if item), "")
        if not base_url:
            return url
        if not base_url.endswith("/"):
            base_url += "/"
        return urljoin(base_url, url.lstrip("/"))

    def _parse_submission_contact(self, submission_contact: str) -> dict[str, Any]:
        warnings: list[str] = []
        emails = EMAIL_RE.findall(submission_contact or "")
        parsed_email = emails[0] if emails else ""

        name_text = submission_contact or ""
        name_text = re.sub(r"mailto:\s*", " ", name_text, flags=re.IGNORECASE)
        for email in emails:
            name_text = name_text.replace(email, " ")
        name_text = CONTACT_LABEL_RE.sub(" ", name_text)
        name_text = re.sub(r"[\[\]\(\)<>【】]", " ", name_text)
        name_text = re.sub(r"[，,；;|/]+", " ", name_text)
        name_text = re.sub(r"\s+", " ", name_text).strip(" .,:;")

        if not parsed_email:
            warnings.append("email_not_found")
        if not name_text:
            warnings.append("author_name_not_found")

        return {
            "submission_contact": submission_contact,
            "parsed_submission_author_name": name_text,
            "parsed_submission_email": parsed_email,
            "parse_contact_success": bool(parsed_email and name_text),
            "warnings": warnings,
        }

    def _apply_email_patch(self, document: Document, patch_data: dict[str, Any], qc_report: dict[str, Any]) -> None:
        email_patch = patch_data.get("email_patch") or {}
        if not email_patch.get("enabled"):
            qc_report["actions"].append({"action": "email_patch", "status": "skipped"})
            return

        target_person = str(email_patch.get("target_person") or "").strip()
        old_email = str(email_patch.get("old_email") or "").strip()
        new_email = str(email_patch.get("new_email") or "").strip()
        if not new_email:
            qc_report["actions"].append({"action": "email_patch", "status": "skipped", "reason": "new_email_missing"})
            qc_report["needs_human_check"].append("email_patch_new_email_missing")
            return

        replaced = False
        for paragraph in document.paragraphs:
            para_text = paragraph.text or ""
            para_norm = self._normalize_name(para_text)
            if target_person and target_person not in para_text and self._normalize_name(target_person) not in para_norm:
                continue

            if old_email and old_email in para_text:
                if self._looks_like_correspondence_paragraph(para_text):
                    self._rebuild_existing_correspondence_paragraph(paragraph, new_email)
                else:
                    self._replace_paragraph_text(paragraph, para_text.replace(old_email, new_email, 1))
                replaced = True
                break

            found_emails = EMAIL_RE.findall(para_text)
            if found_emails and self._looks_like_correspondence_paragraph(para_text):
                self._rebuild_existing_correspondence_paragraph(paragraph, new_email)
                replaced = True
                break

        if not replaced:
            for paragraph in document.paragraphs:
                para_text = paragraph.text or ""
                if old_email and old_email in para_text:
                    self._replace_paragraph_text(paragraph, para_text.replace(old_email, new_email, 1))
                    replaced = True
                    break

        qc_report["actions"].append(
            {
                "action": "email_patch",
                "status": "applied" if replaced else "skipped",
                "target_person": target_person,
                "old_email": old_email,
                "new_email": new_email,
            }
        )
        if not replaced:
            qc_report["needs_human_check"].append("email_patch_target_not_found")

    def _apply_co_corresponding_patch(
        self,
        document: Document,
        patch_data: dict[str, Any],
        parsed_contact: dict[str, Any],
        qc_report: dict[str, Any],
    ) -> None:
        patch = patch_data.get("co_corresponding_author_patch") or {}
        if not patch.get("enabled"):
            qc_report["actions"].append({"action": "co_corresponding_author_patch", "status": "skipped"})
            return

        author_name = str(patch.get("author_name") or parsed_contact.get("parsed_submission_author_name") or "").strip()
        email = str(patch.get("email") or parsed_contact.get("parsed_submission_email") or "").strip()
        affiliation_text = str(patch.get("affiliation_text") or "").strip()
        if not author_name:
            qc_report["actions"].append(
                {"action": "co_corresponding_author_patch", "status": "skipped", "reason": "author_name_missing"}
            )
            qc_report["needs_human_check"].append("co_corresponding_author_name_missing")
            return

        author_line_found = self._add_corresponding_mark_to_author_line(document, author_name, qc_report)
        correspondence_added = self._append_correspondence_paragraph(
            document=document,
            author_name=author_name,
            email=email,
            affiliation_text=affiliation_text,
            qc_report=qc_report,
        )

        qc_report["actions"].append(
            {
                "action": "co_corresponding_author_patch",
                "status": "applied" if (author_line_found or correspondence_added) else "skipped",
                "author_name": author_name,
                "email": email,
                "affiliation_text": affiliation_text,
            }
        )
        if not author_line_found:
            qc_report["needs_human_check"].append("author_line_for_co_corresponding_not_found")
        if not correspondence_added:
            qc_report["needs_human_check"].append("correspondence_anchor_not_found")

    def _apply_back_matter_patch(self, document: Document, patch_data: dict[str, Any], qc_report: dict[str, Any]) -> None:
        patch = patch_data.get("back_matter_patch") or {}
        sections = patch.get("sections") or []
        built_sections: list[tuple[str, str]] = []
        for item in sections:
            heading = str((item or {}).get("heading") or "").strip()
            body = str((item or {}).get("body") or "").strip()
            if heading:
                body = self._normalize_english_identifiers(body)
                if heading.lower() == "ethical statement" and "xxx" in body.lower():
                    resolved = self._resolve_committee_placeholder(document, body)
                    body = resolved
                built_sections.append((heading, body))

        if not built_sections:
            qc_report["actions"].append({"action": "back_matter_patch", "status": "skipped", "reason": "sections_missing"})
            qc_report["needs_human_check"].append("back_matter_sections_missing")
            return

        paragraphs = document.paragraphs
        start_idx = self._find_back_matter_start(paragraphs)
        anchor_paragraph: Paragraph | None = None
        removed_mode = "append"

        heading_template = None
        body_template = None
        if start_idx is not None:
            heading_template = paragraphs[start_idx]
            if start_idx + 1 < len(paragraphs):
                body_template = paragraphs[start_idx + 1]
            anchor_paragraph = paragraphs[start_idx - 1] if start_idx > 0 else None
            for paragraph in list(document.paragraphs[start_idx:]):
                self._delete_paragraph(paragraph)
            removed_mode = "replace_from_existing_back_matter"
        else:
            anchor_paragraph = document.paragraphs[-1] if document.paragraphs else None
            heading_template = anchor_paragraph
            body_template = anchor_paragraph

        if anchor_paragraph is None:
            new_para = document.add_paragraph()
            current_paragraph = new_para
        else:
            current_paragraph = anchor_paragraph

        inserted_headings: list[str] = []
        for heading, body in built_sections:
            heading_para = self._insert_paragraph_after(
                current_paragraph,
                heading,
                style=heading_template.style if heading_template and heading_template.style else None,
                template_paragraph=heading_template,
            )
            self._clear_paragraph_runs(heading_para)
            heading_run = heading_para.add_run(heading)
            self._apply_default_run_format(heading_run, bold=True)
            body_para = self._insert_paragraph_after(
                heading_para,
                body or "None.",
                style=body_template.style if body_template and body_template.style else None,
                template_paragraph=body_template or heading_template,
            )
            self._clear_paragraph_runs(body_para)
            body_run = body_para.add_run(body or "None.")
            self._apply_default_run_format(body_run, bold=False)
            spacer_para = self._insert_paragraph_after(
                body_para,
                "",
                style=body_template.style if body_template and body_template.style else None,
                template_paragraph=body_template or heading_template,
            )
            self._clear_paragraph_runs(spacer_para)
            current_paragraph = spacer_para
            inserted_headings.append(heading)

        qc_report["actions"].append(
            {
                "action": "back_matter_patch",
                "status": "applied",
                "mode": removed_mode,
                "headings": inserted_headings,
            }
        )

    def _find_back_matter_start(self, paragraphs: list[Paragraph]) -> int | None:
        for idx, paragraph in enumerate(paragraphs):
            text = (paragraph.text or "").strip()
            if not text:
                continue
            lowered = self._heading_key(text)
            if lowered in BACK_MATTER_HEADINGS:
                return idx
        return None

    def _add_corresponding_mark_to_author_line(self, document: Document, author_name: str, qc_report: dict[str, Any]) -> bool:
        target = self._normalize_name(author_name)
        for paragraph in document.paragraphs:
            para_text = paragraph.text or ""
            if not para_text.strip():
                continue
            lowered = para_text.lower()
            if self._looks_like_stop_section(lowered):
                break
            if target and target in self._normalize_name(para_text):
                if re.search(re.escape(author_name) + r"\s*[*#†]", para_text):
                    return True
                new_text = re.sub(re.escape(author_name) + r"(\s*[0-9]+[*#†]*)", author_name + r"*\1", para_text, count=1)
                if new_text == para_text:
                    new_text = re.sub(re.escape(author_name), author_name + "*", para_text, count=1)
                if new_text != para_text:
                    self._rebuild_author_paragraph(paragraph, new_text)
                    return True
        qc_report["warnings"].append(f"author_line_not_safely_located_for: {author_name}")
        return False

    def _append_correspondence_paragraph(
        self,
        document: Document,
        author_name: str,
        email: str,
        affiliation_text: str,
        qc_report: dict[str, Any],
    ) -> bool:
        correspondence_indexes = [
            idx for idx, paragraph in enumerate(document.paragraphs) if self._looks_like_correspondence_paragraph(paragraph.text or "")
        ]
        if not correspondence_indexes:
            qc_report["warnings"].append("correspondence_section_not_found")
            return False

        start_idx = correspondence_indexes[-1]
        last_idx = start_idx
        for idx in range(start_idx + 1, len(document.paragraphs)):
            text = (document.paragraphs[idx].text or "").strip()
            lowered = self._heading_key(text)
            if lowered in BACK_MATTER_HEADINGS:
                break
            if not text:
                continue
            last_idx = idx

        anchor = document.paragraphs[last_idx]
        corr_para = self._insert_paragraph_after(anchor, "", style=anchor.style if anchor.style else None, template_paragraph=anchor)
        self._build_correspondence_paragraph(
            corr_para,
            author_name=author_name,
            affiliation_text=affiliation_text,
            email=email,
            tel="",
        )
        return True

    def _replace_paragraph_text(self, paragraph: Paragraph, new_text: str) -> None:
        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run.text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(new_text)

    def _insert_paragraph_after(
        self,
        paragraph: Paragraph,
        text: str,
        style: Any = None,
        template_paragraph: Paragraph | None = None,
    ) -> Paragraph:
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if style is not None:
            try:
                new_para.style = style
            except Exception:  # noqa: BLE001
                pass
        if template_paragraph is not None:
            self._copy_paragraph_format(template_paragraph, new_para)
        if text:
            run = new_para.add_run(text)
            if template_paragraph is not None and template_paragraph.runs:
                self._copy_run_format(template_paragraph.runs[0], run)
        return new_para

    def _delete_paragraph(self, paragraph: Paragraph) -> None:
        p = paragraph._element
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)
        paragraph._p = paragraph._element = None

    def _looks_like_correspondence_paragraph(self, text: str) -> bool:
        lowered = (text or "").lower()
        return "correspondence to" in lowered or lowered.startswith("correspondence:")

    def _looks_like_stop_section(self, lowered_text: str) -> bool:
        stop_keywords = [
            "correspondence",
            "funding",
            "ethical statement",
            "clinical trial registration",
            "consent to publish",
            "acknowledgment",
            "acknowledgement",
            "conflict of interest",
            "data availability",
            "author contributions",
        ]
        return any(keyword in lowered_text for keyword in stop_keywords)

    def _heading_key(self, text: str) -> str:
        cleaned = (text or "").strip()
        if not cleaned:
            return ""
        first_line = cleaned.splitlines()[0].strip().lower()
        first_line = first_line.rstrip(":")
        return first_line

    def _copy_paragraph_format(self, source: Paragraph, target: Paragraph) -> None:
        src = source.paragraph_format
        dst = target.paragraph_format
        dst.left_indent = src.left_indent
        dst.right_indent = src.right_indent
        dst.first_line_indent = src.first_line_indent
        dst.keep_together = src.keep_together
        dst.keep_with_next = src.keep_with_next
        dst.page_break_before = src.page_break_before
        dst.widow_control = src.widow_control
        dst.space_before = src.space_before
        dst.space_after = src.space_after
        dst.line_spacing = src.line_spacing
        dst.line_spacing_rule = src.line_spacing_rule
        target.alignment = source.alignment

    def _copy_run_format(self, source_run, target_run) -> None:
        sf = source_run.font
        tf = target_run.font
        tf.name = sf.name
        tf.size = sf.size
        tf.bold = sf.bold
        tf.italic = sf.italic
        tf.underline = sf.underline
        tf.superscript = sf.superscript
        tf.subscript = sf.subscript
        tf.all_caps = sf.all_caps
        tf.small_caps = sf.small_caps
        tf.strike = sf.strike
        tf.double_strike = sf.double_strike
        tf.highlight_color = sf.highlight_color
        if sf.color.rgb is not None:
            tf.color.rgb = sf.color.rgb

    def _rebuild_author_paragraph(self, paragraph: Paragraph, text: str) -> None:
        template_runs = list(paragraph.runs)
        normal_template = next((run for run in template_runs if not run.font.superscript), template_runs[0] if template_runs else None)
        superscript_template = next((run for run in template_runs if run.font.superscript), normal_template)

        self._clear_paragraph_runs(paragraph)

        parts = [part.strip() for part in text.split(",")]
        for idx, part in enumerate(parts):
            match = re.match(r"^(.*?)([0-9]+[*#†]*|[*#†]+[0-9]*|[0-9]+[*#†]+)$", part)
            if match:
                name_part = match.group(1).strip()
                marker_part = match.group(2).strip()
            else:
                name_part = part.strip()
                marker_part = ""

            if idx > 0:
                run = paragraph.add_run(", ")
                if normal_template is not None:
                    self._copy_run_format(normal_template, run)
            if name_part:
                run = paragraph.add_run(name_part)
                if normal_template is not None:
                    self._copy_run_format(normal_template, run)
            if marker_part:
                run = paragraph.add_run(marker_part)
                if superscript_template is not None:
                    self._copy_run_format(superscript_template, run)
                else:
                    run.font.superscript = True
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

    def _resolve_committee_placeholder(self, document: Document, body: str) -> str:
        org = self._extract_corresponding_org(document)
        if not org:
            return body
        return re.sub(r"Animal Ethics Committee of XXX", f"Animal Ethics Committee of {org}", body, flags=re.IGNORECASE)

    def _normalize_english_identifiers(self, text: str) -> str:
        value = text or ""
        value = value.replace("【", "[").replace("】", "]")
        value = re.sub(
            r"\(\s*No\.\s*医伦第\s*\[\s*([0-9A-Za-z._-]+)\s*\]\s*\)",
            r"(Ethics approval document No. \1)",
            value,
            flags=re.IGNORECASE,
        )
        value = re.sub(
            r"\(\s*No\.\s*医伦第\s*([0-9A-Za-z._-]+)\s*\)",
            r"(Ethics approval document No. \1)",
            value,
            flags=re.IGNORECASE,
        )
        value = re.sub(
            r"医伦第\s*\[\s*([0-9A-Za-z._-]+)\s*\]",
            r"Ethics approval document No. \1",
            value,
            flags=re.IGNORECASE,
        )
        value = re.sub(
            r"医伦第\s*([0-9A-Za-z._-]+)",
            r"Ethics approval document No. \1",
            value,
            flags=re.IGNORECASE,
        )
        value = re.sub(
            r"\(\s*No\.\s*[?]{2,}\s*([0-9A-Za-z._-]+)\s*[?]+\s*\)",
            r"(Ethics approval document No. \1)",
            value,
            flags=re.IGNORECASE,
        )
        value = re.sub(
            r"No\.\s*[?]{2,}\s*([0-9A-Za-z._-]+)\s*[?]+",
            r"Ethics approval document No. \1",
            value,
            flags=re.IGNORECASE,
        )
        value = re.sub(
            r"[?]{2,}\s*([0-9A-Za-z._-]+)\s*[?]+",
            r"Ethics approval document No. \1",
            value,
            flags=re.IGNORECASE,
        )
        value = re.sub(r"\(\s*No\.\s*\[\s*([0-9A-Za-z._-]+)\s*\]\s*\)", r"(No. \1)", value)
        value = re.sub(r"\(\s*No\.\s*([0-9A-Za-z._-]+)\s*\)", r"(No. \1)", value)
        return value

    def _extract_corresponding_org(self, document: Document) -> str:
        for paragraph in document.paragraphs:
            text = (paragraph.text or "").strip()
            if "Correspondence to:" in text:
                m = re.search(r"Correspondence to:\s*[^,]+,\s*(.+)", text)
                if m:
                    tail = m.group(1).strip()
                    tail = re.split(r",\s*No\.", tail, maxsplit=1)[0].strip()
                    if tail.lower().startswith("department of ") and "," in tail:
                        tail = tail.split(",", 1)[1].strip()
                    return tail
        return ""

    def _clear_paragraph_runs(self, paragraph: Paragraph) -> None:
        for run in list(paragraph.runs):
            run._element.getparent().remove(run._element)

    def _rebuild_existing_correspondence_paragraph(self, paragraph: Paragraph, new_email: str) -> None:
        lines = (paragraph.text or "").splitlines()
        first = lines[0].strip() if lines else ""
        second = lines[1].strip() if len(lines) > 1 else ""
        third = lines[2].strip() if len(lines) > 2 else ""

        first = first.lstrip("*").strip()
        first = re.sub(r"^Correspondence to:\s*", "", first, flags=re.IGNORECASE)
        name = ""
        affiliation = ""
        if "," in first:
            name, affiliation = first.split(",", 1)
            name = name.strip()
            affiliation = affiliation.strip()
        else:
            name = first.strip()

        tel = ""
        if third:
            tel = re.sub(r"^Tel\.\s*:\s*", "", third, flags=re.IGNORECASE).strip()
        self._build_correspondence_paragraph(
            paragraph,
            author_name=name,
            affiliation_text=affiliation,
            email=new_email,
            tel=tel,
        )

    def _build_correspondence_paragraph(
        self,
        paragraph: Paragraph,
        *,
        author_name: str,
        affiliation_text: str,
        email: str,
        tel: str,
    ) -> None:
        self._clear_paragraph_runs(paragraph)
        mark_run = paragraph.add_run("*")
        self._apply_default_run_format(mark_run, bold=True, superscript=True, color=RGBColor(0x00, 0x00, 0xFF))
        space_run = paragraph.add_run(" ")
        self._apply_default_run_format(space_run, bold=False, color=RGBColor(0x00, 0x00, 0x00))
        label_run = paragraph.add_run("Correspondence to:")
        self._apply_default_run_format(label_run, bold=True, color=RGBColor(0x00, 0x00, 0x00))
        space2 = paragraph.add_run(" ")
        self._apply_default_run_format(space2, bold=False, color=RGBColor(0x00, 0x00, 0x00))
        name_run = paragraph.add_run(author_name)
        self._apply_default_run_format(name_run, bold=True, color=RGBColor(0x00, 0x00, 0x00))
        if affiliation_text:
            aff_run = paragraph.add_run(", " + affiliation_text.rstrip("."))
            self._apply_default_run_format(aff_run, bold=False, color=RGBColor(0x00, 0x00, 0x00))
        if email:
            newline = paragraph.add_run("\n")
            self._apply_default_run_format(newline, bold=False, color=RGBColor(0x00, 0x00, 0x00))
            email_label = paragraph.add_run("E-mail:")
            self._apply_default_run_format(email_label, bold=True, color=RGBColor(0x00, 0x00, 0x00))
            email_space = paragraph.add_run(" ")
            self._apply_default_run_format(email_space, bold=False, color=RGBColor(0x00, 0x00, 0x00))
            email_run = paragraph.add_run(email)
            self._apply_default_run_format(email_run, bold=False, color=RGBColor(0x00, 0x00, 0x00))
            tail_space = paragraph.add_run(" ")
            self._apply_default_run_format(tail_space, bold=False, color=RGBColor(0x00, 0x00, 0x00))
        if tel:
            newline2 = paragraph.add_run("\n")
            self._apply_default_run_format(newline2, bold=False, color=RGBColor(0x00, 0x00, 0x00))
            tel_label = paragraph.add_run("Tel.:")
            self._apply_default_run_format(tel_label, bold=True, color=RGBColor(0x00, 0x00, 0x00))
            tel_run = paragraph.add_run(" " + tel)
            self._apply_default_run_format(tel_run, bold=False, color=RGBColor(0x00, 0x00, 0x00))

    def _apply_default_run_format(
        self,
        run,
        *,
        bold: bool | None = None,
        superscript: bool | None = None,
        color: RGBColor | None = None,
    ) -> None:
        run.font.name = "Times New Roman"
        if run.font.size is None:
            from docx.shared import Pt
            run.font.size = Pt(12)
        else:
            from docx.shared import Pt
            run.font.size = Pt(12)
        if bold is not None:
            run.font.bold = bold
        if superscript is not None:
            run.font.superscript = superscript
        if color is not None:
            run.font.color.rgb = color

    def _normalize_name(self, text: str) -> str:
        lowered = (text or "").lower()
        lowered = re.sub(r"[^a-z0-9]+", "", lowered)
        return lowered
