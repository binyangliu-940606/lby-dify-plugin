from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ===== 新增：只设置一次全局默认样式，避免每格重复设置导致极慢 =====
def set_doc_defaults(doc: Document):
    """
    设置全局默认样式：Times New Roman，小四(12pt)，1.5行距
    只调用一次（生成 Table.docx 时最开始调用）
    """
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    # eastAsia 字体也设为 Times New Roman，避免中英文混排切换
    if style._element is not None and style._element.rPr is not None:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # 1.5 行距：Normal 段落样式设置一次即可
    # 注意：不同 Word 对 style paragraph_format 支持可能不同，但通常可用
    try:
        style.paragraph_format.line_spacing = 1.5
    except Exception:
        pass


# ===== 保留接口，但不再“每次都循环设置 run 样式” =====
def set_paragraph_style(paragraph):
    """
    保留该函数以兼容你原有调用，但改为“轻量化”：
    - 行距设置一次即可
    - 不再遍历每个 run 设置字体（这是性能灾难）
    """
    paragraph.paragraph_format.line_spacing = 1.5


def _set_cell_text(cell, text, bold=False):
    """
    高性能写单元格：
    - 不用 cell.text=（会重建段落结构，慢）
    - 直接写入 cell.paragraphs[0] 的 run
    - 字体由 Normal 样式兜底，仅表头需要 bold
    """
    # 取第一个段落（Word 表格单元格默认至少有一个段落）
    p = cell.paragraphs[0]

    # 清空已有 runs（不要删除段落节点，避免 XML 开销）
    for r in p.runs:
        r.text = ""

    run = p.add_run("" if text is None else str(text))
    run.bold = bool(bold)

    # 行距轻量设置（可选，不设置也行，因为 Normal 已设）
    p.paragraph_format.line_spacing = 1.5


# ===== 你原有边框逻辑不动 =====
def _set_border_nil(elem):
    elem.set(qn("w:val"), "nil")

def _set_border_single(elem, sz="8"):
    elem.set(qn("w:val"), "single")
    elem.set(qn("w:sz"), sz)
    elem.set(qn("w:space"), "0")
    elem.set(qn("w:color"), "000000")

def _clear_table_borders(tbl):
    tblPr = tbl._tbl.tblPr
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = tblBorders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            tblBorders.append(elem)
        _set_border_nil(elem)

    for row in tbl.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn("w:tcBorders"))
            if tcBorders is None:
                tcBorders = OxmlElement("w:tcBorders")
                tcPr.append(tcBorders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                e = tcBorders.find(qn(f"w:{edge}"))
                if e is None:
                    e = OxmlElement(f"w:{edge}")
                    tcBorders.append(e)
                _set_border_nil(e)

def set_table_three_line(tbl):
    if tbl is None:
        return
    if len(tbl.rows) == 0 or len(tbl.columns) == 0:
        return

    _clear_table_borders(tbl)

    tblPr = tbl._tbl.tblPr
    tblBorders = tblPr.find(qn("w:tblBorders"))
    top = tblBorders.find(qn("w:top"))
    bottom = tblBorders.find(qn("w:bottom"))
    _set_border_single(top, sz="8")
    _set_border_single(bottom, sz="8")

    first_row = tbl.rows[0]
    for cell in first_row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)

        btm = tcBorders.find(qn("w:bottom"))
        if btm is None:
            btm = OxmlElement("w:bottom")
            tcBorders.append(btm)
        _set_border_single(btm, sz="8")


# def add_table_with_title_note(doc: Document, title: str, note: str, grid: list, max_rows: int = 60, max_cols: int = 20):
#     """
#     性能优化版：
#     - 先裁剪 rows/cols，再创建表格（不要创建原始超大表再break）
#     - 用 tbl.rows[i].cells 替代 tbl.cell(i,j)
#     - 单元格写入用 add_run，不用 cell.text
#     - 表格设置 tbl.autofit = False 显著提速
#     """
#     # 表题
#     p_title = doc.add_paragraph(title if title else "")
#     if p_title.runs:
#         p_title.runs[0].bold = True
#     set_paragraph_style(p_title)

#     if not grid or len(grid) == 0:
#         grid = [["(Table content missing)"]]

#     # 先裁剪后建表（关键！）
#     rows = min(len(grid), max_rows)
#     cols = 1
#     for i in range(rows):
#         cols = max(cols, len(grid[i]))
#     cols = min(cols, max_cols)

#     tbl = doc.add_table(rows=rows, cols=cols)
#     tbl.autofit = False  # 关键提速点

#     for i in range(rows):
#         row_cells = tbl.rows[i].cells  # 一次取一行 cells（非常关键）
#         row_data = grid[i]
#         for j in range(cols):
#             val = row_data[j] if j < len(row_data) else ""
#             _set_cell_text(row_cells[j], val, bold=(i == 0))

#     set_table_three_line(tbl)

#     if note:
#         p_note = doc.add_paragraph(note)
#         set_paragraph_style(p_note)
def add_table_with_title_note(
    doc: Document,
    title: str,
    note: str,
    grid: list,
    max_rows: int = 60,
    max_cols: int = 20,
    skip_rows_threshold: int = 200,   # 超过此行数就不建表
    skip_cols_threshold: int = 60     # 超过此列数就不建表
):
    """
    性能优化 + 超大表跳过策略：
    1) 表题
    2) 如果原始表过大：不建表，只写说明（避免卡死）
    3) 否则：裁剪后建表（60x20 默认），写入三线表
    4) 表注
    """
    # 表题
    p_title = doc.add_paragraph(title if title else "")
    if p_title.runs:
        p_title.runs[0].bold = True
    set_paragraph_style(p_title)

    if not grid or len(grid) == 0:
        grid = [["(Table content missing)"]]

    # 计算原始规模（用于跳过策略）
    original_rows = len(grid)
    original_cols = max((len(r) for r in grid), default=1)

    # ===== 超大表：直接跳过建表，写说明 =====
    if original_rows > skip_rows_threshold or original_cols > skip_cols_threshold:
        msg = (
            f"(Table too large to render in Table.docx. "
            f"Original size: {original_rows} rows × {original_cols} cols. "
            f"Please refer to the source attachment for the full table.)"
        )
        p_msg = doc.add_paragraph(msg)
        set_paragraph_style(p_msg)

        # 表注仍然写（有助于保留缩写说明等）
        if note:
            p_note = doc.add_paragraph(note)
            set_paragraph_style(p_note)
        return

    # ===== 正常表：先裁剪后建表（关键） =====
    rows = min(original_rows, max_rows)

    cols = 1
    for i in range(rows):
        cols = max(cols, len(grid[i]))
    cols = min(cols, max_cols)

    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.autofit = False  # 显著提速

    for i in range(rows):
        row_cells = tbl.rows[i].cells
        row_data = grid[i]
        for j in range(cols):
            val = row_data[j] if j < len(row_data) else ""
            _set_cell_text(row_cells[j], val, bold=(i == 0))

    set_table_three_line(tbl)

    # 表注
    if note:
        p_note = doc.add_paragraph(note)
        set_paragraph_style(p_note)

def add_page_break(doc: Document):
    doc.add_page_break()