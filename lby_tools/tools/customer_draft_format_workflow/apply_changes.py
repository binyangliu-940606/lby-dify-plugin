import os
import shutil
import json
import os
import re
import shutil

from typing import Dict, Any, List, Tuple
from collections import OrderedDict

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from PIL import Image

from .docx_ops.run_span_replace import replace_text_preserve_runs
from .docx_ops.remove_tables_ooxml import remove_table_blocks
from .docx_ops.comments_core import create_comment_only
from .docx_ops.comments_anchor import add_comment_precise
from .docx_ops.table_3line import add_table_with_title_note, add_page_break,set_doc_defaults



# from dify_plugin import Tool
# from dify_plugin.entities.tool import ToolInvokeMessage



# class LbyToolsTool(Tool):
#     def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
#         #获取参数如文件url
#         payload_json = json.loads(tool_parameters["payload_json"])

#         sub_docx_path = payload_json["sub_docx_path"]
#         work_dir = payload_json["work_dir"]
#         normalize_pairs_final = payload_json["normalize_pairs_final"]
#         fig_assets = payload_json["fig_assets"]
#         table_assets = payload_json["table_assets"]
#         comment_tasks = payload_json["comment_tasks"]
#         figure_copy_plan = payload_json["figure_copy_plan"]

#         yield self.create_json_message({
#             "return_data": apply_all_changes(
#                             sub_docx_path,
#                             work_dir,
#                             normalize_pairs_final,
#                             fig_assets,
#                             table_assets,
#                             comment_tasks,
#                             figure_copy_plan,
#                         ),
#         })



# ========================
# 工具函数：基础
# ========================
def _ensure_dir(p: str):
    """确保目录存在；不存在则创建（递归创建）。"""
    os.makedirs(p, exist_ok=True)


# 主图+子图的常见格式：Figure 2A / Figure S1B
# 这里用于快速判断是否是这种“主图编号 + 子图字母”的 key
MAIN_FIG_PAT = re.compile(r"^(Figure\s+S?\d+)([A-Z])$")  # Figure 2A / Figure S1B


def normalize_main_figure_key(fk: str) -> str:
    """
    将 figure key 归一化到“主图级别”：
    - 'Figure 10A' -> 'Figure 10'
    - 'Figure 100' -> 'Figure 100'
    说明：只截取 'Figure + 数字' 部分，忽略后续子图字母。
    """
    m = re.match(r'^(Figure\s*\d+)', fk.strip(), flags=re.IGNORECASE)
    return m.group(1) if m else fk


def subfigure_suffix_rank(fk: str) -> int:
    """
    子图排序 rank：用于同一主图内部选择代表项、以及合并图注排序：
    - 没有后缀（如 'Figure 10'）返回 -1，表示最优先
    - 否则 A < B < C < D ...（按首字母）
    """
    m = re.match(r'^Figure\s*\d+([A-Za-z]+)?$', fk.strip())
    if not m or not m.group(1):
        return -1
    return ord(m.group(1)[0].upper()) - ord('A')


def _safe_remove(p: str):
    """安全删除文件：存在则删除；异常直接忽略，避免流程中断。"""
    try:
        if os.path.exists(p):
            os.remove(p)
    except Exception:
        pass


def _copy_file(src: str, dst: str):
    """复制文件到目标路径；目标目录不存在则先创建。"""
    _ensure_dir(os.path.dirname(dst))
    shutil.copyfile(src, dst)


def _move_paragraph_after(new_p: Paragraph, after_p: Paragraph):
    """把 new_p 段落移动到 after_p 后面（通过底层 xml 操作实现插入位置控制）。"""
    new_elm = new_p._p
    after_elm = after_p._p
    parent = after_elm.getparent()

    # 从原位置移除 new_p 的 xml 节点
    new_elm.getparent().remove(new_elm)
    # 插入到 after_p 的后一个位置
    parent.insert(parent.index(after_elm) + 1, new_elm)


# ========================
# 样式：居中 + 五号 + 宋体/Times New Roman
# ========================
def _style_paragraph_center_5(p: Paragraph):
    """
    将段落设置为：
    - 居中
    - 五号字（10.5pt）
    - 西文 Times New Roman
    - 中文宋体（eastAsia）
    适用于图题、图注等格式统一。
    """
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(10.5)  # 五号
        run.font.name = "Times New Roman"
        try:
            # 设置中文字体（eastAsia）为宋体，避免中文显示不一致
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        except Exception:
            pass


# ========================
# 1) normalize_pairs 清洗 + 子图兜底
# ========================
def split_citation_numbers(cit: str) -> List[int]:
    """
    输入形如: [1/2/100], [1,2,3], [1-3], 【1，2；3-5/7】...
    输出数字列表（仅“切割”，不展开区间：1-3 -> [1,3]）
    """
    inner = re.sub(r'^\s*(?:\[|【)\s*|\s*(?:\]|】)\s*$', '', cit)
    nums = [int(x) for x in re.split(r"\D+", inner) if x]
    return nums

def _extract_pmid_from_replace(repl: str) -> str:
    # 支持 "(PMID: 26521728)" / "PMID:26521728" / "26521728" 等
    m = re.search(r"\bPMID\s*:\s*(\d+)\b", repl, flags=re.I)
    if m:
        return m.group(1)
    m = re.search(r"\b(\d{6,})\b", repl)  # 兜底：取一个较长数字
    return m.group(1) if m else ""

def merge_citation_replacements(items: List[Dict[str, Any]]) -> List[Dict[str, str]]:
    """
    将形如:
      {'find': '[100,39, 40]', 'idx': 39, 'kind': 'pmid', 'replace': '(PMID: 26521728)'}
    合并为:
      {'kind': 'pmid', 'find': '[100,39, 40]', 'replace': '(100,PMID: 26521728,35281931)'}
    规则：
    - 按 (kind, find) 分组
    - find 中的数字序列按原顺序输出
    - 若某个数字 = idx 且有 replace，则用 "PMID: xxxx" 替换该数字
    - 若某个数字找不到匹配 idx，则原样输出数字（不加 PMID）
    """
    # (kind, find) -> {idx -> pmid}
    group_maps: "OrderedDict[Tuple[str,str], Dict[int,str]]" = OrderedDict()
    for it in items:
        key = (it["kind"], it["find"])
        m = group_maps.setdefault(key, {})
        pmid = _extract_pmid_from_replace(it.get("replace", ""))
        if pmid:
            m[int(it["idx"])] = pmid

    out: List[Dict[str, str]] = []
    for (kind, find), idx2pmid in group_maps.items():
        nums = split_citation_numbers(find)

        parts: List[str] = []
        for n in nums:
            if n in idx2pmid:
                parts.append(f"PMID: {idx2pmid[n]}")
            else:
                parts.append(str(n))

        out.append({
            "kind": kind,
            "find": find,
            "replace": f"({';'.join(parts)})"
        })
    return out

def _normalize_pairs_sanitize(normalize_pairs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    清洗 normalize_pairs（查找/替换对）：
    - 过滤非法结构
    - find/replace 支持 list 或 str，统一为单条 ("find": str, "replace": str)
    - 去重（kind, find, replace）
    - 按 find 长度从长到短排序，避免短串先替换导致长串匹配失效
    """
    out: List[Dict[str, Any]] = []
    seen = set()

    pmid_list = []
    for pair in normalize_pairs or []:
        if not isinstance(pair, dict):
            continue

        kind = pair.get("kind", "unknown")
        if kind == 'pmid':
            pmid_list.append(pair)

    out.extend(merge_citation_replacements(pmid_list))

    for pair in normalize_pairs or []:
        if not isinstance(pair, dict):
            continue

        kind = pair.get("kind", "unknown")
        if kind != 'pmid':
            fval = pair.get("find")
            rval = pair.get("replace")

            # replace 可能是 list：取第一个非空字符串
            if isinstance(rval, list):
                rval = next((x for x in rval if isinstance(x, str) and x.strip()), "")
            if not isinstance(rval, str):
                rval = "" if rval is None else str(rval)
            rval = rval.strip()
            # replace 为空则跳过
            if not rval:
                continue

            # find 可能是 list 或单字符串：统一到 finds 列表
            finds: List[str] = []
            if isinstance(fval, list):
                finds = [x.strip() for x in fval if isinstance(x, str) and x.strip()]
            elif isinstance(fval, str):
                if fval.strip():
                    finds = [fval.strip()]
            else:
                continue

            # 展开并去重
            for f in finds:
                key = (kind, f, rval)
                if key in seen:
                    continue
                seen.add(key)
                out.append({"kind": kind, "find": f, "replace": rval})

    # 长串优先替换，减少替换互相干扰
    out.sort(key=lambda x: len(x["find"]), reverse=True)
    return out


# 匹配形如 1A / S1B 的 token（用于兜底补回子图字母）
_SUBFIG_TOKEN_PAT = re.compile(r"(S?\d+)([A-Z])\b")  # 1A / S1B


def _fix_subfigure_loss_in_pairs(pairs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    兜底修复一种常见问题：
    - find 里带子图后缀（如 1A）
    - replace 里只替换成 Figure 1（丢掉 A）
    则自动把 A 追加回 replace 末尾，避免子图信息丢失。
    """
    out = []
    for p in pairs or []:
        find_text = p.get("find")
        rep_text = p.get("replace")
        if not isinstance(find_text, str) or not isinstance(rep_text, str):
            out.append(p)
            continue

        m = _SUBFIG_TOKEN_PAT.search(find_text)
        if m:
            num = m.group(1)
            letter = m.group(2)
            if re.search(rf"(?i)\bFigure\s+{re.escape(num)}\b", rep_text) and (letter not in rep_text):
                p = dict(p)
                p["replace"] = rep_text + letter
        out.append(p)
    return out


# ========================
# 2) run级替换（保格式）
# ========================
def _apply_normalization_replacements(docx_in: str, docx_out: str, normalize_pairs: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    对 docx 进行“按段落 run 级别”的文本替换：
    - 使用 replace_text_preserve_runs，尽量保留原有格式（粗体/斜体/上下标等）
    - 收集替换发生的 span 信息（段落索引、起止位置、替换类别等），用于后续定位/统计
    """
    doc = Document(docx_in)
    change_spans: List[Dict[str, Any]] = []

    for p_i, p in enumerate(doc.paragraphs):
        p_text = p.text or ""
        if not p_text:
            continue

        # 对每个 pair 做替换；注意 pairs 已按 find 长度降序
        for pair in normalize_pairs or []:
            find_text = pair["find"]
            replace_text = pair["replace"]
            kind = pair.get("kind", "unknown")

            if find_text in p_text:
                # 返回所有被替换的 span（start/end 相对于段落文本）
                replaced_spans = replace_text_preserve_runs(p, find_text, replace_text)
                for sp in replaced_spans:
                    change_spans.append({
                        "paragraph_index": p_i,
                        "start": sp.start,
                        "end": sp.end,
                        "kind": kind,
                        "replace": replace_text
                    })
                # 更新当前段落文本，避免多轮替换中 p_text 过期
                p_text = p.text or ""

    doc.save(docx_out)
    return {"ok": True, "change_spans": change_spans}


# ========================
# 3) 插入段落/图片
# ========================
def _insert_paragraph_after_title(doc: Document, after_p_index: int, text: str, *, center: bool = False, style5: bool = False) -> int:
    """
    在指定段落索引 after_p_index 之后插入新段落，并返回新段落索引（after_p_index+1）。
    可选：
    - center：段落居中
    - style5：应用“居中+五号+宋体/Times New Roman”的统一样式
    """
    after_p = doc.paragraphs[after_p_index]
    new_p = doc.add_paragraph(text)
    if center:
        new_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _move_paragraph_after(new_p, after_p)
    if style5:
        # _style_paragraph_center_5(new_p)
        #  """
        # 将段落设置为：
        # - 居中
        # - 五号字（10.5pt）
        # - 西文 Times New Roman
        # - 中文宋体（eastAsia）
        # 适用于图题、图注等格式统一。
        # """
        for run in new_p.runs:
            run.font.size = Pt(10.5)  # 五号
            run.font.bold = True
            run.font.name = "Times New Roman"
            try:
                # 设置中文字体（eastAsia）为宋体，避免中文显示不一致
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            except Exception:
                pass
    return after_p_index + 1

def _insert_paragraph_after(doc: Document, after_p_index: int, text: str, *, center: bool = False, style5: bool = False) -> int:
    """
    在指定段落索引 after_p_index 之后插入新段落，并返回新段落索引（after_p_index+1）。
    可选：
    - center：段落居中
    - style5：应用“居中+五号+宋体/Times New Roman”的统一样式
    """
    after_p = doc.paragraphs[after_p_index]
    new_p = doc.add_paragraph(text)
    if center:
        new_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _move_paragraph_after(new_p, after_p)
    if style5:
        # _style_paragraph_center_5(new_p)
        for run in new_p.runs:
            run.font.size = Pt(10.5)  # 五号
            run.font.name = "Times New Roman"
            try:
                # 设置中文字体（eastAsia）为宋体，避免中文显示不一致
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            except Exception:
                pass
    return after_p_index + 1


def _convert_tif_to_png_if_needed(image_path: str, cache_dir: str) -> str:
    """
    docx 对 tif/tiff 支持很差（通常无法直接插入），因此：
    - 若是 tif/tiff，则转换为 png 放在 cache_dir 下，并返回 png 路径
    - 若已转换过且缓存存在，直接复用
    - 非 tif/tiff 直接返回原路径
    """
    ext = os.path.splitext(image_path)[1].lower()
    if ext not in [".tif", ".tiff"]:
        return image_path

    _ensure_dir(cache_dir)
    out_path = os.path.join(cache_dir, os.path.splitext(os.path.basename(image_path))[0] + ".png")

    # 有缓存且非空，直接用缓存
    if os.path.exists(out_path) and os.path.getsize(out_path) > 0:
        return out_path

    # PIL 打开并转 RGB（避免带 alpha/索引色导致兼容问题）
    im = Image.open(image_path)
    try:
        im.seek(0)
    except Exception:
        pass
    im = im.convert("RGB")
    im.save(out_path, format="PNG")
    return out_path


def _is_docx_supported_image(path: str) -> bool:
    """判断图片扩展名是否通常可被 python-docx / Word 支持插入。"""
    ext = os.path.splitext(path)[1].lower()
    return ext in [".png", ".jpg", ".jpeg", ".gif", ".bmp"]


def _insert_picture_after(doc: Document, after_p_index: int, image_path: str, work_dir: str, width_inch: float = 6.0) -> Tuple[int, bool, str]:
    """
    在指定段落后插入居中的图片段落：
    - 自动 tif->png
    - 检查可插入格式
    - 插入图片并按 width_inch 缩放
    返回：(新的 after_p_index, 是否成功, 失败原因字符串)
    """
    if not image_path or not os.path.exists(image_path):
        return after_p_index, False, "image_not_found"

    try:
        ins_path = _convert_tif_to_png_if_needed(image_path, os.path.join(work_dir, "_img_cache"))
    except Exception as e:
        return after_p_index, False, f"tif_convert_failed: {e}"

    if not _is_docx_supported_image(ins_path):
        return after_p_index, False, f"unsupported_image_ext: {os.path.splitext(ins_path)[1].lower()}"

    after_p = doc.paragraphs[after_p_index]

    # 插入一个空段落用于承载图片
    p_pic = doc.add_paragraph("")
    p_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _move_paragraph_after(p_pic, after_p)
    _style_paragraph_center_5(p_pic)


    def normalize_image_for_docx(src_path: str) -> str:
        # 生成一个“确定可被python-docx识别”的文件
        dst_path = os.path.splitext(src_path)[0] + ".__normalized__.jpg"

        with Image.open(src_path) as im:
            # 有些图片是RGBA/LA/P，先转成RGB
            if im.mode not in ("RGB", "L"):
                im = im.convert("RGB")
            elif im.mode == "L":
                # 灰度也可以直接保存成JPEG，这里转RGB更稳
                im = im.convert("RGB")

            # 重新编码成标准JPEG
            im.save(dst_path, format="JPEG", quality=95, optimize=True)

        return dst_path

    # try:
    #     run = p_pic.add_run()
    #     run.add_picture(ins_path, width=Inches(width_inch))
    #     return after_p_index + 1, True, ""
    # except Exception as e:
    #     return after_p_index, False, f"add_picture_failed: {e}"
    

    try:
        run = p_pic.add_run()
        run.add_picture(ins_path, width=Inches(width_inch))
        return after_p_index + 1, True, ""
    except Exception as e:
        # 如果是识别失败，尝试重编码再插入
        if "UnrecognizedImageError" in type(e).__name__ or "UnrecognizedImageError" in str(e):
            try:
                fixed = normalize_image_for_docx(ins_path)
                run = p_pic.add_run()
                run.add_picture(fixed, width=Inches(width_inch))
                return after_p_index + 1, True, ""
            except Exception as e2:
                return after_p_index, False, f"add_picture_failed_after_normalize: {e2}"
        return after_p_index, False, f"add_picture_failed: {e}"


# ========================
# 3.5) Figure 聚合：主图只插一次（图片+标题+合并图注）
# ========================
def _group_fig_assets_by_main(fig_assets: Dict[str, Any]) -> Dict[str, List[Tuple[str, Dict[str, Any]]]]:
    """
    将 fig_assets 按“主图编号”聚合：
    - 输入 key 可能是 Figure 10A / Figure 10B / Figure 10
    - 归一化后按 Figure 10 分组
    每组内部按子图字母排序（无后缀的最优先）。
    """
    groups: Dict[str, List[Tuple[str, Dict[str, Any]]]] = {}
    for fk, asset in (fig_assets or {}).items():
        if not str(fk).startswith("Figure"):
            continue
        main_fk = normalize_main_figure_key(fk)
        groups.setdefault(main_fk, []).append((fk, asset or {}))

    for main_fk in groups:
        groups[main_fk].sort(key=lambda x: subfigure_suffix_rank(x[0]))
    return groups


def _merge_subfigure_captions(items: List[Tuple[str, Dict[str, Any]]]) -> str:
    """
    合并同一主图下各子图的 caption：
    - 若 key 满足 Figure 10A 这种形式，则拼成 "A. caption B. caption ..."
    - 否则直接拼 caption
    返回合并后的单行字符串。
    """
    parts = []
    for fk, a in items:
        cap = (a.get("caption") or "").strip()
        if not cap:
            continue
        m = re.match(r"^(Figure\s+S?\d+)([A-Z])$", fk.strip())
        if m:
            letter = m.group(2)
            parts.append(f"({letter}) {cap}")
        else:
            parts.append(cap)
    parts = " ".join(parts)
    return f'注: {parts}'


def _insert_main_figure_block_after_paragraph(
    doc: Document,
    p_index: int,
    main_fk: str,
    items: List[Tuple[str, Dict[str, Any]]],
    work_dir: str
) -> Dict[str, Any]:
    """
    在文中某段落 p_index 后插入一个“主图块”：
    - 同一主图只插入一次图片（使用排序后的第一个子图作为代表图）
    - 插入主图标题行（Figure X. Title）
    - 合并所有子图 caption 插入为图注
    返回插入结果（图片是否成功等）。
    """
    # items 已按 A/B/C... 排序，取第一个作为代表title与图片
    rep_fk, rep_asset = items[0]
    image_src_path = rep_asset.get("image_src_path")
    title = (rep_asset.get("title") or "").strip()

    merged_caption = _merge_subfigure_captions(items)

    cur = p_index
    result = {"fig_key": main_fk, "pic_ok": False, "pic_reason": ""}

    # 1) 图片（若有）
    if image_src_path:
        cur, ok, reason = _insert_picture_after(doc, cur, image_src_path, work_dir, width_inch=6.0)
        result["pic_ok"] = ok
        result["pic_reason"] = reason

    # 2) 标题（居中+五号）
    # 你如果要严格 "Figure 10.TSNE..."（无空格），用 f"{main_fk}.{title}"
    title_line = f"{main_fk}. {title}".strip() if title else f"{main_fk}."
    cur = _insert_paragraph_after_title(doc, cur, title_line, center=True, style5=True)

    # 3) 合并图注（居中+五号）
    if merged_caption:
        cur = _insert_paragraph_after(doc, cur, merged_caption, center=True, style5=True)

    # 4) （新增）空一行：在图注后再插入一个空段落
    cur = _insert_paragraph_after(doc, cur, "", center=False, style5=False)
    
    return result


# ========================
# 4) comment：替换后重新定位 anchor，再精确锚定
# ========================
def _find_anchor_span_in_doc(doc: Document, anchor: str, p_hint: int = None, window: int = 4):
    """
    在 doc 中查找 anchor 出现的位置，并返回 (段落索引, start, end)：
    - 如果提供 p_hint，则先在其附近 window 范围内查找（更快且更准）
    - 对 "Figure 10" / "Table 2" 这类锚点做特殊处理：避免误匹配 Figure 10A（使用 (?![A-Z])）
    - 若附近找不到，则全量扫描全文
    """
    if not anchor:
        return None

    safe_anchor = re.escape(anchor.strip())
    if re.match(r"(?i)^(Figure|Table)\s+S?\d+$", anchor.strip()):
        # 负向前瞻：后面不是 A-Z，避免 Figure 10 命中 Figure 10A
        pat = re.compile(safe_anchor + r"(?![A-Z])")
    else:
        pat = re.compile(safe_anchor)

    def scan(a: int, b: int):
        for pi in range(a, b):
            text = doc.paragraphs[pi].text or ""
            m = pat.search(text)
            if m:
                return pi, m.start(), m.end()
        return None

    n = len(doc.paragraphs)
    # 先在提示段落附近搜索
    if isinstance(p_hint, int) and 0 <= p_hint < n:
        a = max(0, p_hint - window)
        b = min(n, p_hint + window + 1)
        hit = scan(a, b)
        if hit:
            return hit

    # 再全量搜索
    return scan(0, n)


def _apply_comments_precise_after_locate(docx_in: str, docx_out: str, comment_tasks: List[Dict[str, Any]], work_dir: str) -> Dict[str, Any]:
    """
    给 docx 批量添加批注（comment），并尽量精确锚定到 anchor 的字符范围：
    - 先复制 docx_in -> docx_out
    - 每条任务独立进行：先定位 anchor 的 (p_i, start, end)
      然后：
        1) create_comment_only 创建 comment 记录（返回 comment_id）
        2) add_comment_precise 将 comment_id 精确锚定到段落范围
    - 每条任务都用临时文件串联，成功则覆盖输出文件
    """
    shutil.copyfile(docx_in, docx_out)

    ok_count = 0
    fallback_skip = 0

    for i, task in enumerate(comment_tasks or []):
        anchor = task.get("anchor") or ""
        ctext = task.get("text", "")
        p_hint = task.get("p_hint", task.get("p"))

        # anchor 或评论内容为空则跳过
        if not anchor or not ctext.strip():
            continue

        # 注意：每次循环重新打开 docx_out，是为了拿到“最新版本”去定位
        doc_loc = Document(docx_out)
        loc = _find_anchor_span_in_doc(doc_loc, anchor, p_hint=p_hint, window=4)
        if not loc:
            fallback_skip += 1
            continue
        p_i, start, end = loc

        # step1：只创建 comment（不锚定），拿到 comment_id
        tmp1 = os.path.join(work_dir, f"__tmp_cmt_create_{i}.docx")
        r1 = create_comment_only(docx_out, tmp1, ctext, author="DLAB")
        if not r1.get("ok"):
            _safe_remove(tmp1)
            continue
        cid = r1["comment_id"]

        # step2：将 comment 精确锚定到段落的 start-end 范围
        tmp2 = os.path.join(work_dir, f"__tmp_cmt_anchor_{i}.docx")
        r2 = add_comment_precise(tmp1, tmp2, p_i, start, end, cid)

        if r2.get("ok"):
            # 成功则以 tmp2 覆盖当前输出版本
            shutil.move(tmp2, docx_out)
            ok_count += 1
        else:
            fallback_skip += 1

        # 清理临时文件
        _safe_remove(tmp1)
        _safe_remove(tmp2)

    return {"ok": True, "comment_ok": ok_count, "comment_skip": fallback_skip}

# ========================
# 5) Table.docx 生成
# ========================
def _generate_table_docx(table_docx_path: str, table_assets: Dict[str, Any]) -> Dict[str, Any]:
    """
    根据 table_assets 生成单独的 Table.docx：
    - 每个表：标题行 + 三线表 + note
    - 表与表之间插入分页符
    """
    tdoc = Document()
    set_doc_defaults(tdoc)

    keys = sorted(list((table_assets or {}).keys()))

    for idx, tkey in enumerate(keys):
        asset = table_assets[tkey] or {}
        title = (asset.get("title") or "").strip()
        note = (asset.get("note") or "").strip()
        grid = asset.get("grid") or []

        title_line = f"{tkey}. {title}".strip() if title else f"{tkey}."
        add_table_with_title_note(tdoc, title_line, note, grid)

        # 最后一张表后不加分页
        if idx != len(keys) - 1:
            add_page_break(tdoc)

    tdoc.save(table_docx_path)
    return {"ok": True, "table_count": len(keys)}


# ========================
# 6) 主函数：apply_all_changes
# ========================
def apply_all_changes(
    sub_docx_path: str,
    work_dir: str,
    normalize_pairs: List[Dict[str, Any]],
    fig_assets: Dict[str, Any],
    table_assets: Dict[str, Any],
    comment_tasks: List[Dict[str, Any]],
    figure_copy_plan: List[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """
    主流程：对输入 sub_docx 进行一系列处理，产出 Final.docx / Table.docx / Figure 文件夹：
    1) 清洗 normalize_pairs
    2) docx 归一化替换（run 级保格式）
    3) 删除文中表格块（表格另存到 Table.docx）
    4) 在正文中插入 Figure（按主图聚合：图片只插一次 + 合并图注）
    5) 重新定位并写入 comments（批注）
    6) 生成 Table.docx
    7) 复制 Figure 图片到 Figure 目录（按主图只拷贝一次，优先A）
    """
    _ensure_dir(work_dir)

    # 各步骤输出路径（便于排查问题/分步查看）
    step1_norm = os.path.join(work_dir, "Final.step1.normalized.docx")
    step2_no_tables = os.path.join(work_dir, "Final.step2.no_tables.docx")
    step3_with_figs = os.path.join(work_dir, "Final.step3.with_figures.docx")
    final_docx_path = os.path.join(work_dir, "Final.docx")

    # Table 输出
    table_docx_path = os.path.join(work_dir, "Tables.docx")
    # Figure 资源目录
    final_figure_dir = os.path.join(work_dir, "Figures")
    _ensure_dir(final_figure_dir)

    # 1) normalize_pairs 清洗
    pairs = _normalize_pairs_sanitize(normalize_pairs or [])

    # 2) 归一化替换
    r_norm = _apply_normalization_replacements(sub_docx_path, step1_norm, pairs)
    change_spans = r_norm.get("change_spans", [])

    # 3) 删除表块（正文中移除表格，表格另生成 Table.docx）
    remove_table_blocks(step1_norm, step2_no_tables)


















    # # 4) 插入 Figure 块：按主图聚合（图片只一次 + 图注合并）
    # doc2 = Document(step2_no_tables)

    # # 将 Figure 资产按主图分组
    # groups = _group_fig_assets_by_main(fig_assets or {})

    # # 找每个主图在文中“首次出现段落位置”：取其所有子图 key 命中的最小段落号
    # first_pos_main: Dict[str, int] = {}
    # for main_fk, items in groups.items():
    #     best = None
    #     for fk, _asset in items:
    #         # fk 多为 Figure 10A，不需要(?![A-Z])；但保留原逻辑兼容无后缀的 Figure 10
    #         if re.match(r"(?i)^Figure\s+S?\d+$", fk):
    #             pat = re.compile(re.escape(fk) + r"(?![A-Z])")
    #         else:
    #             pat = re.compile(re.escape(fk))

    #         # 从头扫描段落，找到首次出现的位置
    #         for p_i, p in enumerate(doc2.paragraphs):
    #             t = p.text or ""
    #             if pat.search(t):
    #                 best = p_i if best is None else min(best, p_i)
    #                 break
    #     if best is not None:
    #         first_pos_main[main_fk] = best

    # # 生成插入任务列表：需要插入的主图（已存在于 sub 的则跳过）
    # insert_jobs: List[Tuple[int, str]] = []
    # for main_fk, p_i in first_pos_main.items():
    #     items = groups[main_fk]
    #     # 只要这一组里存在一个未 already_exists_in_sub，就插入主图块
    #     if all((a.get("already_exists_in_sub", False) for _fk, a in items)):
    #         continue
    #     insert_jobs.append((p_i, main_fk))

    # # 倒序插入，避免插入段落导致后续索引偏移
    # fig_insert_reports = []
    # for p_i, main_fk in sorted(insert_jobs, key=lambda x: x[0], reverse=True):
    #     rep = _insert_main_figure_block_after_paragraph(doc2, p_i, main_fk, groups[main_fk], work_dir)
    #     fig_insert_reports.append(rep)







    doc2 = Document(step2_no_tables)

    # 将 Figure 资产按主图分组
    groups = _group_fig_assets_by_main(fig_assets or {})

    # ---------- Results范围定位 + “小段标题”识别（支持纯正文加粗） ----------
    def _norm(s: str) -> str:
        return re.sub(r"\s+", " ", (s or "")).strip().lower()

    def _find_results_range(paragraphs):
        """
        返回 (results_start_idx, results_end_idx_exclusive)
        - start: 标题为 Results / 结果 的那一段的下一段开始
        - end: 下一个顶级板块标题（如 Discussion/讨论/Conclusion...）出现的位置；若无则到文末
        """
        results_title_idx = None
        for i, p in enumerate(paragraphs):
            t = _norm(p.text)
            # 只匹配“单独一行”的 Results/结果（避免匹配到 'Results are expressed...'）
            if t in ("results", "结果"):
                results_title_idx = i
                break

        if results_title_idx is None:
            # 没有 Results：退化为全文
            return 0, len(paragraphs)

        start = results_title_idx + 1

        end = len(paragraphs)
        stop_titles = {"discussion", "讨论", "conclusion", "conclusions", "结论", "总结"}
        for j in range(start, len(paragraphs)):
            t = _norm(paragraphs[j].text)
            if t in stop_titles:
                end = j
                break

        return start, end

    def _is_section_heading(paragraph) -> bool:
        """
        判断是否“小段标题”：
        1) style 含 Heading（最稳）
        2) 或者：该段落文本主要为加粗（纯正文加粗标题）
        3) 再兜底：短句 + 不以句号结尾
        """
        style_name = (getattr(paragraph.style, "name", "") or "").lower()
        if "heading" in style_name:
            return True

        text = (paragraph.text or "").strip()
        if not text:
            return False

        # 2) 纯正文加粗标题识别：统计非空 runs 的加粗比例
        runs = [r for r in (paragraph.runs or []) if (r.text or "").strip()]
        if runs:
            bold_cnt = 0
            for r in runs:
                if r.bold is True:
                    bold_cnt += 1
            bold_ratio = bold_cnt / len(runs)

            if bold_ratio >= 0.8:
                # 过滤 'Figure 2A' 这类
                if not re.match(r"(?i)^\s*figure\s+\w+", text):
                    return True

        # 3) 兜底：短句且不以句号结尾
        if len(text) <= 120 and not re.search(r"[\.。]$", text):
            if len(text.split()) <= 20:
                # 过滤常见正文开头句式（按需增补）
                if not re.match(r"(?i)^(in this study|to evaluate|to investigate|we |data were )", text):
                    return True

        return False

    results_start, results_end = _find_results_range(doc2.paragraphs)

    # 找到 Results 内每个“小段标题”索引，并构建小段边界
    # sections: [(heading_idx, body_start_idx, body_end_idx_exclusive)]
    heading_idxs = [i for i in range(results_start, results_end) if _is_section_heading(doc2.paragraphs[i])]
    sections = []
    if heading_idxs:
        for k, h_i in enumerate(heading_idxs):
            body_start = h_i + 1
            body_end = heading_idxs[k + 1] if k + 1 < len(heading_idxs) else results_end
            sections.append((h_i, body_start, body_end))
    else:
        # Results 没有可识别标题：整体视作一个 section
        sections.append((results_start - 1, results_start, results_end))

    def _last_nonempty_par_idx(start: int, end_excl: int) -> int:
        """在 [start, end_excl) 内从后往前找最后一个非空段落索引；找不到则返回 start-1"""
        for i in range(end_excl - 1, start - 1, -1):
            if (doc2.paragraphs[i].text or "").strip():
                return i
        return start - 1

    # Results兜底插入点：Results最后一个非空段落之后
    results_fallback_insert_pos = _last_nonempty_par_idx(results_start, results_end)
    if results_fallback_insert_pos < 0:
        results_fallback_insert_pos = max(results_start - 1, 0)
    # ---------- 结束：Results范围定位 + section识别 ----------


    # ---------- 从 Results 扫描定位，插到“下一小标题的上一行”（本小段末尾） ----------
    # first_pos_main: 主图 -> 插入到哪个段落之后 (after_idx)
    first_pos_main: Dict[str, int] = {}

    for main_fk, items in groups.items():
        best_after_idx = None

        for fk, _asset in items:
            # fk 多为 Figure 10A；兼容无后缀的 Figure 10
            if re.match(r"(?i)^Figure\s+S?\d+$", fk):
                pat = re.compile(re.escape(fk) + r"(?![A-Z])")
            else:
                pat = re.compile(re.escape(fk))

            # 只在 Results 范围内扫描
            hit_paragraph_idx = None
            for p_i in range(results_start, results_end):
                t = doc2.paragraphs[p_i].text or ""
                if pat.search(t):
                    hit_paragraph_idx = p_i
                    break

            if hit_paragraph_idx is None:
                continue

            # 命中 -> 找所属 section，并把插入点定到该 section 的末尾（下一小标题上一行）
            chosen_after_idx = None
            for (h_i, body_start, body_end) in sections:
                if body_start <= hit_paragraph_idx < body_end:
                    last_body_idx = _last_nonempty_par_idx(body_start, body_end)
                    chosen_after_idx = last_body_idx if last_body_idx >= body_start else (body_end - 1)
                    break

            if chosen_after_idx is None:
                chosen_after_idx = results_fallback_insert_pos

            best_after_idx = chosen_after_idx if best_after_idx is None else min(best_after_idx, chosen_after_idx)

        # Results 内完全找不到引用：直接放 Results 板块最后
        if best_after_idx is None:
            best_after_idx = results_fallback_insert_pos

        first_pos_main[main_fk] = best_after_idx


    # 生成插入任务列表：需要插入的主图（已存在于 sub 的则跳过）
    insert_jobs: List[Tuple[int, str]] = []
    for main_fk, after_idx in first_pos_main.items():
        items = groups[main_fk]
        if all((a.get("already_exists_in_sub", False) for _fk, a in items)):
            continue
        insert_jobs.append((after_idx, main_fk))

    # 同一小段内多个主图：按 Figure 1,2,3 顺序插入（同 after_idx 排序）
    def _fig_sort_key(fig_label: str):
        m = re.match(r"(?i)figure\s+(s?)\s*(\d+)", (fig_label or "").strip())
        if not m:
            return (1, 10**9, fig_label)
        is_supp = 1 if m.group(1).lower() == "s" else 0  # 正文图优先于补充图（可调整）
        num = int(m.group(2))
        return (is_supp, num, fig_label)

    # 倒序插入避免索引漂移；同一插入点内按 Figure 编号升序
    fig_insert_reports = []
    for after_idx, main_fk in sorted(
        insert_jobs,
        key=lambda x: (x[0], _fig_sort_key(x[1])),
        reverse=True
    ):
        rep = _insert_main_figure_block_after_paragraph(
            doc2, after_idx, main_fk, groups[main_fk], work_dir
        )
        fig_insert_reports.append(rep)
















        

    doc2.save(step3_with_figs)

    # 5) 写 comments（替换后重新定位 anchor）
    _apply_comments_precise_after_locate(step3_with_figs, final_docx_path, comment_tasks or [], work_dir)

    # 6) 生成 Table.docx
    _generate_table_docx(table_docx_path, table_assets or {})

    # 7) Figure 文件夹拷贝（保留原后缀）——按主图只拷贝一次（优先A）
    copied = 0
    copied_main = set()

    # 优先使用 fig_assets 直接生成拷贝计划（更可靠）
    groups_for_copy = _group_fig_assets_by_main(fig_assets or {})
    for main_fk, items in groups_for_copy.items():
        if main_fk in copied_main:
            continue
        # 取排序后的第一个子图作为代表文件（优先 A，无后缀最优先）
        rep_fk, rep_asset = items[0]
        src = rep_asset.get("image_src_path")
        if not src or not os.path.exists(src):
            continue

        # 输出文件名按主图编号命名，并保留原始后缀
        ext = os.path.splitext(src)[1]  # 保留原后缀
        dst_main_name = main_fk + ext
        _copy_file(src, os.path.join(final_figure_dir, dst_main_name))
        copied_main.add(main_fk)
        copied += 1

    # 返回产物路径与统计信息
    return {
        "final_docx_path": final_docx_path,
        "table_docx_path": table_docx_path,
        "final_figure_dir": final_figure_dir,
        "change_spans": change_spans,
        "figure_insert_reports": fig_insert_reports,
        "stats": {
            "normalized_changes": len(change_spans),
            "figures_inserted": len(insert_jobs),
            "figures_copied": copied
        }
    }