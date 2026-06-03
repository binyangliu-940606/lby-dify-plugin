# -*- coding: utf-8 -*-
import copy
import io
import re
import json
import zipfile
import requests

from collections.abc import Generator
from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage

from typing import Any
from docx import Document
from docx.shared import Pt
from xml.etree import ElementTree as ET

# ---------------------------------------------------------------------------
# Tool 入口
# ---------------------------------------------------------------------------


class LbyToolsTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage, None, None]:
        payload = json.loads(tool_parameters["payload_json"])

        records = payload["antibody_json"]
        wb_text = payload["wb_text"]
        new_wb_text = payload["new_wb_text"]
        doc_url = payload["doc_url"]
        doc_id = payload["doc_id"]
        file_name = payload["file_name"]

        enriched = enrich_wb_text_rule(wb_text, records)

        if wb_text and not new_wb_text:
            new_wb_text = enriched

        logs: list[str] = []

        if not new_wb_text and enriched:
            new_wb_text = enriched
            logs.append("new_wb_text 为空，已使用规则引擎 enriched_wb_text")
        elif (
            new_wb_text
            and wb_text
            and count_matches(new_wb_text, records) < count_matches(enriched, records)
        ):
            new_wb_text = enriched
            logs.append("LLM 嵌入条数少于规则引擎，已回退为 enriched_wb_text")

        raw = load_doc_bytes(doc_url)
        updated, rep_log = replace_western_blot_in_docx(raw, wb_text, new_wb_text)
        logs.extend(rep_log)

        mw_bytes = build_mw_list_docx(records, doc_id)

        file_name_mane_docx = file_name
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        yield self.create_blob_message(
            blob=updated,
            meta={
                "mime_type": mime_type,
                "filename": file_name_mane_docx,
            },
        )

        file_name_wm_list = "单独分子量文件.docx"
        yield self.create_blob_message(
            blob=mw_bytes,
            meta={
                "mime_type": mime_type,
                "filename": file_name_wm_list,
            },
        )


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}

# ---------------------------------------------------------------------------
# JSON 解析
# ---------------------------------------------------------------------------


def norm_catalog(cat: str) -> str:
    c = (cat or "").strip().lower()
    c = c.lstrip("#")
    return c


# ---------------------------------------------------------------------------
# 分子量精简（写入正文 & 列表）
# ---------------------------------------------------------------------------
def shorten_mw(mw_official: str) -> str:
    """从 mw_official 提取适合写入 WB 段落/列表的简短 kDa 描述。"""
    s = (mw_official or "").strip()
    if not s:
        return ""

    # 优先：Predicted: 44 kDa
    m = re.search(r"Predicted:\s*([^;]+)", s, re.I)
    if m:
        pred = m.group(1).strip()
        det = re.search(r"Detected[^:]*:\s*([^;]+)", s, re.I)
        if det:
            d = det.group(1).strip()
            # 若检测到多条带，取 "and" 前后形成范围
            nums = re.findall(r"(\d+)\s*kDa", d, re.I)
            if len(nums) >= 2:
                lo, hi = min(int(x) for x in nums), max(int(x) for x in nums)
                pred_n = re.search(r"(\d+)\s*kDa", pred, re.I)
                if pred_n:
                    return f"{lo}-{hi}kDa"
            return normalize_kda_token(d.split(";")[0])
        return normalize_kda_token(pred)

    # 单值 44 kDa / 54kDa
    m = re.search(r"(\d+(?:\s*-\s*\d+)?)\s*kDa", s, re.I)
    if m:
        return normalize_kda_token(m.group(0))

    return s[:80]


def normalize_kda_token(token: str) -> str:
    t = token.strip()
    t = re.sub(r"\s*kDa", "kDa", t, flags=re.I)
    t = re.sub(r"\s+", "", t)
    if "kDa" not in t and re.search(r"\d", t):
        t += "kDa"
    return t


def display_antibody_name(rec: dict, fallback_from_text: str = "") -> str:
    """分子量列表用的短名称。"""
    name = (rec.get("target_name") or "").strip()
    name = re.sub(r"\s+antibody\s*$", "", name, flags=re.I)
    name = re.sub(r"\s+antibod(y|ies)\s*$", "", name, flags=re.I)
    if name:
        return name
    return fallback_from_text or "Unknown"


def product_note(rec: dict) -> str:
    """图1括号内产品描述：(target | vendor)。"""
    target = (rec.get("target_name") or "").strip()
    vendor = (rec.get("vendor") or "").strip()
    notes = (rec.get("notes") or "").strip()
    if target and vendor:
        base = f"{target} | {vendor}"
    else:
        base = target or vendor or notes[:120]
    return base


# ---------------------------------------------------------------------------
# 规则化嵌入（LLM 兜底 / 校验）
# ---------------------------------------------------------------------------
def enrich_wb_text_rule(wb_text: str, records: list[dict]) -> str:
    """
    在每个一抗条目 vendor 之后、稀释比之前插入：，{mw}，({product_note})
    匹配依据：catalog_no（及 # 前缀变体），按货号逐条替换，避免复杂抗体名误匹配。
    """
    text = wb_text or ""
    if not text or not records:
        return text

    for rec in records:
        cat = (rec.get("catalog_no") or "").strip()
        if not cat:
            continue
        cat_esc = re.escape(cat.lstrip("#"))
        mw = shorten_mw(rec.get("mw_official", ""))
        note = product_note(rec)
        if not mw and not note:
            continue

        insert = ""
        if mw:
            insert += f"，{mw}"
        if note:
            insert += f"，({note})"

        # 仅匹配「货号，供应商，稀释比」三段式，且尚未含 kDa
        pat = re.compile(
            rf"（\s*#?{cat_esc}\s*，\s*([^，]+?)\s*，\s*(\d+:\d+)\s*）"
        )

        def repl_keep_hash(m: re.Match) -> str:
            vendor = m.group(1)
            dilution = m.group(2)
            full = m.group(0)
            if re.search(r"\d\s*kDa", full, re.I):
                return full
            # 还原原始货号写法（含 #）
            cat_in_doc = re.search(
                rf"（\s*(#?{cat_esc})\s*，", full, re.I
            )
            cat_token = cat_in_doc.group(1) if cat_in_doc else cat
            return f"（{cat_token}，{vendor}{insert}，{dilution}）"

        text = pat.sub(repl_keep_hash, text)  # 同一货号多处（如 GAPDH）全部替换

    return text


def count_matches(text: str, records: list[dict]) -> int:
    if not text:
        return 0
    n = 0
    for rec in records:
        cat = norm_catalog(rec.get("catalog_no", ""))
        if cat and re.search(rf"（\s*#?{re.escape(cat)}\s*，", text, re.I):
            n += 1
    return n


# ---------------------------------------------------------------------------
# DOCX：下载与 Western blot 段落替换（保留 run 样式）
# ---------------------------------------------------------------------------


def load_doc_bytes(doc_url: str) -> bytes:
    url = (doc_url or "").strip()

    try:
        resp = requests.get(url, timeout=120, verify=True)
        resp.raise_for_status()
        data = resp.content
    except requests.RequestException as e:
        raise RuntimeError(f"下载文档失败: {e}") from e

    if len(data) < 1000:
        raise RuntimeError("下载的文件过小，可能不是有效的 docx")
    return data


def replace_western_blot_in_docx(
    docx_bytes: bytes, old_wb_text: str, new_wb_text: str
) -> tuple[bytes, list[str]]:
    """
    用 old_wb_text 在文档中定位连续段落，整段替换为 new_wb_text。
    old_wb_text / new_wb_text 均为完整 WB 正文，按换行对应 Word 中的连续 w:p。
    """
    logs: list[str] = []
    old_text = (old_wb_text or "").replace("\r\n", "\n").strip()
    new_text = (new_wb_text or "").replace("\r\n", "\n").strip()

    if not old_text:
        logs.append("警告：old_wb_text 为空，未执行替换")
        return docx_bytes, logs

    old_parts = split_body_paragraphs(old_text)
    new_parts = split_body_paragraphs(new_text) if new_text else [""]

    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zin:
        xml = zin.read("word/document.xml")
        root = ET.fromstring(xml)
        body = root.find("w:body", NS)
        if body is None:
            raise RuntimeError("document.xml 中未找到 w:body")

        paragraphs = body.findall("w:p", NS)
        replaced = False

        found = find_consecutive_paragraph_match(paragraphs, old_parts)
        if found:
            start, chunk_ps = found
            replace_paragraph_range(body, chunk_ps, new_parts)
            end = start + len(old_parts) - 1
            if len(old_parts) == 1:
                logs.append(f"已替换 WB 段落（索引 {start}）")
            else:
                logs.append(
                    f"已替换 WB 连续段落（索引 {start}-{end}，共 {len(old_parts)} 段）"
                )
            replaced = True
        else:
            combined_old = "".join(old_parts)
            for i, p in enumerate(paragraphs):
                pt = get_paragraph_text(p)
                if not pt:
                    continue
                if old_text in pt or combined_old in pt or fuzzy_wb_match(pt, old_text):
                    if len(new_parts) > 1:
                        replace_paragraph_range(body, [p], new_parts)
                        logs.append(
                            f"已单段命中 old_wb_text，展开替换为 {len(new_parts)} 段（索引 {i}）"
                        )
                    else:
                        set_paragraph_text_preserve_runs(p, new_parts[0])
                        logs.append(f"已替换 WB 段落（索引 {i}）")
                    replaced = True
                    break

        if not replaced:
            logs.append("警告：未定位到与 old_wb_text 匹配的段落，请检查 wb_text 是否与文档一致")

        new_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)

        out_buf = io.BytesIO()
        with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == "word/document.xml":
                    zout.writestr(item, new_xml)
                else:
                    zout.writestr(item, zin.read(item.filename))

    return out_buf.getvalue(), logs


def split_body_paragraphs(text: str) -> list[str]:
    """将 body 按换行拆成多个段落（忽略空行）。"""
    return [part.strip() for part in (text or "").split("\n") if part.strip()]


def collapse_ws(s: str) -> str:
    return re.sub(r"\s+", "", s or "")


def paragraph_text_match(doc_text: str, old_text: str) -> bool:
    if not old_text:
        return False
    if old_text in doc_text or doc_text in old_text:
        return True
    return fuzzy_wb_match(doc_text, old_text)


def consecutive_paragraphs_match(doc_parts: list[str], old_parts: list[str]) -> bool:
    if len(doc_parts) != len(old_parts):
        return False
    if all(paragraph_text_match(doc, old) for doc, old in zip(doc_parts, old_parts)):
        return True
    return fuzzy_wb_match("".join(doc_parts), "".join(old_parts))


def find_consecutive_paragraph_match(
    paragraphs: list[ET.Element], old_parts: list[str]
) -> tuple[int, list[ET.Element]] | None:
    """在文档中定位与 old_parts 一一对应的连续 w:p 段落。"""
    if not old_parts:
        return None
    n = len(old_parts)
    for i in range(0, len(paragraphs) - n + 1):
        chunk_ps = paragraphs[i : i + n]
        chunk_texts = [get_paragraph_text(p) for p in chunk_ps]
        if consecutive_paragraphs_match(chunk_texts, old_parts):
            return i, chunk_ps
    return None


def clone_paragraph_element(p: ET.Element) -> ET.Element:
    new_p = copy.deepcopy(p)
    set_paragraph_text_preserve_runs(new_p, "")
    return new_p


def replace_paragraph_range(
    body: ET.Element,
    matched_ps: list[ET.Element],
    new_parts: list[str],
) -> None:
    """将连续若干 w:p 替换为 new_parts（每段对应一个 w:p，可多可少）。"""
    if not matched_ps:
        return
    if not new_parts:
        new_parts = [""]

    for i, text in enumerate(new_parts):
        if i < len(matched_ps):
            set_paragraph_text_preserve_runs(matched_ps[i], text)
            continue
        template = matched_ps[-1]
        new_p = clone_paragraph_element(template)
        set_paragraph_text_preserve_runs(new_p, text)
        insert_at = list(body).index(matched_ps[i - 1]) + 1
        body.insert(insert_at, new_p)

    for p in matched_ps[len(new_parts) :]:
        body.remove(p)


def fuzzy_wb_match(doc_para: str, old_body: str) -> bool:
    if not old_body:
        return False
    doc_compact = collapse_ws(doc_para)
    old_compact = collapse_ws(old_body)
    a = doc_compact[:200]
    b = old_compact[:200]
    return (
        a == b
        or (len(b) > 80 and b in doc_compact)
        or (len(old_compact) > 80 and old_compact in doc_compact)
    )


def get_paragraph_text(p: ET.Element) -> str:
    parts: list[str] = []
    for t in p.iter(f"{{{W_NS}}}t"):
        if t.text:
            parts.append(t.text)
        if t.tail:
            parts.append(t.tail)
    return "".join(parts)


def set_paragraph_text_preserve_runs(p: ET.Element, new_text: str) -> None:
    """
    将全部新文本写入第一个 w:t，其余 w:t 置空，保留 w:r 及其 w:rPr（字体、加粗、下划线等）。
    """
    texts = list(p.iter(f"{{{W_NS}}}t"))
    if not texts:
        r = ET.SubElement(p, f"{{{W_NS}}}r")
        t = ET.SubElement(r, f"{{{W_NS}}}t")
        t.text = new_text
        return
    texts[0].text = new_text
    for t in texts[1:]:
        t.text = ""
        if t.tail:
            t.tail = ""


# ---------------------------------------------------------------------------
# 分子量列表 DOCX
# ---------------------------------------------------------------------------
def build_mw_list_docx(records: list[dict], doc_id: str) -> bytes:
    doc = Document()
    title = doc.add_paragraph()
    title_run = "XG 分子量"
    if doc_id:
        title_run = f"XG({doc_id})分子量"
    run = title.add_run(title_run)
    run.bold = True
    run.font.size = Pt(16)

    for rec in records:
        if (rec.get("retrieval_status") or "").lower() == "skip":
            continue
        name = display_antibody_name(rec)
        cat = (rec.get("catalog_no") or "").strip().lstrip("#")
        mw = shorten_mw(rec.get("mw_official", "")) or (rec.get("mw_in_doc") or "").strip() or "未检索"
        # GAPDH 等有时不写「抗体」
        ab_word = "抗体" if "igg" not in name.lower() and "gapdh" not in name.lower() else ""
        if ab_word:
            line = f"{name} 抗体： 货号 {cat}, 分子量： {mw}"
        else:
            line = f"{name} ： 货号 {cat}, 分子量： {mw}"
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
