# -*- coding: utf-8 -*-
import copy
import io
import re
import json
import zipfile
import requests

from typing import Any
from docx import Document
from docx.shared import Pt
from xml.etree import ElementTree as ET

# ---------------------------------------------------------------------------
# Tool 入口
# ---------------------------------------------------------------------------

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}

# ---------------------------------------------------------------------------
# JSON 解析
# ---------------------------------------------------------------------------


def norm_catalog(cat: str) -> str:
    c = (cat or "").strip().lower()
    c = c.lstrip("#")
    return c


# ---------------------------------------------------------------------------
# 分子量精简（写入正文 & 列表）
# ---------------------------------------------------------------------------
def shorten_mw(mw_official: str) -> str:
    """从 mw_official 提取适合写入 WB 段落/列表的简短 kDa 描述。"""
    s = (mw_official or "").strip()
    if not s:
        return ""

    # 优先：Predicted: 44 kDa
    m = re.search(r"Predicted:\s*([^;]+)", s, re.I)
    if m:
        pred = m.group(1).strip()
        det = re.search(r"Detected[^:]*:\s*([^;]+)", s, re.I)
        if det:
            d = det.group(1).strip()
            # 若检测到多条带，取 "and" 前后形成范围
            nums = re.findall(r"(\d+)\s*kDa", d, re.I)
            if len(nums) >= 2:
                lo, hi = min(int(x) for x in nums), max(int(x) for x in nums)
                pred_n = re.search(r"(\d+)\s*kDa", pred, re.I)
                if pred_n:
                    return f"{lo}-{hi}kDa"
            return normalize_kda_token(d.split(";")[0])
        return normalize_kda_token(pred)

    # 单值 44 kDa / 54kDa
    m = re.search(r"(\d+(?:\s*-\s*\d+)?)\s*kDa", s, re.I)
    if m:
        return normalize_kda_token(m.group(0))

    return s[:80]


def normalize_kda_token(token: str) -> str:
    t = token.strip()
    t = re.sub(r"\s*kDa", "kDa", t, flags=re.I)
    t = re.sub(r"\s+", "", t)
    if "kDa" not in t and re.search(r"\d", t):
        t += "kDa"
    return t


def display_antibody_name(rec: dict, fallback_from_text: str = "") -> str:
    """分子量列表用的短名称。"""
    name = (rec.get("target_name") or "").strip()
    name = re.sub(r"\s+antibody\s*$", "", name, flags=re.I)
    name = re.sub(r"\s+antibod(y|ies)\s*$", "", name, flags=re.I)
    if name:
        return name
    return fallback_from_text or "Unknown"


def product_note(rec: dict) -> str:
    """图1括号内产品描述：(target | vendor)。"""
    target = (rec.get("target_name") or "").strip()
    vendor = (rec.get("vendor") or "").strip()
    notes = (rec.get("notes") or "").strip()
    if target and vendor:
        base = f"{target} | {vendor}"
    else:
        base = target or vendor or notes[:120]
    return base


# ---------------------------------------------------------------------------
# 规则化嵌入（LLM 兜底 / 校验）
# ---------------------------------------------------------------------------
def enrich_wb_text_rule(wb_text: str, records: list[dict]) -> str:
    """
    在每个一抗条目 vendor 之后、稀释比之前插入：，{mw}，({product_note})
    匹配依据：catalog_no（及 # 前缀变体），按货号逐条替换，避免复杂抗体名误匹配。
    """
    text = wb_text or ""
    if not text or not records:
        return text

    for rec in records:
        cat = (rec.get("catalog_no") or "").strip()
        if not cat:
            continue
        cat_esc = re.escape(cat.lstrip("#"))
        mw = shorten_mw(rec.get("mw_official", ""))
        note = product_note(rec)
        if not mw and not note:
            continue

        insert = ""
        if mw:
            insert += f"，{mw}"
        if note:
            insert += f"，({note})"

        # 仅匹配「货号，供应商，稀释比」三段式，且尚未含 kDa
        pat = re.compile(
            rf"（\s*#?{cat_esc}\s*，\s*([^，]+?)\s*，\s*(\d+:\d+)\s*）"
        )

        def repl_keep_hash(m: re.Match) -> str:
            vendor = m.group(1)
            dilution = m.group(2)
            full = m.group(0)
            if re.search(r"\d\s*kDa", full, re.I):
                return full
            # 还原原始货号写法（含 #）
            cat_in_doc = re.search(
                rf"（\s*(#?{cat_esc})\s*，", full, re.I
            )
            cat_token = cat_in_doc.group(1) if cat_in_doc else cat
            return f"（{cat_token}，{vendor}{insert}，{dilution}）"

        text = pat.sub(repl_keep_hash, text)  # 同一货号多处（如 GAPDH）全部替换

    return text


def count_matches(text: str, records: list[dict]) -> int:
    if not text:
        return 0
    n = 0
    for rec in records:
        cat = norm_catalog(rec.get("catalog_no", ""))
        if cat and re.search(rf"（\s*#?{re.escape(cat)}\s*，", text, re.I):
            n += 1
    return n


# ---------------------------------------------------------------------------
# DOCX：下载与 Western blot 段落替换（保留 run 样式）
# ---------------------------------------------------------------------------


def load_doc_bytes(doc_url: str) -> bytes:
    url = (doc_url or "").strip()

    try:
        resp = requests.get(url, timeout=120, verify=True)
        resp.raise_for_status()
        data = resp.content
    except requests.RequestException as e:
        raise RuntimeError(f"下载文档失败: {e}") from e

    if len(data) < 1000:
        raise RuntimeError("下载的文件过小，可能不是有效的 docx")
    return data


def replace_western_blot_in_docx(
    docx_bytes: bytes, old_wb_text: str, new_wb_text: str
) -> tuple[bytes, list[str]]:
    """
    用 old_wb_text 在文档中定位连续段落，整段替换为 new_wb_text。
    old_wb_text / new_wb_text 均为完整 WB 正文，按换行对应 Word 中的连续 w:p。
    """
    logs: list[str] = []
    old_text = (old_wb_text or "").replace("\r\n", "\n").strip()
    new_text = (new_wb_text or "").replace("\r\n", "\n").strip()

    if not old_text:
        logs.append("警告：old_wb_text 为空，未执行替换")
        return docx_bytes, logs

    old_parts = split_body_paragraphs(old_text)
    new_parts = split_body_paragraphs(new_text) if new_text else [""]

    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zin:
        xml = zin.read("word/document.xml")
        root = ET.fromstring(xml)
        body = root.find("w:body", NS)
        if body is None:
            raise RuntimeError("document.xml 中未找到 w:body")

        paragraphs = body.findall("w:p", NS)
        replaced = False

        found = find_consecutive_paragraph_match(paragraphs, old_parts)
        if found:
            start, chunk_ps = found
            replace_paragraph_range(body, chunk_ps, new_parts)
            end = start + len(old_parts) - 1
            if len(old_parts) == 1:
                logs.append(f"已替换 WB 段落（索引 {start}）")
            else:
                logs.append(
                    f"已替换 WB 连续段落（索引 {start}-{end}，共 {len(old_parts)} 段）"
                )
            replaced = True
        else:
            combined_old = "".join(old_parts)
            for i, p in enumerate(paragraphs):
                pt = get_paragraph_text(p)
                if not pt:
                    continue
                if old_text in pt or combined_old in pt or fuzzy_wb_match(pt, old_text):
                    if len(new_parts) > 1:
                        replace_paragraph_range(body, [p], new_parts)
                        logs.append(
                            f"已单段命中 old_wb_text，展开替换为 {len(new_parts)} 段（索引 {i}）"
                        )
                    else:
                        set_paragraph_text_preserve_runs(p, new_parts[0])
                        logs.append(f"已替换 WB 段落（索引 {i}）")
                    replaced = True
                    break

        if not replaced:
            logs.append("警告：未定位到与 old_wb_text 匹配的段落，请检查 wb_text 是否与文档一致")

        new_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)

        out_buf = io.BytesIO()
        with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == "word/document.xml":
                    zout.writestr(item, new_xml)
                else:
                    zout.writestr(item, zin.read(item.filename))

    return out_buf.getvalue(), logs


def split_body_paragraphs(text: str) -> list[str]:
    """将 body 按换行拆成多个段落（忽略空行）。"""
    return [part.strip() for part in (text or "").split("\n") if part.strip()]


def collapse_ws(s: str) -> str:
    return re.sub(r"\s+", "", s or "")


def paragraph_text_match(doc_text: str, old_text: str) -> bool:
    if not old_text:
        return False
    if old_text in doc_text or doc_text in old_text:
        return True
    return fuzzy_wb_match(doc_text, old_text)


def consecutive_paragraphs_match(doc_parts: list[str], old_parts: list[str]) -> bool:
    if len(doc_parts) != len(old_parts):
        return False
    if all(paragraph_text_match(doc, old) for doc, old in zip(doc_parts, old_parts)):
        return True
    return fuzzy_wb_match("".join(doc_parts), "".join(old_parts))


def find_consecutive_paragraph_match(
    paragraphs: list[ET.Element], old_parts: list[str]
) -> tuple[int, list[ET.Element]] | None:
    """在文档中定位与 old_parts 一一对应的连续 w:p 段落。"""
    if not old_parts:
        return None
    n = len(old_parts)
    for i in range(0, len(paragraphs) - n + 1):
        chunk_ps = paragraphs[i : i + n]
        chunk_texts = [get_paragraph_text(p) for p in chunk_ps]
        if consecutive_paragraphs_match(chunk_texts, old_parts):
            return i, chunk_ps
    return None


def clone_paragraph_element(p: ET.Element) -> ET.Element:
    new_p = copy.deepcopy(p)
    set_paragraph_text_preserve_runs(new_p, "")
    return new_p


def replace_paragraph_range(
    body: ET.Element,
    matched_ps: list[ET.Element],
    new_parts: list[str],
) -> None:
    """将连续若干 w:p 替换为 new_parts（每段对应一个 w:p，可多可少）。"""
    if not matched_ps:
        return
    if not new_parts:
        new_parts = [""]

    for i, text in enumerate(new_parts):
        if i < len(matched_ps):
            set_paragraph_text_preserve_runs(matched_ps[i], text)
            continue
        template = matched_ps[-1]
        new_p = clone_paragraph_element(template)
        set_paragraph_text_preserve_runs(new_p, text)
        insert_at = list(body).index(matched_ps[i - 1]) + 1
        body.insert(insert_at, new_p)

    for p in matched_ps[len(new_parts) :]:
        body.remove(p)


def fuzzy_wb_match(doc_para: str, old_body: str) -> bool:
    if not old_body:
        return False
    doc_compact = collapse_ws(doc_para)
    old_compact = collapse_ws(old_body)
    a = doc_compact[:200]
    b = old_compact[:200]
    return (
        a == b
        or (len(b) > 80 and b in doc_compact)
        or (len(old_compact) > 80 and old_compact in doc_compact)
    )


def get_paragraph_text(p: ET.Element) -> str:
    parts: list[str] = []
    for t in p.iter(f"{{{W_NS}}}t"):
        if t.text:
            parts.append(t.text)
        if t.tail:
            parts.append(t.tail)
    return "".join(parts)


def set_paragraph_text_preserve_runs(p: ET.Element, new_text: str) -> None:
    """
    将全部新文本写入第一个 w:t，其余 w:t 置空，保留 w:r 及其 w:rPr（字体、加粗、下划线等）。
    """
    texts = list(p.iter(f"{{{W_NS}}}t"))
    if not texts:
        r = ET.SubElement(p, f"{{{W_NS}}}r")
        t = ET.SubElement(r, f"{{{W_NS}}}t")
        t.text = new_text
        return
    texts[0].text = new_text
    for t in texts[1:]:
        t.text = ""
        if t.tail:
            t.tail = ""


# ---------------------------------------------------------------------------
# 分子量列表 DOCX
# ---------------------------------------------------------------------------
def build_mw_list_docx(records: list[dict], doc_id: str) -> bytes:
    doc = Document()
    title = doc.add_paragraph()
    title_run = "XG 分子量"
    if doc_id:
        title_run = f"XG({doc_id})分子量"
    run = title.add_run(title_run)
    run.bold = True
    run.font.size = Pt(16)

    for rec in records:
        if (rec.get("retrieval_status") or "").lower() == "skip":
            continue
        name = display_antibody_name(rec)
        cat = (rec.get("catalog_no") or "").strip().lstrip("#")
        mw = shorten_mw(rec.get("mw_official", "")) or (rec.get("mw_in_doc") or "").strip() or "未检索"
        # GAPDH 等有时不写「抗体」
        ab_word = "抗体" if "igg" not in name.lower() and "gapdh" not in name.lower() else ""
        if ab_word:
            line = f"{name} 抗体： 货号 {cat}, 分子量： {mw}"
        else:
            line = f"{name} ： 货号 {cat}, 分子量： {mw}"
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()



payload_json= "{\"antibody_json\": [{\"target_name\": \"CD41 antibody\", \"vendor\": \"Abcam\", \"catalog_no\": \"ab134131\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u6297CD41\\uff08Rabbit, Abcam, Cat#ab134131, 1:1000\\uff09\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u4e3b\\u8981\\u4e00\\u6297\\u5217\\u8868\", \"confidence\": 0.98, \"search_query\": \"CD41 antibody Abcam ab134131 molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"114 kDa\", \"mw_type\": \"predicted\", \"evidence_url\": \"https://www.abcam.com/en-us/products/primary-antibodies/cd41-antibody-epr4330-ab134131\", \"evidence_quote\": \"Molecular weight: 114 kDa\", \"retrieval_status\": \"found\", \"notes\": \"Molecular weight taken from Abcam product page for ab134131; corresponds to the predicted molecular weight of the target Integrin alpha-IIb (CD41, UniProt P08514).\"}, {\"target_name\": \"CD61 antibody\", \"vendor\": \"Abcam\", \"catalog_no\": \"ab179475\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u6297CD61\\uff08Rabbit, Abcam, Cat#ab179475, 1:1000\\uff09\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u4e3b\\u8981\\u4e00\\u6297\\u5217\\u8868\", \"confidence\": 0.98, \"search_query\": \"CD61 antibody Abcam ab179475 molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"113 kDa\", \"mw_type\": \"predicted\", \"evidence_url\": \"https://www.abcam.co.jp/products/primary-antibodies/integrin-alpha-v-antibody-epr16800-ab179475\", \"evidence_quote\": \"Anti-Integrin alpha V antibody [EPR16800] (ab179475) specifically detects Integrin alpha V (UniProt ID: P06756; Molecular weight: 113kDa) and is sold in 100 \\u00b5L\", \"retrieval_status\": \"found\", \"notes\": \"The Abcam product page for catalog number ab179475 identifies the target as Integrin alpha V (not CD61/integrin beta 3) and lists a molecular weight of 113 kDa. I used the vendor (Abcam) product page matching the exact catalog number as the evidence. The supplied target_name (CD61 antibody) appears to be a mismatch with the catalog number.\"}, {\"target_name\": \"CD42b antibody\", \"vendor\": \"Abcam\", \"catalog_no\": \"ab183345\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u6297CD42b\\uff08Rabbit, Abcam, Cat#ab183345, 1:1000\\uff09\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u4e3b\\u8981\\u4e00\\u6297\\u5217\\u8868\", \"confidence\": 0.98, \"search_query\": \"CD42b antibody Abcam ab183345 molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"\", \"mw_type\": \"unknown\", \"evidence_url\": \"\", \"evidence_quote\": \"\", \"retrieval_status\": \"not_found\", \"notes\": \"\\u5df2\\u68c0\\u7d22\\u5b98\\u65b9\\u4ea7\\u54c1\\u9875\\u53ca\\u53ef\\u5f97\\u7684\\u7b2c\\u4e09\\u65b9\\u805a\\u5408\\u9875\\uff1a\\u68c0\\u67e5\\u4e86Abcam\\u5b98\\u65b9\\u4ea7\\u54c1\\u9875\\uff08https://www.abcam.com/en-us/products/primary-antibodies/cd42b-antibody-sp219-ab183345\\uff09\\u3001Abcam\\u76f8\\u5173\\u4ea7\\u54c1\\u53d8\\u4f53\\u9875\\u3001\\u7ecf\\u9500/\\u5206\\u9500\\u4fe1\\u606f\\uff08iright\\u3001Danaher lifesciences\\u3001CiteAb\\uff09\\u3002\\u5728\\u53ef\\u6293\\u53d6\\u7684\\u9875\\u9762\\u6458\\u8981\\u4e0e\\u5185\\u5bb9\\u7247\\u6bb5\\u4e2d\\u672a\\u53d1\\u73b0\\u4efb\\u4f55\\u5173\\u4e8e\\u76ee\\u6807\\u86cb\\u767d/\\u8be5\\u6297\\u4f53\\u7684\\u5206\\u5b50\\u91cf\\u3001\\u9884\\u671f\\u6761\\u5e26\\u5927\\u5c0f\\u6216\\u89c2\\u5bdf\\u5230\\u7684\\u8868\\u89c2\\u5206\\u5b50\\u91cf\\u7684\\u660e\\u786e\\u8868\\u8ff0\\u3002\\u672a\\u627e\\u5230\\u5b98\\u65b9datasheet\\u6216\\u8bf4\\u660e\\u4e66\\u4e2d\\u5305\\u542b\\u5206\\u5b50\\u91cf\\u4fe1\\u606f\\u7684\\u53ef\\u8ffd\\u6eaf\\u8bc1\\u636e\\u3002\\u82e5\\u9700\\u7ee7\\u7eed\\uff0c\\u53ef\\u76f4\\u63a5\\u6253\\u5f00Abcam\\u4ea7\\u54c1\\u9875\\u7684\\u5b8c\\u6574\\u8bf4\\u660e\\u4e66\\u6216\\u8054\\u7cfb\\u5382\\u5546\\u83b7\\u53d6WB\\u9884\\u671f\\u6761\\u5e26/\\u76ee\\u6807\\u86cb\\u767d\\u5206\\u5b50\\u91cf\\u4fe1\\u606f\\u3002\"}, {\"target_name\": \"HIF-2\\u03b1 antibody\", \"vendor\": \"Abcam\", \"catalog_no\": \"ab199\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u6297HIF-2\\u03b1\\uff08Rabbit, Abcam, Cambridge, UK, Cat#ab199, 1:1000\\uff09\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u4e3b\\u8981\\u4e00\\u6297\\u5217\\u8868\", \"confidence\": 0.98, \"search_query\": \"HIF-2\\u03b1 antibody Abcam ab199 molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"Approximately 118 kDa (observed band); predicted molecular weight: 100 kDa\", \"mw_type\": \"band_size\", \"evidence_url\": \"https://doc.abcam.com/legacy-unpublished/datasheets/com/datasheet_199.pdf\", \"evidence_quote\": \"Detects a band of approximately 118 kDa (predicted molecular weight: 100 kDa).\", \"retrieval_status\": \"found\", \"notes\": \"Abcam legacy product datasheet (PDF) explicitly reports an observed WB band of ~118 kDa and a predicted molecular weight of 100 kDa for ab199.\"}, {\"target_name\": \"ATG5 antibody\", \"vendor\": \"Cell Signaling Technology\", \"catalog_no\": \"12994\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u6297ATG5\\uff08Rabbit, Cell Signaling Technology, Cat#12994, 1:1000\\uff09\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u4e3b\\u8981\\u4e00\\u6297\\u5217\\u8868\", \"confidence\": 0.98, \"search_query\": \"ATG5 antibody Cell Signaling Technology 12994 molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"55 kDa\", \"mw_type\": \"predicted\", \"evidence_url\": \"https://www.cellsignal.com/products/primary-antibodies/atg5-d5f5u-rabbit-monoclonal-antibody/12994\", \"evidence_quote\": \"MW (kDa) | 55\", \"retrieval_status\": \"found\", \"notes\": \"\\u5b98\\u65b9\\u4ea7\\u54c1\\u9875\\u5217\\u51fa\\u201cMW (kDa) 55\\u201d\\u3002\\u6839\\u636eCell Signaling\\u4ea7\\u54c1\\u9875\\u9762\\u7684\\u5b57\\u6bb5\\u6807\\u7b7e\\uff08MW (kDa)\\uff09\\uff0c\\u5c06\\u7c7b\\u578b\\u6807\\u6ce8\\u4e3apredicted\\uff08\\u7406\\u8bba/\\u6807\\u6ce8\\u5206\\u5b50\\u91cf\\uff09\\u3002\\u540c\\u65f6\\u5b58\\u5728\\u5b98\\u65b9datasheet\\uff08#12994\\uff09\\uff0c\\u4f46\\u4ea7\\u54c1\\u9875\\u5df2\\u660e\\u786e\\u7ed9\\u51fa\\u5206\\u5b50\\u91cf\\u3002\"}, {\"target_name\": \"phospho-ULK1 (Ser757) antibody\", \"vendor\": \"Cell Signaling Technology\", \"catalog_no\": \"14202\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u6297ULK1\\u3001p-ULK1 Ser555\\u53caSer757\\uff08Rabbit, Cell Signaling Technology, Cat#8054/5869/14202, 1:1000\\uff09\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u4e3b\\u8981\\u4e00\\u6297\\u5217\\u8868\", \"confidence\": 0.96, \"search_query\": \"phospho-ULK1 (Ser757) antibody Cell Signaling Technology 14202 molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"140-150 kDa\", \"mw_type\": \"apparent\", \"evidence_url\": \"https://www.cellsignal.com/products/primary-antibodies/phospho-ulk1-ser757-d7o6u-rabbit-monoclonal-antibody/14202\", \"evidence_quote\": \"MW (kDa) | 140-150\", \"retrieval_status\": \"found\", \"notes\": \"Information obtained from the vendor product page for catalog #14202. Cell Signaling's datasheet for a related ULK1 phospho antibody (#6888) also lists MW (kDa) 140-150, supporting consistency across CST listings.\"}, {\"target_name\": \"AMPK antibody\", \"vendor\": \"Cell Signaling Technology\", \"catalog_no\": \"2532\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u6297AMPK\\u53cap-AMPK Thr172\\uff08Rabbit, Cell Signaling Technology, Cat#2532/2535, 1:1000\\uff09\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u4e3b\\u8981\\u4e00\\u6297\\u5217\\u8868\", \"confidence\": 0.95, \"search_query\": \"AMPK antibody Cell Signaling Technology 2532 molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"62 kDa\", \"mw_type\": \"apparent\", \"evidence_url\": \"https://media.cellsignal.com/pdf/2532.pdf\", \"evidence_quote\": \"Sensitivity:   Endogenous   MW (kDa):   62\", \"retrieval_status\": \"found\", \"notes\": \"Official CST datasheet (PDF) lists endogenous MW as 62 kDa (used as apparent/observed molecular weight). CST technical support article also notes the specific AMPK\\u03b1 signal is a prominent band around 62 kDa, though occasional non-specific bands near ~50 kDa may appear.\"}, {\"target_name\": \"phospho-AMPK (Thr172) antibody\", \"vendor\": \"Cell Signaling Technology\", \"catalog_no\": \"2535\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u6297AMPK\\u53cap-AMPK Thr172\\uff08Rabbit, Cell Signaling Technology, Cat#2532/2535, 1:1000\\uff09\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u4e3b\\u8981\\u4e00\\u6297\\u5217\\u8868\", \"confidence\": 0.96, \"search_query\": \"phospho-AMPK (Thr172) antibody Cell Signaling Technology 2535 molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"62 kDa\", \"mw_type\": \"apparent\", \"evidence_url\": \"https://media.cellsignal.com/pdf/2535.pdf\", \"evidence_quote\": \"Endogenous MW (kDa): 62\", \"retrieval_status\": \"found\", \"notes\": \"Source: official Cell Signaling Technology datasheet (2535.pdf) which lists the endogenous molecular weight as 62 kDa; mw_type set to apparent because the datasheet reports 'Endogenous MW' (observed).\"}, {\"target_name\": \"LC3B antibody\", \"vendor\": \"Cell Signaling Technology\", \"catalog_no\": \"2775\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u4ee5\\u53ca\\u81ea\\u566c\\u76f8\\u5173\\u86cb\\u767d\\u6297LC3B\\uff08Rabbit, Cell Signaling Technology, Cat#2775, 1:1000\\uff09\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u4e3b\\u8981\\u4e00\\u6297\\u5217\\u8868\", \"confidence\": 0.98, \"search_query\": \"LC3B antibody Cell Signaling Technology 2775 molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"14, 16 kDa\", \"mw_type\": \"band_size\", \"evidence_url\": \"https://www.cellsignal.com/products/primary-antibodies/lc3b-antibody/2775\", \"evidence_quote\": \"MW (kDa):   14, 16\", \"retrieval_status\": \"found\", \"notes\": \"Official Cell Signaling Technology product page and datasheet (https://media.cellsignal.com/pdf/2775.pdf) list MW (kDa) as 14 and 16 kDa. The two bands correspond to LC3B type I/II (multiple forms); datasheet notes stronger reactivity with the type II form.\"}, {\"target_name\": \"phospho-4EBP1 antibody\", \"vendor\": \"Cell Signaling Technology\", \"catalog_no\": \"2855\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u6297p-4EBP1\\uff08Rabbit, Cell Signaling Technology, Cat#2855, 1:1000\\uff09\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u4e3b\\u8981\\u4e00\\u6297\\u5217\\u8868\", \"confidence\": 0.97, \"search_query\": \"phospho-4EBP1 antibody Cell Signaling Technology 2855 molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"Calculated: 12 kDa; Observed (apparent): 15\\u201320 kDa (data from alternative vendor Product page for 4E-BP1)\", \"mw_type\": \"apparent\", \"evidence_url\": \"https://www.ptglab.com/products/Phospho-4EBP1-Thr37-Antibody-81812-4-RR.htm\", \"evidence_quote\": \"Calculated Molecular Weight | 118 aa, 12 kDa\\nObserved Molecular Weight | 15-20 kDa\", \"retrieval_status\": \"uncertain\", \"notes\": \"Could not find explicit molecular weight on the Cell Signaling Technology product page for catalog 2855 in the provided search snippets (CST product URL present but no MW in content). Used an alternative vendor (Proteintech) product page for phospho-4EBP1 which lists calculated (12 kDa) and observed (15\\u201320 kDa) molecular weights for 4E-BP1. This is not the official CST datasheet; please verify on the CST product page or official datasheet for catalog 2855.\"}, {\"target_name\": \"mTOR antibody\", \"vendor\": \"Cell Signaling Technology\", \"catalog_no\": \"2983\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u6297mTOR\\u53cap-mTOR Ser2448\\uff08Rabbit, Cell Signaling Technology, Cat#2983/5536, 1:1000\\uff09\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u4e3b\\u8981\\u4e00\\u6297\\u5217\\u8868\", \"confidence\": 0.95, \"search_query\": \"mTOR antibody Cell Signaling Technology 2983 molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"289 kDa\", \"mw_type\": \"predicted\", \"evidence_url\": \"https://media.cellsignal.com/pdf/2983.pdf\", \"evidence_quote\": \"Sensitivity: Endogenous MW (kDa): 289\", \"retrieval_status\": \"found\", \"notes\": \"Official Cell Signaling Technology datasheet for product #2983 lists 'MW (kDa): 289'. mw_type set to 'predicted' because the datasheet provides the molecular weight value without explicitly stating 'expected band size' or 'observed/apparent' band; interpretation as theoretical/target molecular weight. Source: CST product datasheet PDF (2983).\"}, {\"target_name\": \"Beclin-1 antibody\", \"vendor\": \"Cell Signaling Technology\", \"catalog_no\": \"3495\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u6297Beclin-1\\uff08Rabbit, Cell Signaling Technology, Cat#3495, 1:1000\\uff09\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u4e3b\\u8981\\u4e00\\u6297\\u5217\\u8868\", \"confidence\": 0.98, \"search_query\": \"Beclin-1 antibody Cell Signaling Technology 3495 molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"60 kDa\", \"mw_type\": \"band_size\", \"evidence_url\": \"https://www.cellsignal.com/products/primary-antibodies/beclin-1-d40c5-rabbit-monoclonal-antibody/3495\", \"evidence_quote\": \"MW (kDa) | 60\", \"retrieval_status\": \"found\", \"notes\": \"\\u4fe1\\u606f\\u6765\\u81eaCell Signaling Technology\\u5b98\\u7f51\\u4ea7\\u54c1\\u9875/\\u8bf4\\u660e\\u4e66\\uff0c\\u9875\\u9762\\u5217\\u51fa\\u201cMW (kDa): 60\\u201d\\uff0c\\u4e3a\\u5382\\u5bb6\\u63d0\\u4f9b\\u7684\\u9884\\u671f/\\u6761\\u5e26\\u5206\\u5b50\\u91cf\\uff08\\u56e0\\u6b64\\u6807\\u6ce8\\u4e3a band_size\\uff09\\u3002\"}, {\"target_name\": \"HIF-1\\u03b1 antibody\", \"vendor\": \"Cell Signaling Technology\", \"catalog_no\": \"36169\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u4e3b\\u8981\\u4e00\\u6297\\u5305\\u62ec\\uff1a\\u6297HIF-1\\u03b1\\uff08Rabbit, Cell Signaling Technology, Danvers, MA, USA, Cat#36169, 1:1000\\uff09\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u4e3b\\u8981\\u4e00\\u6297\\u5217\\u8868\", \"confidence\": 0.98, \"search_query\": \"HIF-1\\u03b1 antibody Cell Signaling Technology 36169 molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"\", \"mw_type\": \"unknown\", \"evidence_url\": \"\", \"evidence_quote\": \"\", \"retrieval_status\": \"not_found\", \"notes\": \"\\u5df2\\u68c0\\u67e5Cell Signaling Technology\\u7684\\u4ea7\\u54c1\\u8bf4\\u660e\\u4e66/\\u6570\\u636e\\u8868\\uff08Tavily\\u7ed3\\u679c\\u4e2d\\u7684\\u4e09\\u4e2a\\u5b98\\u65b9datasheet\\u94fe\\u63a5\\uff1ahttps://www.cellsignal.com/products/36169/datasheet?images=1 \\uff0c https://www.cellsignal.com/products/36169/datasheet?images=0&protocol=0 \\uff0c https://www.cellsignal.com/products/36169/datasheet?images=1&protocol=0&size=A4\\uff09\\uff0c\\u4f46\\u6293\\u53d6\\u5185\\u5bb9\\u4e2d\\u672a\\u5305\\u542b\\u660e\\u786e\\u7684\\u5206\\u5b50\\u91cf/\\u6761\\u5e26\\u5927\\u5c0f\\u8bf4\\u660e\\uff08\\u9875\\u9762\\u7247\\u6bb5\\u4ec5\\u63cf\\u8ff0\\u4e86Simple Western\\u5b9e\\u9a8c\\u4e0e\\u5206\\u5b50\\u91cf\\u8303\\u56f4\\u6a21\\u5757\\u4f46\\u672a\\u7ed9\\u51fa\\u76ee\\u6807\\u5206\\u5b50\\u91cf\\u6570\\u5b57\\uff09\\u3002\\u4ea6\\u68c0\\u67e5\\u4e86\\u975e\\u5b98\\u65b9\\u4f9b\\u5e94\\u5546/\\u8d44\\u6e90\\uff08\\u5982Abclonal A11945 PDF\\u4e0eAbcam HIF-1\\u03b1\\u9875\\u9762\\uff09\\uff0c\\u5b83\\u4eec\\u5206\\u522b\\u62a5\\u544a\\u4e86Observed MW 120 kDa / Predicted 83\\u201396 kDa\\u6216Predicted 92 kDa\\u3001Observed 110 kDa\\uff0c\\u4f46\\u8fd9\\u4e9b\\u4e0d\\u662f\\u5bf9\\u5e94\\u4e8eCell Signaling\\u8d27\\u53f736169\\u7684\\u5b98\\u65b9\\u6765\\u6e90\\u3002\\u56e0\\u672a\\u5728\\u5b98\\u65b9\\u4ea7\\u54c1\\u9875\\u6216\\u5b98\\u65b9datasheet\\u4e2d\\u627e\\u5230\\u4e0e\\u8d27\\u53f7\\u4e25\\u683c\\u5339\\u914d\\u7684\\u5206\\u5b50\\u91cf\\u58f0\\u660e\\uff0c\\u6545\\u6807\\u8bb0\\u4e3anot_found\\u3002\\u5c1d\\u8bd5\\uff1a1) \\u641c\\u7d22\\u5e76\\u67e5\\u770b\\u5b98\\u65b9datasheet\\uff08\\u591a\\u79cd\\u663e\\u793a\\u683c\\u5f0f\\uff09\\uff1b2) \\u6bd4\\u5bf9\\u7b2c\\u4e09\\u65b9\\u4f9b\\u5e94\\u5546\\u4fe1\\u606f\\u4ee5\\u9a8c\\u8bc1\\u5e38\\u89c1\\u62a5\\u544a\\u503c\\uff08\\u5df2\\u8bb0\\u5f55\\u4f46\\u672a\\u4f5c\\u4e3a\\u5b98\\u65b9\\u8bc1\\u636e\\uff09\\u3002\"}, {\"target_name\": \"phospho-mTOR (Ser2448) antibody\", \"vendor\": \"Cell Signaling Technology\", \"catalog_no\": \"5536\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u6297mTOR\\u53cap-mTOR Ser2448\\uff08Rabbit, Cell Signaling Technology, Cat#2983/5536, 1:1000\\uff09\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u4e3b\\u8981\\u4e00\\u6297\\u5217\\u8868\", \"confidence\": 0.96, \"search_query\": \"phospho-mTOR (Ser2448) antibody Cell Signaling Technology 5536 molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"289 kDa\", \"mw_type\": \"predicted\", \"evidence_url\": \"https://media.cellsignal.com/pdf/5536.pdf\", \"evidence_quote\": \"MW (kDa):   289\", \"retrieval_status\": \"found\", \"notes\": \"Official Cell Signaling Technology product datasheet (PDF) for #5536 lists 'MW (kDa): 289'. The same value appears on the CST product page for #5536. Interpreted as predicted/theoretical molecular weight of full-length mTOR (UniProt P42345) because the datasheet lists a single MW field rather than an observed/expected band size.\"}, {\"target_name\": \"phospho-ULK1 (Ser555) antibody\", \"vendor\": \"Cell Signaling Technology\", \"catalog_no\": \"5869\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u6297ULK1\\u3001p-ULK1 Ser555\\u53caSer757\\uff08Rabbit, Cell Signaling Technology, Cat#8054/5869/14202, 1:1000\\uff09\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u4e3b\\u8981\\u4e00\\u6297\\u5217\\u8868\", \"confidence\": 0.96, \"search_query\": \"phospho-ULK1 (Ser555) antibody Cell Signaling Technology 5869 molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"140-150 kDa; additional bands observed between 90\\u2013100 kDa\", \"mw_type\": \"apparent\", \"evidence_url\": \"http://media.cellsignal.com/pdf/5869.pdf\", \"evidence_quote\": \"Bands of unknown origin are detected between 90 and 100 kDa. Applications: W, IP Reactivity: H M Sensitivity: Endogenous MW (kDa): 140-150\", \"retrieval_status\": \"found\", \"notes\": \"Official Cell Signaling Technology datasheet (PDF for catalog #5869) lists 'MW (kDa): 140-150' as the endogenous molecular weight and also notes additional bands of unknown origin between 90\\u2013100 kDa. MW taken as apparent/observed from manufacturer datasheet.\"}, {\"target_name\": \"HRP-conjugated Goat anti-Rabbit IgG secondary antibody\", \"vendor\": \"Cell Signaling Technology\", \"catalog_no\": \"7074\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u4e0eHRP\\u6807\\u8bb0\\u7684\\u7f8a\\u6297\\u5154\\u6216\\u7f8a\\u6297\\u9f20\\u4e8c\\u6297\\uff08Cell Signaling Technology, Cat#7074/7076, 1:5000\\uff09\\u5ba4\\u6e29\\u907f\\u5149\\u5b75\\u80b21 h\\u3002\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u4e8c\\u6297\", \"confidence\": 0.98, \"search_query\": \"HRP-conjugated Goat anti-Rabbit IgG secondary antibody Cell Signaling Technology 7074 molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"\", \"mw_type\": \"unknown\", \"evidence_url\": \"\", \"evidence_quote\": \"\", \"retrieval_status\": \"not_found\", \"notes\": \"\\u5df2\\u68c0\\u7d22\\u5b98\\u65b9\\u4ea7\\u54c1\\u9875\\u4e0e\\u76f8\\u5173\\u5b98\\u65b9\\u8d44\\u6599\\uff0c\\u4f46\\u672a\\u627e\\u5230\\u660e\\u786e\\u7684\\u5206\\u5b50\\u91cf\\u4fe1\\u606f\\u3002\\u68c0\\u7d22\\u8bb0\\u5f55\\uff1a1) \\u5b98\\u65b9\\u4ea7\\u54c1\\u9875 https://www.cellsignal.com/products/secondary-antibodies/anti-rabbit-igg-hrp-linked-antibody/7074 \\uff08\\u9875\\u9762Product Specifications\\u4e2d\\u6709\\u201cMW (kDa)\\u201d\\u5217\\u4f46\\u672a\\u586b\\u5199\\u5177\\u4f53\\u6570\\u503c\\uff09\\uff1b2) \\u5b98\\u65b9datasheet \\u9875\\u9762 https://www.cellsignal.com/products/7074/datasheet?images=1&protocol=0\\uff08\\u672a\\u89c1\\u5206\\u5b50\\u91cf\\u6587\\u672c\\uff09\\uff1b3) \\u5b98\\u65b9SDS PDF https://media.cellsignal.cn/s3sds/7074-sds-EGHS-EN-20230221104237000.pdf\\uff08SECTION 9: \\u201cMolecular Weight  No information available\\u201d\\uff09\\uff1b\\u6b64\\u5916\\u68c0\\u7d22\\u4e86CiteAb\\u4e0eAntibody Registry\\u6761\\u76ee\\u4ea6\\u65e0\\u5206\\u5b50\\u91cf\\u63cf\\u8ff0\\u3002\\u4f9d\\u636e\\u62bd\\u53d6\\u89c4\\u5219\\uff0c\\u672a\\u627e\\u5230\\u53ef\\u4f5c\\u4e3a\\u8bc1\\u636e\\u7684\\u5206\\u5b50\\u91cf\\u4fe1\\u606f\\uff0c\\u6545\\u8fd4\\u56de not_found\\u3002 \\u641c\\u7d22\\u5173\\u952e\\u8bcd\\uff1a\\\"HRP-conjugated Goat anti-Rabbit IgG secondary antibody Cell Signaling Technology 7074 molecular weight\\\"\\u3002\"}, {\"target_name\": \"HRP-conjugated Goat anti-Mouse IgG secondary antibody\", \"vendor\": \"Cell Signaling Technology\", \"catalog_no\": \"7076\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u4e0eHRP\\u6807\\u8bb0\\u7684\\u7f8a\\u6297\\u5154\\u6216\\u7f8a\\u6297\\u9f20\\u4e8c\\u6297\\uff08Cell Signaling Technology, Cat#7074/7076, 1:5000\\uff09\\u5ba4\\u6e29\\u907f\\u5149\\u5b75\\u80b21 h\\u3002\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u4e8c\\u6297\", \"confidence\": 0.98, \"search_query\": \"HRP-conjugated Goat anti-Mouse IgG secondary antibody Cell Signaling Technology 7076 molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"\", \"mw_type\": \"unknown\", \"evidence_url\": \"\", \"evidence_quote\": \"\", \"retrieval_status\": \"not_found\", \"notes\": \"\\u5df2\\u68c0\\u7d22\\u5b98\\u65b9\\u4ea7\\u54c1\\u9875\\u4e0e\\u76f8\\u5173\\u5b98\\u65b9\\u6587\\u6863\\u4f46\\u672a\\u627e\\u5230\\u660e\\u786e\\u7684\\u5206\\u5b50\\u91cf\\u4fe1\\u606f\\u3002\\u67e5\\u9605\\u5185\\u5bb9\\u5305\\u62ec\\uff1a1) \\u5b98\\u65b9\\u4ea7\\u54c1\\u9875 https://www.cellsignal.com/products/secondary-antibodies/anti-mouse-igg-hrp-linked-antibody/7076\\uff08\\u9875\\u9762Product Specifications\\u4e2d\\u5217\\u6709\\u201cMW (kDa)\\u201d\\u5b57\\u6bb5\\u4f46\\u65e0\\u6570\\u503c\\uff09\\uff1b2) \\u5b98\\u65b9datasheet \\u9875\\u9762 https://www.cellsignal.com/products/7076/datasheet?images=1&protocol=0\\uff1b3) \\u5b98\\u65b9SDS PDF https://media.cellsignal.com/s3sds/7076-sds-EGHS-EN-20230221104301000.pdf\\uff1b\\u53e6\\u67e5\\u9605CiteAb\\u548cAntibody Registry\\u6761\\u76ee\\uff08\\u975e\\u5b98\\u65b9\\u6c47\\u603b\\uff09\\u3002\\u4e0a\\u8ff0\\u5b98\\u65b9\\u6765\\u6e90\\u4e2d\\u5747\\u672a\\u7ed9\\u51fa\\u660e\\u786e\\u7684\\u5206\\u5b50\\u91cf/expected band size\\u4fe1\\u606f\\uff0c\\u6545\\u672a\\u80fd\\u63d0\\u53d6\\u5230\\u53ef\\u8ffd\\u6eaf\\u7684\\u5206\\u5b50\\u91cf\\u6570\\u636e\\u3002\"}, {\"target_name\": \"p62 antibody\", \"vendor\": \"Cell Signaling Technology\", \"catalog_no\": \"8025\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u6297p62\\uff08Rabbit, Cell Signaling Technology, Cat#8025, 1:1000\\uff09\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u4e3b\\u8981\\u4e00\\u6297\\u5217\\u8868\", \"confidence\": 0.98, \"search_query\": \"p62 antibody Cell Signaling Technology 8025 molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"62 kDa\", \"mw_type\": \"band_size\", \"evidence_url\": \"https://www.cellsignal.com/products/primary-antibodies/sqstm1-p62-d5e2-rabbit-monoclonal-antibody/8025\", \"evidence_quote\": \"MW (kDa) | 62\", \"retrieval_status\": \"found\", \"notes\": \"Information obtained from the official Cell Signaling Technology product page for catalog #8025, which lists MW as 62 kDa.\"}, {\"target_name\": \"ULK1 antibody\", \"vendor\": \"Cell Signaling Technology\", \"catalog_no\": \"8054\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u6297ULK1\\u3001p-ULK1 Ser555\\u53caSer757\\uff08Rabbit, Cell Signaling Technology, Cat#8054/5869/14202, 1:1000\\uff09\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u4e3b\\u8981\\u4e00\\u6297\\u5217\\u8868\", \"confidence\": 0.95, \"search_query\": \"ULK1 antibody Cell Signaling Technology 8054 molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"150 kDa\", \"mw_type\": \"band_size\", \"evidence_url\": \"https://media.cellsignal.com/pdf/8054.pdf\", \"evidence_quote\": \"Applications: W, W-S, IP Reactivity: H M R Mk Sensitivity: Endogenous MW (kDa): 150 Source/Isotype: Rabbit IgG\", \"retrieval_status\": \"found\", \"notes\": \"Official Cell Signaling product page and the official datasheet (PDF) for catalog #8054 both list MW (kDa) = 150. mw_type set to 'band_size' because CST presents this as the expected MW in the product specifications.\"}, {\"target_name\": \"phospho-S6K antibody\", \"vendor\": \"Cell Signaling Technology\", \"catalog_no\": \"9205\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u6297p-S6K\\uff08Rabbit, Cell Signaling Technology, Cat#9205, 1:1000\\uff09\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u4e3b\\u8981\\u4e00\\u6297\\u5217\\u8868\", \"confidence\": 0.97, \"search_query\": \"phospho-S6K antibody Cell Signaling Technology 9205 molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"70, 85 kDa\", \"mw_type\": \"band_size\", \"evidence_url\": \"https://www.cellsignal.com/products/primary-antibodies/phospho-p70-s6-kinase-thr389-antibody/9205\", \"evidence_quote\": \"MW (kDa)   70, 85\", \"retrieval_status\": \"found\", \"notes\": \"Cell Signaling Technology product page and official datasheet for catalog #9205 list MW (kDa) as 70, 85 \\u2014 reported as expected/observed band sizes (likely corresponding to p70 S6K isoforms/modified forms). Evidence taken from CST product page/datasheet.\"}, {\"target_name\": \"GAPDH antibody\", \"vendor\": \"Proteintech\", \"catalog_no\": \"60004-1-Ig\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u548c\\u5185\\u53c2GAPDH\\uff08Mouse, Proteintech, Wuhan, China, Cat#60004-1-Ig, 1:5000\\uff09\\u3002\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u5185\\u53c2\\u6297\\u4f53\", \"confidence\": 0.99, \"search_query\": \"GAPDH antibody Proteintech 60004-1-Ig molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"Calculated Molecular Weight | 36 kDa; Observed Molecular Weight | 36 kDa\", \"mw_type\": \"apparent\", \"evidence_url\": \"https://www.ptglab.com/products/GAPDH-Antibody-60004-1-Ig.htm\", \"evidence_quote\": \"Calculated Molecular Weight | 36 kDa; Observed Molecular Weight | 36 kDa\", \"retrieval_status\": \"found\", \"notes\": \"Official Proteintech product page lists both calculated and observed molecular weight as 36 kDa. A matching PDF datasheet (https://www.ptglab.com/products/Pictures/pdf/60004-1-Ig.pdf) contains the same entries. mw_type set to 'apparent' because an observed MW is provided; calculated MW is also included in mw_official.\"}, {\"target_name\": \"ECL ultra-sensitive chemiluminescent substrate\", \"vendor\": \"Thermo Fisher Scientific\", \"catalog_no\": \"32106\", \"lot_no\": \"\", \"mw_in_doc\": \"\", \"context_snippet\": \"\\u6700\\u540e\\uff0c\\u4f7f\\u7528ECL\\u8d85\\u654f\\u5316\\u5b66\\u53d1\\u5149\\u8bd5\\u5242\\uff08Thermo Fisher Scientific, Cat#32106\\uff09\\u663e\\u5f71\", \"location_hint\": \"Western blot\\u65b9\\u6cd5-\\u663e\\u5f71\\u8bd5\\u5242\", \"confidence\": 0.99, \"search_query\": \"ECL ultra-sensitive chemiluminescent substrate Thermo Fisher Scientific 32106 molecular weight kDa datasheet\", \"search_skip\": false, \"mw_official\": \"\", \"mw_type\": \"unknown\", \"evidence_url\": \"\", \"evidence_quote\": \"\", \"retrieval_status\": \"not_found\", \"notes\": \"Searched Thermo Fisher product pages and manuals for Cat. No. 32106: \\u2022 Thermo Fisher product page (https://www.thermofisher.com/order/catalog/product/32106) \\u2014 product is listed as Pierce\\u2122 ECL Western Blotting Substrate (500 mL) but page contains no molecular weight information. \\u2022 Thermo Fisher user guide PDF (https://documents.thermofisher.com/TFS-Assets/LSG/manuals/MAN0011536_Pierce_ECL_West_Blot_Subs_UG.pdf) \\u2014 lists components and storage for Cat. No. 32106 but no molecular weight. \\u2022 Fisher Scientific manual PDF (https://assets.fishersci.com/TFS-Assets/LSG/manuals/MAN0011669_Pierce_Fast_West_Blot_ECL_Subs_UG.pdf) \\u2014 references Cat. No. 32106 but no MW. Conclusion: no official molecular weight found for this reagent/kit. Likely not applicable because this product is a chemical two-component chemiluminescent substrate (mixture), not a single defined protein with a molecular weight. Attempts included searching vendor product pages and official manuals; no MW or expected band size information present.\"}], \"wb_text\": \"\\u5c01\\u95ed\\u540e\\u7684\\u819c\\u4f7f\\u7528TBST\\u6d17\\u6da43\\u6b21\\uff0c\\u6bcf\\u6b2110 min\\uff0c\\u968f\\u540e\\u7f6e\\u4e8e4\\u00b0C\\u6447\\u5e8a\\u4e0a\\u4e0e\\u7279\\u5f02\\u6027\\u4e00\\u6297\\u5b75\\u80b2\\u8fc7\\u591c\\u3002\\u4e3b\\u8981\\u4e00\\u6297\\u5305\\u62ec\\uff1a\\u6297HIF-1\\u03b1\\uff08Rabbit, Cell Signaling Technology, Danvers, MA, USA, Cat#36169, 1:1000\\uff09\\u3001\\u6297HIF-2\\u03b1\\uff08Rabbit, Abcam, Cambridge, UK, Cat#ab199, 1:1000\\uff09\\u3001\\u6297CD41\\uff08Rabbit, Abcam, Cat#ab134131, 1:1000\\uff09\\u3001\\u6297CD61\\uff08Rabbit, Abcam, Cat#ab179475, 1:1000\\uff09\\u3001\\u6297CD42b\\uff08Rabbit, Abcam, Cat#ab183345, 1:1000\\uff09\\u3001\\u6297AMPK\\u53cap-AMPK Thr172\\uff08Rabbit, Cell Signaling Technology, Cat#2532/2535, 1:1000\\uff09\\u3001\\u6297mTOR\\u53cap-mTOR Ser2448\\uff08Rabbit, Cell Signaling Technology, Cat#2983/5536, 1:1000\\uff09\\u3001\\u6297p-S6K\\uff08Rabbit, Cell Signaling Technology, Cat#9205, 1:1000\\uff09\\u3001\\u6297p-4EBP1\\uff08Rabbit, Cell Signaling Technology, Cat#2855, 1:1000\\uff09\\u3001\\u6297ULK1\\u3001p-ULK1 Ser555\\u53caSer757\\uff08Rabbit, Cell Signaling Technology, Cat#8054/5869/14202, 1:1000\\uff09\\uff0c\\u4ee5\\u53ca\\u81ea\\u566c\\u76f8\\u5173\\u86cb\\u767d\\u6297LC3B\\uff08Rabbit, Cell Signaling Technology, Cat#2775, 1:1000\\uff09\\u3001\\u6297p62\\uff08Rabbit, Cell Signaling Technology, Cat#8025, 1:1000\\uff09\\u3001\\u6297Beclin-1\\uff08Rabbit, Cell Signaling Technology, Cat#3495, 1:1000\\uff09\\u3001\\u6297ATG5\\uff08Rabbit, Cell Signaling Technology, Cat#12994, 1:1000\\uff09\\u548c\\u5185\\u53c2GAPDH\\uff08Mouse, Proteintech, Wuhan, China, Cat#60004-1-Ig, 1:5000\\uff09\\u3002\\u6b21\\u65e5\\uff0c\\u819c\\u7ecfTBST\\u6d17\\u6da43\\u6b21\\u540e\\uff0c\\u4e0eHRP\\u6807\\u8bb0\\u7684\\u7f8a\\u6297\\u5154\\u6216\\u7f8a\\u6297\\u9f20\\u4e8c\\u6297\\uff08Cell Signaling Technology, Cat#7074/7076, 1:5000\\uff09\\u5ba4\\u6e29\\u907f\\u5149\\u5b75\\u80b21 h\\u3002\\u6700\\u540e\\uff0c\\u4f7f\\u7528ECL\\u8d85\\u654f\\u5316\\u5b66\\u53d1\\u5149\\u8bd5\\u5242\\uff08Thermo Fisher Scientific, Cat#32106\\uff09\\u663e\\u5f71\\uff0c\\u901a\\u8fc7ChemiDoc MP\\u51dd\\u80f6\\u6210\\u50cf\\u7cfb\\u7edf\\uff08Bio-Rad, Hercules, CA, USA\\uff09\\u91c7\\u96c6\\u5316\\u5b66\\u53d1\\u5149\\u4fe1\\u53f7\\uff0c\\u5e76\\u5229\\u7528ImageJ 1.53\\u8f6f\\u4ef6\\u5bf9\\u76ee\\u6807\\u86cb\\u767d\\u6761\\u5e26\\u8fdb\\u884c\\u7070\\u5ea6\\u5b9a\\u91cf\\u5f52\\u4e00\\u5316\\u5206\\u6790\\uff1a\\u78f7\\u9178\\u5316\\u86cb\\u767d\\u91c7\\u7528\\u5bf9\\u5e94\\u603b\\u86cb\\u767d\\u8fdb\\u884c\\u5f52\\u4e00\\u5316\\uff0c\\u5373p-AMPK/AMPK\\u3001p-mTOR/mTOR\\u3001p-ULK1/ULK1\\uff1b\\u975e\\u78f7\\u9178\\u5316\\u86cb\\u767d\\u91c7\\u7528GAPDH\\u5f52\\u4e00\\u5316\\u3002LC3B\\u86cb\\u767d\\u540c\\u65f6\\u62a5\\u544aLC3B-II/LC3B-I\\u6bd4\\u503c\\u53caLC3B-II/GAPDH\\u6bd4\\u503c\\uff08PMID\\uff1a36555516\\uff09\\u3002\\n\", \"new_wb_text\": \"Western blot\\n\\u5c01\\u95ed\\u540e\\u7684\\u819c\\u4f7f\\u7528TBST\\u6d17\\u6da43\\u6b21\\uff0c\\u6bcf\\u6b2110 min\\uff0c\\u968f\\u540e\\u7f6e\\u4e8e4\\u00b0C\\u6447\\u5e8a\\u4e0a\\u4e0e\\u7279\\u5f02\\u6027\\u4e00\\u6297\\u5b75\\u80b2\\u8fc7\\u591c\\u3002\\u4e3b\\u8981\\u4e00\\u6297\\u5305\\u62ec\\uff1a\\u6297HIF-1\\u03b1\\uff08Rabbit, Cell Signaling Technology, Danvers, MA, USA, Cat#36169, 1:1000\\uff09\\u3001\\u6297HIF-2\\u03b1\\uff08Rabbit, Abcam, Cambridge, UK, Cat#ab199, 100kDa\\uff0c\\uff08HIF-2\\u03b1/EPAS1 antibody | Abcam\\uff09\\uff0c 1:1000\\uff09\\u3001\\u6297CD41\\uff08Rabbit, Abcam, Cat#ab134131, 114kDa\\uff0c\\uff08Integrin alpha-IIb (CD41) antibody | Abcam\\uff09\\uff0c 1:1000\\uff09\\u3001\\u6297CD61\\uff08Rabbit, Abcam, Cat#ab179475, 113kDa\\uff0c\\uff08Integrin alpha V antibody | Abcam\\uff09\\uff0c 1:1000\\uff09\\u3001\\u6297CD42b\\uff08Rabbit, Abcam, Cat#ab183345, 1:1000\\uff09\\u3001\\u6297AMPK\\u53cap-AMPK Thr172\\uff08Rabbit, Cell Signaling Technology, Cat#2532/2535, 62kDa\\uff0c\\uff08AMPK\\u03b1 and phospho-AMPK\\u03b1 (Thr172) antibodies | Cell Signaling Technology\\uff09\\uff0c 1:1000\\uff09\\u3001\\u6297mTOR\\u53cap-mTOR Ser2448\\uff08Rabbit, Cell Signaling Technology, Cat#2983/5536, 289kDa\\uff0c\\uff08mTOR and phospho-mTOR (Ser2448) antibodies | Cell Signaling Technology\\uff09\\uff0c 1:1000\\uff09\\u3001\\u6297p-S6K\\uff08Rabbit, Cell Signaling Technology, Cat#9205, 70\\u201385kDa\\uff0c\\uff08Phospho-p70 S6 Kinase (Thr389) antibody | Cell Signaling Technology\\uff09\\uff0c 1:1000\\uff09\\u3001\\u6297p-4EBP1\\uff08Rabbit, Cell Signaling Technology, Cat#2855, 1:1000\\uff09\\u3001\\u6297ULK1\\u3001p-ULK1 Ser555\\u53caSer757\\uff08Rabbit, Cell Signaling Technology, Cat#8054/5869/14202, 140\\u2013150kDa\\uff0c\\uff08ULK1 and phospho-ULK1 (Ser555/Ser757) antibodies | Cell Signaling Technology\\uff09\\uff0c 1:1000\\uff09\\uff0c\\u4ee5\\u53ca\\u81ea\\u566c\\u76f8\\u5173\\u86cb\\u767d\\u6297LC3B\\uff08Rabbit, Cell Signaling Technology, Cat#2775, 14\\u201316kDa\\uff0c\\uff08LC3B (LC3B-I/II) antibody | Cell Signaling Technology\\uff09\\uff0c 1:1000\\uff09\\u3001\\u6297p62\\uff08Rabbit, Cell Signaling Technology, Cat#8025, 62kDa\\uff0c\\uff08SQSTM1/p62 antibody | Cell Signaling Technology\\uff09\\uff0c 1:1000\\uff09\\u3001\\u6297Beclin-1\\uff08Rabbit, Cell Signaling Technology, Cat#3495, 60kDa\\uff0c\\uff08Beclin-1 antibody | Cell Signaling Technology\\uff09\\uff0c 1:1000\\uff09\\u3001\\u6297ATG5\\uff08Rabbit, Cell Signaling Technology, Cat#12994, 55kDa\\uff0c\\uff08ATG5 antibody | Cell Signaling Technology\\uff09\\uff0c 1:1000\\uff09\\u548c\\u5185\\u53c2GAPDH\\uff08Mouse, Proteintech, Wuhan, China, Cat#60004-1-Ig, 36kDa\\uff0c\\uff08GAPDH antibody | Proteintech\\uff09\\uff0c 1:5000\\uff09\\u3002\\u6b21\\u65e5\\uff0c\\u819c\\u7ecfTBST\\u6d17\\u6da43\\u6b21\\u540e\\uff0c\\u4e0eHRP\\u6807\\u8bb0\\u7684\\u7f8a\\u6297\\u5154\\u6216\\u7f8a\\u6297\\u9f20\\u4e8c\\u6297\\uff08Cell Signaling Technology, Cat#7074/7076, 1:5000\\uff09\\u5ba4\\u6e29\\u907f\\u5149\\u5b75\\u80b21 h\\u3002\\u6700\\u540e\\uff0c\\u4f7f\\u7528ECL\\u8d85\\u654f\\u5316\\u5b66\\u53d1\\u5149\\u8bd5\\u5242\\uff08Thermo Fisher Scientific, Cat#32106\\uff09\\u663e\\u5f71\\uff0c\\u901a\\u8fc7ChemiDoc MP\\u51dd\\u80f6\\u6210\\u50cf\\u7cfb\\u7edf\\uff08Bio-Rad, Hercules, CA, USA\\uff09\\u91c7\\u96c6\\u5316\\u5b66\\u53d1\\u5149\\u4fe1\\u53f7\\uff0c\\u5e76\\u5229\\u7528ImageJ 1.53\\u8f6f\\u4ef6\\u5bf9\\u76ee\\u6807\\u86cb\\u767d\\u6761\\u5e26\\u8fdb\\u884c\\u7070\\u5ea6\\u5b9a\\u91cf\\u5f52\\u4e00\\u5316\\u5206\\u6790\\uff1a\\u78f7\\u9178\\u5316\\u86cb\\u767d\\u91c7\\u7528\\u5bf9\\u5e94\\u603b\\u86cb\\u767d\\u8fdb\\u884c\\u5f52\\u4e00\\u5316\\uff0c\\u5373p-AMPK/AMPK\\u3001p-mTOR/mTOR\\u3001p-ULK1/ULK1\\uff1b\\u975e\\u78f7\\u9178\\u5316\\u86cb\\u767d\\u91c7\\u7528GAPDH\\u5f52\\u4e00\\u5316\\u3002LC3B\\u86cb\\u767d\\u540c\\u65f6\\u62a5\\u544aLC3B-II/LC3B-I\\u6bd4\\u503c\\u53caLC3B-II/GAPDH\\u6bd4\\u503c\\uff08PMID\\uff1a36555516\\uff09\\u3002\", \"doc_url\": \"http://218.77.106.67:64110/files/e2f0cb1a-b2b3-4bc5-8cd8-2022880ccf26/file-preview?timestamp=1780470298&nonce=459d0bb9cf928e65f2573b160595acb3&sign=CbSpNE-owue3zptbsJDGc_cVTs_Ros43p8p1HY8pPlA=\", \"doc_id\": null, \"file_name\": \"XG(JAN-26-046)\\u5206\\u5b50\\u91cf.docx\"}"

payload = json.loads(payload_json)

records = payload["antibody_json"]
wb_text = payload["wb_text"]
new_wb_text = payload["new_wb_text"]
doc_url = payload["doc_url"]
doc_id = payload["doc_id"]
file_name = payload["file_name"]

enriched = enrich_wb_text_rule(wb_text, records)

if wb_text and not new_wb_text:
    new_wb_text = enriched

logs: list[str] = []

if not new_wb_text and enriched:
    new_wb_text = enriched
    logs.append("new_wb_text 为空，已使用规则引擎 enriched_wb_text")
elif (
    new_wb_text
    and wb_text
    and count_matches(new_wb_text, records) < count_matches(enriched, records)
):
    new_wb_text = enriched
    logs.append("LLM 嵌入条数少于规则引擎，已回退为 enriched_wb_text")

raw = load_doc_bytes(doc_url)
updated, rep_log = replace_western_blot_in_docx(raw, wb_text, new_wb_text)
logs.extend(rep_log)

mw_bytes = build_mw_list_docx(records, doc_id)

file_name_mane_docx = file_name
mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
       
