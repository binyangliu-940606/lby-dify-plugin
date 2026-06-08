# -*- coding: utf-8 -*-
import io
import re
import json

from collections.abc import Generator
from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage

from typing import Any
from docx import Document
from docx.shared import Pt
# ---------------------------------------------------------------------------
# Tool 入口
# ---------------------------------------------------------------------------


class LbyToolsTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage, None, None]:
        payload = json.loads(tool_parameters["payload_json"])

        records = payload["antibody_json"]
        doc_id = payload["doc_id"]

        mw_bytes = build_mw_list_docx(records, doc_id)
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        file_name_wm_list = "单独分子量文件.docx"
        yield self.create_blob_message(
            blob=mw_bytes,
            meta={
                "mime_type": mime_type,
                "filename": file_name_wm_list,
            },
        )

# ---------------------------------------------------------------------------
# 分子量列表 DOCX
# ---------------------------------------------------------------------------
def build_mw_list_docx(records: list[dict], doc_id: str) -> bytes:
    doc = Document()
    title = doc.add_paragraph()
    title_run = "XG 分子量"
    if doc_id:
        title_run = f"XG({doc_id})分子量"
    run = title.add_run(title_run)
    run.bold = True
    run.font.size = Pt(16)

    for rec in records:
        name = display_antibody_name(rec)
        cat = (rec.get("catalog_no") or "").strip().lstrip("#")
        vendor = rec.get("vendor") or ""

        mw = (rec.get("mw_short") or "").strip() 
        if not mw:
            mw = shorten_mw(rec.get("mw_official", ""))
        if not mw:    
            mw = "未检索"
        
        # GAPDH 等有时不写「抗体」
        ab_word = "抗体" if "igg" not in name.lower() and "抗体" not in name.lower() else ""
        if ab_word:
            name = f"{name} 抗体"

        line = f"{name}：货号 {cat}, 分子量：{mw}, 品牌：{vendor}"
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def display_antibody_name(rec: dict, fallback_from_text: str = "") -> str:
    """分子量列表用的短名称。"""
    name = (rec.get("target_name") or "").strip()
    name = re.sub(r"\s+antibody\s*$", "", name, flags=re.I)
    name = re.sub(r"\s+antibod(y|ies)\s*$", "", name, flags=re.I)
    if name:
        return name
    return fallback_from_text or "Unknown"


# ---------------------------------------------------------------------------
# 分子量精简（写入正文 & 列表）
# ---------------------------------------------------------------------------
def shorten_mw(mw_official: str) -> str:
    """从 mw_official 提取适合写入 WB 段落/列表的简短 kDa 描述。"""
    s = (mw_official or "").strip()
    if not s:
        return ""
    return s








