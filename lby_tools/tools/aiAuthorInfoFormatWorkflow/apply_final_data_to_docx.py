# -*- coding: utf-8 -*-
import io
import os
import re
import json
import zipfile

from collections.abc import Generator
from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage

from datetime import datetime, timezone
from typing import Any, Optional, Tuple

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor
from lxml import etree as LET

# ---------------------------------------------------------------------------
# Tool 入口
# ---------------------------------------------------------------------------

class LbyToolsTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage, None, None]:
        payload = json.loads(tool_parameters["payload_json"])
       
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        final_data_json = payload["final_data_json"]
        mode = payload["mode"]
        file_name_clean = payload["file_name_clean"]
        file_name_track = payload["file_name_track"]
        old_data_json = payload["old_data_json"]

        out_bytes_clear,out_bytes_track = generate_title_page_files(
            final_data_json,                      # dict
            mode=mode,                 # 或其它字符串则只出 Clean
            old_data_json=old_data_json,         # revise_format 时用于生成批注对比
        )
        
        yield self.create_blob_message(
            blob=out_bytes_clear,
            meta={
                "mime_type": mime_type,
                "filename": file_name_clean,
            },
        )

        if mode=='revise_format':
            yield self.create_blob_message(
                blob=out_bytes_track,
                meta={
                    "mime_type": mime_type,
                    "filename": file_name_track,
                },
            )

# -*- coding: utf-8 -*-
"""
根据 final_data_json / old_data_json 生成论文标题页作者信息区 Word 文档。

- mode == "revise_format": 输出 Clean + 痕迹版（内容一致，痕迹版带 Word 批注说明相对 old 的增删改）
- 其他 mode: 仅输出 Clean

依赖: python-docx, lxml
"""

# Word 批注命名空间
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NSMAP = {"w": W_NS}
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"

_DIGIT_SUPER = str.maketrans("0123456789", "0123456789")
LINK_BLUE_HEX = "0000FF"


def _sanitize(s: str | None) -> str:
    if s is None:
        return ""
    return re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", str(s))

def _unicode_superscript_num(n: int) -> str:
    return str(n).translate(_DIGIT_SUPER)

def _collect_affiliations_final(final_data: dict[str, Any]) -> list[tuple[int, str, dict]]:
    """返回 [(num, formatted_line, raw), ...] 按 num 排序。"""
    by_num: dict[int, tuple[str, dict]] = {}
    for line in final_data.get("affiliations") or []:
        num = int(line.get("idx") or 0)

        department = line.get('department','')
        institution = line.get('institution','')
        city = line.get('city','')
        country = line.get('country','')
        province = line.get('province','')
        postal_code = line.get('postal_code','')
        street = line.get('street','')
        tel = line.get('tel','')
        
        arr_instruction = []
        if department:
            arr_instruction.append(f"{department}")
        
        if institution:
            arr_instruction.append(f"{institution}")
        
        if city or postal_code:
            arr_instruction.append(f"{city} {postal_code}")

        if province:
            arr_instruction.append(f"{province}")

        if country:
            arr_instruction.append(f"{country}")

        txt = ', '.join(arr_instruction)

        by_num[num] = (_sanitize(txt), line)

    return [(n, by_num[n][0]) for n in sorted(by_num.keys())]


def _build_author_line_runs_final(final_data: dict[str, Any]) -> str:
    """纯文本作者行（用于 diff）；上标以 Unicode 数字与 *# 字符表示。"""
    authors = final_data.get("authors") or []
    chunks: list[str] = []
    for a in authors:
        name = _sanitize(a.get("full_name") or "")
        if not name:
            continue
        nums = a.get("affiliation_refs") or 0
        # sup = "".join(_unicode_superscript_num(int(x)) for x in nums if x is not None)
        sup = nums
        marks = list(a.get("footnote_marks") or [])
        if a.get("is_corresponding").lower() in ("yes", "true", "1") and '*' not in marks:
            marks.append("*")
        if a.get("is_cofirst").lower() in ("yes", "true", "1") and '#' not in marks:
            marks.append("#")
        mark_s = ""
        for m in marks:
            m = str(m).strip()
            if m in ("*", "#", "†", "§"):
                mark_s += m
        chunks.append(f"{name}{sup}{mark_s}")
    return ", ".join(chunks)

def _add_hyperlink(paragraph, url: str, text: str, color_hex: str | None = LINK_BLUE_HEX) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)

    if color_hex:
        hx = color_hex.strip().lstrip("#")
        if len(hx) == 6:
            c = OxmlElement("w:color")
            c.set(qn("w:val"), hx.upper())
            r_pr.append(c)

    new_run.append(r_pr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _set_run_superscript(run, on: bool = True) -> None:
    # 设置上标状态
    run.font.superscript = on
    
    run.font.bold = True
    # 上标颜色设为 RGB(0,0,255)
    run.font.color.rgb = RGBColor(0, 0, 255)


def _populate_author_paragraph(paragraph, final_data: dict[str, Any]) -> None:
    authors = final_data.get("authors") or []
    has_cofirst = False

    count_mark = 0
    for i, a in enumerate(authors):
        marks = list(a.get("footnote_marks") or [])
        if '#' in marks:
            count_mark += 1
            continue

        if a.get("is_cofirst").lower() in ("yes", "true", "1"):
            count_mark += 1
            continue

    for i, a in enumerate(authors):
        family_name = _sanitize(a.get("family_name") or "")
        given_name = _sanitize(a.get("given_name") or "")
        if given_name:
            given_name = "".join(given_name.split()).capitalize()

        name = _sanitize(a.get("full_name") or "")

        if family_name and given_name:
            name = f"{given_name} {family_name}"
        
        if not name:
            continue
        if i:
            paragraph.add_run(", ")

        r_name = paragraph.add_run(name)
        r_name.bold = True

        # nums = a.get("affiliation_refs") or []
        # for num in nums:
        #     r = paragraph.add_run(_unicode_superscript_num(int(num)))
        #     _set_run_superscript(r)
        num = a.get("affiliation_refs")
        r = paragraph.add_run(_unicode_superscript_num(int(num)))
        _set_run_superscript(r)

        marks = list(a.get("footnote_marks") or [])
        if a.get("is_corresponding").lower() in ("yes", "true", "1") and '*' not in marks:
            marks.append("*")
        if count_mark>1:
            if a.get("is_cofirst").lower() in ("yes", "true", "1") and '#' not in marks:
                marks.append("#")
        for m in marks:
            m = str(m).strip()
            if not m:
                continue
            r = paragraph.add_run(m)
            _set_run_superscript(r)
        
        if count_mark>1:
            # if '#' in marks:
            has_cofirst = True
    return has_cofirst


def _populate_affiliation_paragraph(paragraph, num: int, line_text: str) -> None:
    r_sup = paragraph.add_run(_unicode_superscript_num(num))
    _set_run_superscript(r_sup)
    paragraph.add_run(" ")
    paragraph.add_run(_sanitize(line_text))


def _correspondence_paragraphs_body(final_data: dict[str, Any],cb) -> tuple[str, str, str, str]:
    """返回 (通讯作者后的地址整句（不含姓名）, 邮箱, E-mail 标签, Tel 整行)。"""
    
    contact_person = cb.get('contact_person') or ''
    org = cb.get('org') or ''
    email = cb.get('email') or ''
    tel = cb.get('tel') or ''
    street = cb.get('street') or ''
    city = cb.get('city') or ''
    province = cb.get('province') or ''
    postal_code = cb.get('postal_code') or ''
    country = cb.get('country') or ''


    authors = final_data.get('authors') or []
    affiliation_refs = 0
    for author in authors:
        author_name = author.get('full_name') or ''
        if author_name!='' and contact_person==author_name:
            affiliation_refs = author.get('affiliation_refs') or 0
            try:
                affiliation_refs = int(affiliation_refs)
            except ValueError:
                affiliation_refs = 0
    
    if affiliation_refs!=0:
        affiliation_list = final_data.get('affiliations') or []
        for affiliation in affiliation_list:
            # 方式3：有默认值
            idx = affiliation.get('idx') or 0
            try:
                idx = int(idx)
            except ValueError:
                idx = 0
            if idx==affiliation_refs:
                #单位部门优先取instraction中的值
                if affiliation.get('institution') and affiliation.get('department'):
                    org = f'{affiliation.get('department')}, {affiliation.get('institution')}'

                if not tel:
                    tel = affiliation.get('tel')

                if not street:
                    street = affiliation.get('street')

                if not city:
                    city = affiliation.get('city')

                if not province:
                    province = affiliation.get('province')

                if not postal_code:
                    postal_code = affiliation.get('postal_code')

                if not country:
                    country = affiliation.get('country')
    
    # 通讯地址：科室，单位，详细地址（什么路，什么区之类的），城市+邮编，省份，国家
    corr_info = []
    if org:
        corr_info.append(org)
    if street:
        corr_info.append(street)
    if city or province or postal_code:
        corr_info.append(f'{city} {postal_code}')
    if province:
        corr_info.append(province)
    if country:
        corr_info.append(country)

    if tel:
        if not bool(re.match(r"^\+\s*86", str(tel).strip())):
            tel = f"+86-{tel}" 

    return {contact_person:f'{', '.join(corr_info)}.'},{contact_person:email},tel


def _funding_text_final(final_data: dict[str, Any]) -> str:
    items = final_data.get("funding") or []
    funding_len = len(items)
    parts = ""
    for i in range(funding_len):
        it = items[i]
        agency = _sanitize((it or {}).get("agency"))
        grant_no = _sanitize((it or {}).get("grant_no"))

        if i==0:
            parts += "This study was supported by "
            if not agency.strip().lower().startswith("the"):
                parts += "the "

        if not grant_no:
            continue
        
        if not agency:
            agency = "XXX"

        parts += f"{agency} (No. {grant_no})"
        if i==funding_len-1:
            parts += "."
        else:
            parts += ", "
    return parts

def build_title_page_document(final_data: dict[str, Any]) -> Document:
    # helper to set paragraph spacing to 0
    def _zero_para_spacing(p):
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        return p

    doc = Document()

    # 设置页边距为 2cm
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # 设置默认样式：英文 Times New Roman，中文（East Asia）宋体，字号 12pt，段落行距双倍
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    # 设置中文字体（East Asia）
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    # 设置默认段落双倍行距
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    # 设置默认段前段后为 0
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # 作者行
    p_auth = _zero_para_spacing(doc.add_paragraph())
    p_auth.paragraph_format.alignment = 0  # LEFT
    p_auth.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    has_cofirst = _populate_author_paragraph(p_auth, final_data)

    # 单位
    list_instruction = _collect_affiliations_final(final_data)
    for num, line_text in list_instruction:
        p = _zero_para_spacing(doc.add_paragraph())
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        _populate_affiliation_paragraph(p, num, line_text)

    if has_cofirst:
        p_jing = _zero_para_spacing(doc.add_paragraph())
        p_jing.paragraph_format.alignment = 0  # LEFT
        p_jing.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        r_jing = p_jing.add_run('#')
        _set_run_superscript(r_jing)
        p_jing.add_run(" ")
        p_jing.add_run(_sanitize('These authors are regarded as co-first authors'))


    # 小的空行（仍设为段前段后为0）
    tmp = _zero_para_spacing(doc.add_paragraph())


    
    # 准备corr数据
    corr_author_list = []
    corr_email_list = []
    corr_tels = ''
    cb_list = final_data.get("correspondence") or []

    for cb in cb_list:
        corr_author, corr_email, corr_tels = _correspondence_paragraphs_body(final_data,cb)
        
        corr_author_list.append(corr_author)
        corr_email_list.append(corr_email)

    #写入corr内容
    p_corr = _zero_para_spacing(doc.add_paragraph())
    p_corr.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    r_ast = p_corr.add_run("*")
    _set_run_superscript(r_ast)
    p_corr.add_run(" ")
    r_lbl = p_corr.add_run("Correspondence to:")
    r_lbl.bold = True


    #联系人
    for corr_author in corr_author_list:
        p_author_b = _zero_para_spacing(doc.add_paragraph())
        p_author_b.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        contact_person, corr_org = next(iter(corr_author.items()))
        r_author = p_author_b.add_run(contact_person)
        r_author.bold = True
        p_author_b.add_run(", ")
        p_author_b.add_run(corr_org)

    #email
    p_corr_email = _zero_para_spacing(doc.add_paragraph())
    p_corr_email.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    p_email_label = p_corr_email.add_run("E-mail: ")
    p_email_label.bold = True

    i_e = 0
    for corr_email in corr_email_list:
        contact_person, email = next(iter(corr_email.items()))
            
        if email:
            href = email if email.lower().startswith("mailto:") else f"mailto:{email}"
            _add_hyperlink(p_corr_email, href, email, LINK_BLUE_HEX)
        else:
            p_corr_email.add_run("")

        if i_e==len(corr_email_list)-1:
            p_corr_email.add_run(f" ({contact_person})")
        else:
            p_corr_email.add_run(f" ({contact_person}), ")
        i_e += 1

    #tel
    if corr_tels:
        p_corr_tel = _zero_para_spacing(doc.add_paragraph())
        p_corr_tel.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p_tel_label = p_corr_tel.add_run("Tel: ")
        p_tel_label.bold = True

        p_corr_tel.add_run(corr_tels)

    # 另一空行
    tmp2 = _zero_para_spacing(doc.add_paragraph())

    p_fund_h = _zero_para_spacing(doc.add_paragraph())
    p_fund_h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    r_f = p_fund_h.add_run("Funding")
    r_f.bold = True

    fund_txt = _funding_text_final(final_data)
    if fund_txt:
        p_fund_b = _zero_para_spacing(doc.add_paragraph())
        p_fund_b.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p_fund_b.add_run(fund_txt)

    return doc

# ---------------------------------------------------------------------------
# old_data_json → 纯文本快照（用于对比）
# ---------------------------------------------------------------------------


def _old_author_line(old: dict[str, Any]) -> str:
    authors = old.get("authors") or []
    chunks = []
    for au in authors:
        name = _sanitize(au.get("full_name"))
        nums = au.get("affiliation_refs") or 0
        # sup = "".join(_unicode_superscript_num(int(n)) for n in nums)
        sup = nums
        marks = au.get("footnote_marks") or []
        ic = au.get("is_corresponding")
        if isinstance(ic, str) and ic.lower() in ("yes", "true", "1"):
            if "*" not in marks:
                marks = list(marks) + ["*"]
        elif ic is True and "*" not in marks:
            marks = list(marks) + ["*"]
        mark_s = "".join(str(m) for m in marks if m)
        chunks.append(f"{name}{sup}{mark_s}")
    return ", ".join(chunks)


def _old_affiliation_block(old: dict[str, Any]) -> str:
    lines = []
    for a in sorted(old.get("affiliations") or [], key=lambda x: int(x.get("idx") or 0)):
        num = int(a.get("idx") or 0)
        base = _sanitize(a.get("name_en") or a.get("institution"))
        addr = _sanitize(a.get("address"))
        city = _sanitize(a.get("city"))
        st = _sanitize(a.get("state"))
        pc = _sanitize(a.get("postal_code"))
        co = _sanitize(a.get("country"))
        tail = ", ".join(x for x in [addr, city, st, pc, co] if x)
        if base and tail:
            lines.append(f"{_unicode_superscript_num(num)} {base}, {tail}")
        elif base:
            lines.append(f"{_unicode_superscript_num(num)} {base}")
    return "\n".join(lines)


def _old_correspondence_block(old: dict[str, Any]) -> str:
    correspondence = old.get("correspondence") or []
    parts_list = []
    for c in correspondence:
        parts = []
        name = _sanitize(c.get("contact_person"))
        org = _sanitize(c.get("org"))
        addr = _sanitize(c.get("address"))
        city = _sanitize(c.get("city"))
        st = _sanitize(c.get("state"))
        pc = _sanitize(c.get("postal_code"))
        co = _sanitize(c.get("country"))
        email = _sanitize(c.get("email"))
        tel = _sanitize(c.get("tel"))
        line1 = f"* Correspondence to: {name}, {org}"
        if addr or city or pc:
            line1 += ", " + ", ".join(x for x in [addr, city, st, pc, co] if x)
        parts.append(line1)
        parts.append(f"E-mail: {email}")
        parts.append(f"Tel.: {tel}" if tel else "Tel.:")
        parts_list.append("\n".join(parts))
    return "\n".join(parts_list)


def _old_funding_block(old: dict[str, Any]) -> str:
    items = old.get("funding") or []
    return "\n".join(_sanitize((x or {}).get("raw")) for x in items if (x or {}).get("raw"))


def snapshot_old_for_diff(old: dict[str, Any] | None) -> dict[str, str]:
    if not old:
        return {"authors": "", "affiliations": "", "correspondence": "", "funding": ""}
    return {
        "authors": _old_author_line(old),
        "affiliations": _old_affiliation_block(old),
        "correspondence": _old_correspondence_block(old),
        "funding": _old_funding_block(old),
    }

def snapshot_new_for_diff(final_data: dict[str, Any]) -> dict[str, str]:
    aff_lines = []
    for num, line_text in _collect_affiliations_final(final_data):
        aff_lines.append(f"{_unicode_superscript_num(num)} {_sanitize(line_text)}")
    correspondence_list = final_data.get("correspondence") or []
    c_lines_list = []
    for cb in correspondence_list:
        corr_author, corr_email, corr_tels = _correspondence_paragraphs_body(final_data,cb)

        name, middle = next(iter(corr_author.items()))
        email_label, email = next(iter(corr_email.items()))
        tel_line = corr_tels
        c_lines = [
            f"* Correspondence to: {name}, {middle}".replace(", ,", ",").strip().rstrip(","),
            (email_label + " " + email).strip(),
            tel_line,
        ]
        c_lines_list.append("\n".join(c_lines))

    return {
        "authors": _build_author_line_runs_final(final_data),
        "affiliations": "\n".join(aff_lines),
        "correspondence": "\n".join(c_lines_list),
        "funding": _funding_text_final(final_data),
    }


def _diff_comment(old_s: str, new_s: str) -> str | None:
    o = (old_s or "").strip()
    n = (new_s or "").strip()
    if o == n:
        return None
    if not o and n:
        return "【增加】\n" + n
    if o and not n:
        return "【删除】\n" + o
    return "【修改】\n原：" + o + "\n新：" + n


def collect_section_comments(
    final_data: dict[str, Any], old_data: dict[str, Any] | None
) -> list[tuple[str, str]]:
    """返回 [(section_key, comment_text), ...]"""
    old_snap = snapshot_old_for_diff(old_data)
    new_snap = snapshot_new_for_diff(final_data)
    order = ["authors", "affiliations", "correspondence", "funding"]
    out: list[tuple[str, str]] = []
    for key in order:
        msg = _diff_comment(old_snap.get(key, ""), new_snap.get(key, ""))
        if msg:
            out.append((key, msg))
    return out


# ---------------------------------------------------------------------------
# OOXML 批注注入（痕迹版）
# ---------------------------------------------------------------------------


def _next_rid(rels_root: LET._Element) -> str:
    max_id = 0
    for rel in rels_root:
        rid = rel.get("Id") or ""
        if rid.startswith("rId"):
            try:
                max_id = max(max_id, int(rid[3:]))
            except ValueError:
                pass
    return f"rId{max_id + 1}"


def _patch_content_types(zf: zipfile.ZipFile) -> bytes:
    data = zf.read("[Content_Types].xml")
    root = LET.fromstring(data)
    part = "/word/comments.xml"
    exists = False
    for ov in root.iter():
        if ov.tag.endswith("Override") and ov.get("PartName") == part:
            exists = True
            break
    if not exists:
        ov = LET.Element(f"{{{CT_NS}}}Override")
        ov.set("PartName", part)
        ov.set(
            "ContentType",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
        )
        root.append(ov)
    return LET.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _merge_comments_xml(existing_bytes: bytes | None, new_fragment: LET._Element) -> bytes:
    if not existing_bytes:
        return LET.tostring(new_fragment, xml_declaration=True, encoding="UTF-8", standalone=True)
    root = LET.fromstring(existing_bytes)
    for child in list(new_fragment):
        root.append(child)
    return LET.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _paragraph_indices_for_sections(num_aff: int) -> dict[str, int]:
    """
    与 build_title_page_document 段落顺序一致（估算，供基于段落索引的批注使用，痕迹版依然可用）。
    0 作者, 1..1+num_aff-1 单位, 然后可能有 co-first 行, 空行, correspondence 等。
    由于有些段落是可选的（co-first），基于内容的定位更稳健（例如我们对 funding 的 XXX 使用内容搜索）。
    """
    i = 0
    m: dict[str, int] = {}
    m["authors"] = i
    i += 1
    # affiliations start immediately after authors
    m["affiliations"] = i
    i += num_aff
    # after affiliations, optional co-first (unknown) and one blank; rough estimate:
    i += 2
    m["correspondence"] = i
    i += 3
    i += 2
    m["funding"] = i
    return m

def inject_comments_into_docx_bytes(
    docx_bytes: bytes,
    section_messages: list[tuple[str, str]],
    num_affiliations: int,
    author_display: str = "格式修订",
) -> bytes:
    """在内存中向 docx 注入批注；section_messages: [(authors|affiliations|correspondence|funding, text)]"""
    if not section_messages:
        return docx_bytes

    indices = _paragraph_indices_for_sections(num_affiliations)
    sec_to_idx = {k: indices[k] for k in ["authors", "affiliations", "correspondence", "funding"]}

    buf = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zin:
        names = set(zin.namelist())
        rels_path = "word/_rels/document.xml.rels"

        rels_root = LET.fromstring(zin.read(rels_path))
        new_rid = _next_rid(rels_root)
        has_rel = False
        for rel in rels_root:
            if (rel.get("Type") or "").endswith("/comments"):
                has_rel = True
                new_rid = rel.get("Id")
                break
        if not has_rel:
            rel_el = LET.Element(f"{{{REL_NS}}}Relationship")
            rel_el.set("Id", new_rid)
            rel_el.set(
                "Type",
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
            )
            rel_el.set("Target", "comments.xml")
            rels_root.append(rel_el)

        doc_root = LET.fromstring(zin.read("word/document.xml"))
        body = doc_root.find(f".//{{{W_NS}}}body")
        if body is None:
            return docx_bytes

        old_comments = zin.read("word/comments.xml") if "word/comments.xml" in names else None
        max_id = -1
        if old_comments:
            cr = LET.fromstring(old_comments)
            for c in cr:
                if c.tag == f"{{{W_NS}}}comment":
                    try:
                        max_id = max(max_id, int(c.get(f"{{{W_NS}}}id") or -1))
                    except ValueError:
                        pass

        new_comments_root = LET.Element(f"{{{W_NS}}}comments", nsmap=_NSMAP)
        cid = max_id + 1

        for sec, msg in section_messages:
            p_idx = sec_to_idx.get(sec)
            if p_idx is None:
                continue
            paras = [c for c in body if c.tag == f"{{{W_NS}}}p"]
            if p_idx >= len(paras):
                # 如果估算索引超出，则尝试跳过该项
                continue
            p_el = paras[p_idx]

            # create comment element
            c_el = LET.SubElement(new_comments_root, f"{{{W_NS}}}comment")
            c_el.set(f"{{{W_NS}}}id", str(cid))
            c_el.set(f"{{{W_NS}}}author", author_display)
            c_el.set(
                f"{{{W_NS}}}date",
                datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
            )
            c_el.set(f"{{{W_NS}}}initials", "RV")
            cp = LET.SubElement(c_el, f"{{{W_NS}}}p")
            cr_el = LET.SubElement(cp, f"{{{W_NS}}}r")
            ct = LET.SubElement(cr_el, f"{{{W_NS}}}t")
            ct.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            ct.text = msg

            start = LET.Element(f"{{{W_NS}}}commentRangeStart")
            start.set(f"{{{W_NS}}}id", str(cid))
            end = LET.Element(f"{{{W_NS}}}commentRangeEnd")
            end.set(f"{{{W_NS}}}id", str(cid))
            ref_run = LET.Element(f"{{{W_NS}}}r")
            ref = LET.SubElement(ref_run, f"{{{W_NS}}}commentReference")
            ref.set(f"{{{W_NS}}}id", str(cid))

            # 简单策略：在目标段落整体前后插入批注范围（因为这里主要用于痕迹版较大段落批注）
            p_el.insert(0, start)
            p_el.append(end)
            p_el.append(ref_run)

            cid += 1

        merged_comments = _merge_comments_xml(old_comments, new_comments_root)

        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                name = item.filename
                if name in (rels_path, "word/document.xml", "word/comments.xml", "[Content_Types].xml"):
                    continue
                zout.writestr(item, zin.read(name))

            zout.writestr(rels_path, LET.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True))
            zout.writestr("word/document.xml", LET.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True))
            zout.writestr("word/comments.xml", merged_comments)

            zout.writestr("[Content_Types].xml", _patch_content_types(zin))

    return buf.getvalue()

def inject_comments_to_runs_matching_text(
    docx_bytes: bytes,
    target: str,
    max_count: int,
    comment_text: str,
    author_display: str = "格式修订",
) -> bytes:
    """
    在文档中搜索包含 target 的 run (<w:t>)，对每个匹配的 run 精确插入一条批注（最多 max_count 条）。
    返回修改后的 docx bytes。
    """
    if max_count <= 0:
        return docx_bytes

    buf = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zin:
        names = set(zin.namelist())
        rels_path = "word/_rels/document.xml.rels"

        rels_root = LET.fromstring(zin.read(rels_path))
        new_rid = _next_rid(rels_root)
        has_rel = False
        for rel in rels_root:
            if (rel.get("Type") or "").endswith("/comments"):
                has_rel = True
                new_rid = rel.get("Id")
                break
        if not has_rel:
            rel_el = LET.Element(f"{{{REL_NS}}}Relationship")
            rel_el.set("Id", new_rid)
            rel_el.set(
                "Type",
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
            )
            rel_el.set("Target", "comments.xml")
            rels_root.append(rel_el)

        doc_root = LET.fromstring(zin.read("word/document.xml"))
        body = doc_root.find(f".//{{{W_NS}}}body")
        if body is None:
            return docx_bytes

        old_comments = zin.read("word/comments.xml") if "word/comments.xml" in names else None
        max_id = -1
        if old_comments:
            cr = LET.fromstring(old_comments)
            for c in cr:
                if c.tag == f"{{{W_NS}}}comment":
                    try:
                        max_id = max(max_id, int(c.get(f"{{{W_NS}}}id") or -1))
                    except ValueError:
                        pass

        new_comments_root = LET.Element(f"{{{W_NS}}}comments", nsmap=_NSMAP)
        cid = max_id + 1

        placed = 0
        paras = [c for c in body if c.tag == f"{{{W_NS}}}p"]
        for p_el in paras:
            if placed >= max_count:
                break
            # find runs under this paragraph
            runs = [r for r in p_el if r.tag == f"{{{W_NS}}}r"]
            for rnode in runs:
                if placed >= max_count:
                    break
                # find text node
                tnode = None
                for ch in rnode:
                    if ch.tag == f"{{{W_NS}}}t":
                        tnode = ch
                        break
                if tnode is None:
                    continue
                txt = (tnode.text or "")
                if target in txt:
                    # create comment element
                    c_el = LET.SubElement(new_comments_root, f"{{{W_NS}}}comment")
                    c_el.set(f"{{{W_NS}}}id", str(cid))
                    c_el.set(f"{{{W_NS}}}author", author_display)
                    c_el.set(
                        f"{{{W_NS}}}date",
                        datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
                    )
                    c_el.set(f"{{{W_NS}}}initials", "RV")
                    cp = LET.SubElement(c_el, f"{{{W_NS}}}p")
                    cr_el = LET.SubElement(cp, f"{{{W_NS}}}r")
                    ct = LET.SubElement(cr_el, f"{{{W_NS}}}t")
                    ct.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                    ct.text = comment_text

                    start = LET.Element(f"{{{W_NS}}}commentRangeStart")
                    start.set(f"{{{W_NS}}}id", str(cid))
                    end = LET.Element(f"{{{W_NS}}}commentRangeEnd")
                    end.set(f"{{{W_NS}}}id", str(cid))
                    ref_run = LET.Element(f"{{{W_NS}}}r")
                    ref = LET.SubElement(ref_run, f"{{{W_NS}}}commentReference")
                    ref.set(f"{{{W_NS}}}id", str(cid))

                    # insert start before this run, end and ref after this run
                    children = list(p_el)
                    try:
                        child_idx = children.index(rnode)
                    except ValueError:
                        child_idx = None
                    if child_idx is None:
                        # fallback to paragraph-level insertion
                        p_el.insert(0, start)
                        p_el.append(end)
                        p_el.append(ref_run)
                    else:
                        p_el.insert(child_idx, start)
                        # re-evaluate index after insertion
                        children = list(p_el)
                        try:
                            new_idx = children.index(rnode)
                        except ValueError:
                            new_idx = None
                        if new_idx is None:
                            p_el.insert(0, start)
                            p_el.append(end)
                            p_el.append(ref_run)
                        else:
                            p_el.insert(new_idx + 1, end)
                            p_el.insert(new_idx + 2, ref_run)

                    cid += 1
                    placed += 1

        if placed == 0:
            # nothing found, return original bytes
            return docx_bytes

        merged_comments = _merge_comments_xml(old_comments, new_comments_root)

        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                name = item.filename
                if name in (rels_path, "word/document.xml", "word/comments.xml", "[Content_Types].xml"):
                    continue
                zout.writestr(item, zin.read(name))

            zout.writestr(rels_path, LET.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True))
            zout.writestr("word/document.xml", LET.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True))
            zout.writestr("word/comments.xml", merged_comments)

            zout.writestr("[Content_Types].xml", _patch_content_types(zin))

    return buf.getvalue()

def generate_title_page_files(
    final_data_json: dict[str, Any],
    *,
    mode: str,
    old_data_json: dict[str, Any] | None = None,
    comment_author: str = "格式修订",
) -> Tuple[bytes, Optional[bytes]]:
    """
    按 mode 生成 Clean（必选）；mode == "revise_format" 时额外生成痕迹版（批注）。

    返回一个二元组 (clean_bytes, tracked_bytes_or_none)：
      - clean_bytes: 清稿的 docx 二进制数据（bytes）
      - tracked_bytes_or_none: 若 mode == "revise_format" 则为痕迹版的 docx bytes，否则为 None
    """
    # 生成清稿文档对象并保存到内存 BytesIO 获取 bytes
    doc = build_title_page_document(final_data_json)
    bio = io.BytesIO()
    doc.save(bio)
    clean_bytes = bio.getvalue()

    # --- 新增：如果 funding 中有 agency 为空（被显示为 XXX），则在 Clean 版 funding 的 XXX 处添加批注 ---
    # 仅对那些有 grant_no 的 funding item 生效（与 _funding_text_final 行为一致）
    items = final_data_json.get("funding")  or []
    missing_agency_count = 0
    for it in items:
        grant_no = _sanitize((it or {}).get("grant_no"))
        if not grant_no:
            continue
        agency = _sanitize((it or {}).get("agency"))
        if not agency:
            missing_agency_count += 1

    if missing_agency_count > 0:
        try:
            clean_bytes = inject_comments_to_runs_matching_text(
                clean_bytes,
                target="XXX",
                max_count=missing_agency_count,
                comment_text="请提供基金名称",
                author_display=comment_author,
            )
        except Exception:
            # 注入批注为辅助功能，若失败不应阻塞主流程，故忽略异常
            pass
    # --- 新增结束 ---

    if (mode or "").strip() != "revise_format":
        return clean_bytes, None

    # 生成痕迹版所需的注释信息
    msgs = collect_section_comments(final_data_json, old_data_json)
    num_aff = len(_collect_affiliations_final(final_data_json))

    # 将清稿 bytes 注入批注，得到痕迹版 bytes
    tracked_bytes = inject_comments_into_docx_bytes(
        clean_bytes,
        msgs,
        num_affiliations=num_aff,
        author_display=comment_author,
    )
    return clean_bytes, tracked_bytes