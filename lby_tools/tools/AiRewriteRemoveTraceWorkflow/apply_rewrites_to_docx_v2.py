# -*- coding: utf-8 -*-
"""
Dify 插件：按段落改写 docx，并尽量保留超链接结构。
改进：摘要/正文分段维护缩写状态；斜体拆分时为新 run 复制字号/字体；写入时清除误继承的上下标；
      含超链段落避免按比例切碎；
      原文手动标蓝的连续片段若在改写后仍存在则恢复同色，不存在则不标蓝（避免误标记）。
      正文段落清除误继承的加粗（w:b），Heading / Word「Title」样式或 iter2 中 title 为论文题目（Title 等）时保留加粗。
"""

from __future__ import annotations

from collections.abc import Generator
from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from typing import Any
from copy import deepcopy
import json
import os
import re
import tempfile
import unicodedata
import requests
from urllib.parse import urlparse


# ---------------------------------------------------------------------------
# Tool 入口
# ---------------------------------------------------------------------------


class LbyToolsTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage, None, None]:
        payload = json.loads(tool_parameters["payload_json"])
        iter2_list = payload["iter2_list"]
        orgin_docx_url = payload["orgin_docx_url"]
        abbrev_map = payload.get("abbrev_map") or []

        tmp_path = download_to_temp(orgin_docx_url)
        para_items = build_para_items(iter2_list)

        out_bytes = apply_rewrites_to_docx(
            tmp_path,
            para_items,
            abbrev_map=abbrev_map,
        )
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = "final.docx"

        yield self.create_blob_message(
            blob=out_bytes,
            meta={
                "mime_type": mime_type,
                "filename": filename,
            },
        )


# ---------------------------------------------------------------------------
# iter2 / 缩写 scope
# ---------------------------------------------------------------------------


def build_para_items(iter2_list: list[dict]) -> dict[str, dict[str, Any]]:
    """
    para_id -> { rewritten_text, title?, orig_text? }
    """
    out: dict[str, dict[str, Any]] = {}
    for item in iter2_list or []:
        pid = item.get("para_id")
        txt = (item.get("rewritten_text") or "").strip()
        if not pid or not txt:
            continue
        out[pid] = {
            "rewritten_text": txt,
            "title": (item.get("title") or "").strip() or None,
            "orig_text": item.get("orig_text"),
        }
    return out


def abbrev_scope_for_title(title: str | None) -> str:
    """
    期刊常见：摘要里首次定义缩写；正文（含 Introduction）再算一轮。
    Keywords 等跟随正文 scope，避免被摘要的 seen 状态影响。
    """
    if not title:
        return "body"
    t = title.strip().lower()
    if t == "abstract":
        return "abstract"
    return "body"


# ---------------------------------------------------------------------------
# 文本消毒 / 缩写规范化
# ---------------------------------------------------------------------------

_ILLEGAL_XML_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")


def sanitize_for_docx_xml(s: str) -> str:
    if s is None:
        return ""
    s = _ILLEGAL_XML_RE.sub("", s)
    s = s.encode("utf-8", "surrogatepass").decode("utf-8", "ignore")
    return s


def normalize_with_abbrev_map(
    text: str,
    abbrev_map: list[dict],
    seen_abbr: set[str],
) -> str:
    if not text:
        return text

    items = sorted(
        [x for x in (abbrev_map or []) if x.get("abbr") and x.get("full")],
        key=lambda d: len(d["full"]),
        reverse=True,
    )

    for item in items:
        full = item["full"].strip()
        abbr = item["abbr"].strip()

        pat_full_paren = re.compile(
            rf"\b{re.escape(full)}\s*\(\s*{re.escape(abbr)}\s*\)",
            flags=re.IGNORECASE,
        )
        pat_full_only = re.compile(rf"\b{re.escape(full)}\b", flags=re.IGNORECASE)

        if abbr in seen_abbr:
            text = pat_full_paren.sub(abbr, text)
            text = pat_full_only.sub(abbr, text)
        else:
            if pat_full_paren.search(text) or pat_full_only.search(text):
                seen_abbr.add(abbr)

    return text


# ---------------------------------------------------------------------------
# URL：归一化与去重（超链 run 与正文重复显示时清空超链 run）
# ---------------------------------------------------------------------------

_URL_RE = re.compile(r"https?://[^\s)\]>]+", flags=re.IGNORECASE)


def canonical_url_for_dedup(u: str) -> str:
    """
    归一化 URL，用于判断「原文与改写是否仍包含同一链接」。
    消除末尾 /、大小写、http/https 等导致的误判，避免重复拼接。
    """
    s = (u or "").strip().rstrip(".,);」』】]’'\"")
    if not s:
        return ""
    try:
        p = urlparse(s)
        raw_scheme = (p.scheme or "http").lower()
        # 仅用于判断是否「同一链接」：http 与 https 视为等价
        scheme = "http" if raw_scheme in ("http", "https") else raw_scheme
        netloc = (p.netloc or "").lower()
        path = (p.path or "").rstrip("/")
        q = f"?{p.query}" if p.query else ""
        return f"{scheme}://{netloc}{path}{q}".lower()
    except Exception:
        return re.sub(r"/+$", "", s.lower())


def _collect_canonical_urls_in_text(text: str) -> set[str]:
    return {canonical_url_for_dedup(m.group(0)) for m in _URL_RE.finditer(text or "")}


def _normalize_hyperlink_display_text(s: str) -> str:
    """去掉 Word 常见不可见字符，避免 fullmatch URL 失败。"""
    s = (s or "").replace("\u00a0", " ")
    for ch in ("\u200b", "\u200c", "\u200d", "\ufeff"):
        s = s.replace(ch, "")
    return s.strip()


def _hyperlink_run_text_is_redundant_url(t: str, new_canon: set[str]) -> bool:
    """超链/域 run 的显示文本仅为某 URL，且该 URL 已在改写正文文本中出现过。"""
    t = _normalize_hyperlink_display_text(t)
    if not t or not new_canon:
        return False
    t_stripped = t.rstrip(".,;，、）)]>'\"")
    for candidate in (t, t_stripped):
        if not candidate:
            continue
        m = _URL_RE.fullmatch(candidate)
        if m is not None and canonical_url_for_dedup(m.group(0)) in new_canon:
            return True
        found = list(_URL_RE.finditer(candidate))
        if len(found) == 1:
            u = found[0].group(0)
            rest = candidate[found[0].end() :].strip(".,;，、）)]>'\" \t")
            if not rest and canonical_url_for_dedup(u) in new_canon:
                return True
    return False


def _iter_w_t_document_order(paragraph) -> Generator:
    """段内文档顺序的所有 w:t（含超链、域、修订等），与 Word 可见字符顺序一致。"""
    yield from paragraph._p.iter(qn("w:t"))


def _paragraph_plain_text_from_wt(paragraph) -> str:
    """用 w:t 拼接整段文本；比 paragraph.runs 拼接更可靠（避免 tab/分拆节点与 runs API 不一致）。"""
    return "".join((t.text or "") for t in _iter_w_t_document_order(paragraph))


def _remove_wt_char_range(paragraph, start: int, end: int) -> None:
    """按 _paragraph_plain_text_from_wt 的下标，从各 w:t 删掉 [start, end)。"""
    if start >= end:
        return
    pos = 0
    for t in _iter_w_t_document_order(paragraph):
        tx = t.text or ""
        n = len(tx)
        s0, e0 = pos, pos + n
        pos = e0
        if e0 <= start or s0 >= end:
            continue
        lo = max(start, s0) - s0
        hi = min(end, e0) - s0
        t.text = tx[:lo] + tx[hi:]


def _iter_w_r_elements_document_order(paragraph) -> Generator:
    """
    段内文档顺序的 w:r（含 w:hyperlink 下的 run），与 Word 显示顺序一致。
    """
    for child in paragraph._p:
        if child.tag == qn("w:r"):
            yield child
        elif child.tag == qn("w:hyperlink"):
            for sub in child:
                if sub.tag == qn("w:r"):
                    yield sub


def _run_has_fld_separate_only_or_with_rpr(r_el) -> bool:
    """
    True：该 w:r 仅承载 fldChar separate（可有 rPr），无 w:t、无指令。
    用于对接「半截域」：丢失了 begin/instr/end，只剩 separate + 结果。
    """
    has_sep = False
    for node in list(r_el):
        if node.tag == qn("w:rPr"):
            continue
        if node.tag == qn("w:fldChar") and node.get(qn("w:fldCharType")) == "separate":
            has_sep = True
            continue
        return False
    return has_sep


def _combined_w_t_after_index(children: list, start_run_idx: int) -> tuple[str, int]:
    """
    从段落子元素列表中下标 start_run_idx 起，沿连续的顶层 w:r 合并 w:t，
    遇到非 w:r 或其它域节点即停止。
    返回 (合并文本, 下一个未消费的下标)。
    """
    parts: list[str] = []
    j = start_run_idx
    while j < len(children):
        node = children[j]
        if node.tag != qn("w:r"):
            break
        if node.find(qn("w:fldChar")) is not None or node.find(qn("w:instrText")) is not None:
            break
        for t_el in node.findall(qn("w:t")):
            parts.append(t_el.text or "")
        j += 1
    return "".join(parts), j


def remove_orphan_fld_separate_redundant_url_tail(paragraph, new_text: str) -> None:
    """
    处理「半截 HYPERLINK 域」常见于保存/改写后：本段已无 begin/instr/end，
    却仍残留 ``fldChar separate`` + **仅为一串 URL** 的 w:r。此时 _run_is_inside_hyperlink_field 无法识别，
    strip_redundant_hyperlink_field_result_urls 的栈亦为 empty，无法在 ``separate`` 时进入结果区。

    若 ``separate`` 之后（可跨越若干空壳 w:r）合并文本仅为某已在 new_text 出现过的 URL，则删除：
    **承载 separate 的 w:r**，以及 **_combined_w_t_after_index** 消费的连续顶层 w:r。

    仅遍历 **w:p 直接子结点**下的 w:r；常见 Methods 段落即此结构。
    """
    if not new_text:
        return
    new_canon = _collect_canonical_urls_in_text(new_text)
    if not new_canon:
        return

    body = paragraph._p
    for _ in range(32):
        children = list(body)
        removed_any = False
        for i, cur in enumerate(children):
            if cur.tag != qn("w:r") or cur.getparent() is not body:
                continue
            if not _run_has_fld_separate_only_or_with_rpr(cur):
                continue

            trailing, end_j = _combined_w_t_after_index(children, i + 1)
            if not _hyperlink_run_text_is_redundant_url(trailing, new_canon):
                continue

            to_drop = [cur]
            for j in range(i + 1, end_j):
                cand = children[j]
                if cand.tag == qn("w:r") and cand.getparent() is body:
                    to_drop.append(cand)

            for el in to_drop:
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)
            removed_any = True
            break

        if not removed_any:
            break


def strip_redundant_hyperlink_field_result_urls(paragraph, new_text: str) -> None:
    """
    按 Word 域结构（fldChar begin → instrText → separate → 结果区 → end）扫描段落，
    若「结果区」的合并文本仅为已在改写正文 new_text 中出现的 URL（归一化后），则清空对应 w:t。

    不依赖 python-docx 的 paragraph.runs 与域 run 的对应关系；亦不要求 instr 必含 HYPERLINK
    （仍用 _hyperlink_run_text_is_redundant_url 约束，避免误删非 URL 域结果）。
    """
    if not new_text:
        return
    new_canon = _collect_canonical_urls_in_text(new_text)
    if not new_canon:
        return

    stack: list[dict[str, Any]] = []

    def _flush_field_end() -> None:
        if not stack:
            return
        frame = stack.pop()
        result_text = "".join((t.text or "") for t in frame["result_wts"])
        if _hyperlink_run_text_is_redundant_url(result_text, new_canon):
            for t in frame["result_wts"]:
                t.text = ""

    for r_elem in _iter_w_r_elements_document_order(paragraph):
        for node in r_elem:
            if node.tag == qn("w:fldChar"):
                ftype = node.get(qn("w:fldCharType"))
                if ftype == "begin":
                    stack.append(
                        {
                            "instr_parts": [],
                            "result_wts": [],
                            "phase": "header",
                        }
                    )
                elif ftype == "separate":
                    if stack:
                        stack[-1]["phase"] = "result"
                elif ftype == "end":
                    _flush_field_end()
            elif node.tag == qn("w:instrText"):
                if stack and stack[-1]["phase"] == "header":
                    stack[-1]["instr_parts"].append(node.text or "")
            elif node.tag == qn("w:t"):
                if stack and stack[-1]["phase"] == "result":
                    stack[-1]["result_wts"].append(node)


def strip_trailing_duplicate_url_if_repeated_earlier(paragraph) -> None:
    """
    若段末的 URL 与段中更早出现的某 URL 归一化后相同，则删掉段末这一次。
    使用 w:t 级拼接与删除，避免 paragraph.runs 与底层 XML 不一致导致删不掉。
    """
    for _ in range(16):
        s = unicodedata.normalize("NFC", _paragraph_plain_text_from_wt(paragraph))
        matches = list(_URL_RE.finditer(s))
        if len(matches) < 2:
            return
        last_m = matches[-1]
        last_c = canonical_url_for_dedup(last_m.group(0))
        if not last_c:
            return
        prev_canons = {canonical_url_for_dedup(m.group(0)) for m in matches[:-1]}
        if last_c not in prev_canons:
            return
        _remove_wt_char_range(paragraph, last_m.start(), last_m.end())


def strip_redundant_hyperlink_xml(paragraph, new_text: str) -> None:
    """
    直接在 w:hyperlink 下清空显示文本：不依赖 python-docx 能否把 run 识别为「在超链内」。
    """
    if not new_text:
        return
    new_canon = _collect_canonical_urls_in_text(new_text)
    if not new_canon:
        return
    for hyp in paragraph._p.iter(qn("w:hyperlink")):
        combined = "".join((t.text or "") for t in hyp.iter(qn("w:t")))
        if not _hyperlink_run_text_is_redundant_url(combined, new_canon):
            continue
        for t in hyp.iter(qn("w:t")):
            t.text = ""


# ---------------------------------------------------------------------------
# 超链接 / 域 检测
# ---------------------------------------------------------------------------


def _run_is_inside_hyperlink_field(run) -> bool:
    el = run._element
    node = el
    saw_begin = False
    saw_hyperlink_instr = False

    for _ in range(200):
        node = node.getprevious()
        if node is None:
            break
        # 同一 w:r 内可能同时含 fldChar begin 与 instrText；必须先扫 instr，避免 find(begin) 先匹配就 break。
        for instr in node.iter(qn("w:instrText")):
            if "HYPERLINK" in (instr.text or ""):
                saw_hyperlink_instr = True
                break
        fld = node.find(qn("w:fldChar"))
        if fld is not None and fld.get(qn("w:fldCharType")) == "begin":
            saw_begin = True
            break

    if not (saw_begin and saw_hyperlink_instr):
        return False

    node = el
    for _ in range(400):
        node = node.getnext()
        if node is None:
            break
        fld = node.find(qn("w:fldChar"))
        if fld is not None and fld.get(qn("w:fldCharType")) == "end":
            return True
    return False


def run_is_in_hyperlink_container(run) -> bool:
    el = run._element
    parent = el.getparent()
    while parent is not None:
        tag = parent.tag
        if tag == qn("w:hyperlink"):
            return True
        if isinstance(tag, str) and tag.endswith("}hyperlink"):
            return True
        parent = parent.getparent()
    return False


def iter_runs_skip_hyperlink(paragraph):
    for run in paragraph.runs:
        if run_is_in_hyperlink_container(run):
            continue
        if _run_is_inside_hyperlink_field(run):
            continue
        yield run


def paragraph_has_hyperlink_or_field(paragraph) -> bool:
    for run in paragraph.runs:
        if run_is_in_hyperlink_container(run):
            return True
        if _run_is_inside_hyperlink_field(run):
            return True
    return False


def strip_duplicate_url_in_hyperlinks_and_fields(paragraph, new_text: str) -> None:
    """
    改写文已在普通 run 中包含完整 URL 时，原 Word 里保留的超链/域 run 仍会再显示一遍链接。
    若超链 run 的显示文本整段即某 URL，且该 URL（归一化后）已在正文中出现，则清空该 run，避免重复。
    """
    if not new_text:
        return
    strip_redundant_hyperlink_field_result_urls(paragraph, new_text)
    # 半截域（仅余 separate + URL 显示 run）无法用栈配对；在去重链路中追加物理删除
    remove_orphan_fld_separate_redundant_url_tail(paragraph, new_text)
    strip_redundant_hyperlink_xml(paragraph, new_text)
    new_canon = _collect_canonical_urls_in_text(new_text)

    for run in paragraph.runs:
        if not (run_is_in_hyperlink_container(run) or _run_is_inside_hyperlink_field(run)):
            continue
        t = run.text or ""
        if not t.strip():
            continue
        if _hyperlink_run_text_is_redundant_url(t, new_canon):
            run.text = ""


# ---------------------------------------------------------------------------
# run 格式：清除上下标误继承；清除颜色（沿用你原逻辑）
# ---------------------------------------------------------------------------


def reset_run_vert_align_to_baseline(run) -> None:
    """避免按比例塞字时继承 superscript/subscript。"""
    r = run._element
    rPr = r.rPr
    if rPr is None:
        return
    va = rPr.find(qn("w:vertAlign"))
    if va is not None:
        rPr.remove(va)


def reset_run_color_to_default(run) -> None:
    r = run._element
    rPr = r.rPr
    if rPr is None:
        return
    color = rPr.find(qn("w:color"))
    if color is not None:
        rPr.remove(color)


def reset_paragraph_run_colors_to_default(paragraph) -> None:
    for run in paragraph.runs:
        rPr = run._element.rPr
        if rPr is None:
            continue
        if rPr.find(qn("w:color")) is not None:
            reset_run_color_to_default(run)


# 判定是否为「应保留标题式加粗」的段落（Heading / 标题 / Title 等），避免把正文误伤
_HEADING_STYLE_RE = re.compile(
    r"^(heading\s*\d+|标题\s*\d+|toc\s*heading)$",
    re.IGNORECASE,
)


def _paragraph_has_outline_level(paragraph) -> bool:
    try:
        pPr = paragraph._p.pPr
    except Exception:
        pPr = None
    if pPr is None:
        return False
    return pPr.find(qn("w:outlineLvl")) is not None


def paragraph_is_heading_style(paragraph) -> bool:
    """
    True：该段使用 Word 内置/常见标题样式或有大纲级别，改写后保留样式带来的加粗。
    False：按正文处理，清除各 run 上误继承的粗体。
    """
    try:
        raw = (paragraph.style.name or "").strip()
    except Exception:
        raw = ""
    if not raw:
        return _paragraph_has_outline_level(paragraph)

    low = raw.lower()
    if low in ("title", "subtitle"):
        return True
    if _HEADING_STYLE_RE.match(low):
        return True
    # 英文简写 Heading1、中文「标题」+ 数字
    if low.startswith("heading") and any(c.isdigit() for c in raw):
        return True
    if "标题" in raw and any(c.isdigit() for c in raw):
        return True

    return _paragraph_has_outline_level(paragraph)


def payload_marks_document_title_row(item: dict | None) -> bool:
    """
    iter2 中 title 表示段落所属区块名；为 Title / 论文题目 等时指「文题」行，
    即时样式为正文也应保留加粗，避免误清除。
    """
    if not item:
        return False
    t = (item.get("title") or "").strip().lower()
    return t in (
        "title",
        "paper title",
        "article title",
        "manuscript title",
        "论文题目",
        "文章标题",
    )


def reset_run_bold_off(run) -> None:
    """去掉 run 上的显式加粗（含东亚/复杂文种 w:bCs），不改变样式里其它字符格式。"""
    try:
        run.bold = False
    except Exception:
        pass
    rPr = run._element.rPr
    if rPr is None:
        return
    for tag in (qn("w:b"), qn("w:bCs")):
        el = rPr.find(tag)
        if el is not None:
            rPr.remove(el)


def reset_body_paragraph_run_bold(paragraph, item: dict | None = None) -> None:
    """非标题段：全文所有 run（含超链内）取消加粗，避免整段/半截继承粗体。"""
    if paragraph_is_heading_style(paragraph) or payload_marks_document_title_row(item):
        return
    for run in paragraph.runs:
        reset_run_bold_off(run)


# ---------------------------------------------------------------------------
# 原文「标蓝」采集与条件恢复（新文中仍有该子串才上色）
# ---------------------------------------------------------------------------

_BLUE_HEX_WHITELIST = {
    "0000FF",
    "0563C1",  # 常见超链蓝
    "0070C0",
    "4472C4",
    "4F81BD",
    "17365D",
    "365F91",
    "31859B",
}


def _color_element_hex(color_el) -> str | None:
    """从 w:color 取出 RRGGBB；主题色 hyperlink 按常见蓝色近似。"""
    v = getattr(color_el, "val", None)
    if v is None:
        v = color_el.get(qn("w:val"))
    if isinstance(v, str) and len(v) == 6 and v.lower() != "auto":
        try:
            int(v, 16)
            return v.upper()
        except ValueError:
            pass
    tc = getattr(color_el, "themeColor", None)
    if tc is None:
        tc = color_el.get(qn("w:themeColor"))
    if tc == "hyperlink":
        return "0563C1"
    return None


def _is_blue_like_hex(hex6: str) -> bool:
    if not hex6 or len(hex6) != 6:
        return False
    hu = hex6.upper()
    if hu in _BLUE_HEX_WHITELIST:
        return True
    try:
        r = int(hex6[0:2], 16)
        g = int(hex6[2:4], 16)
        b = int(hex6[4:6], 16)
    except ValueError:
        return False
    if b >= 160 and b >= r + 35 and b >= g + 35:
        return True
    return False


def collect_blue_fragments(paragraph) -> list[tuple[str, str]]:
    """
    扫描该段在改写前的 runs：连续、且为「蓝色系」的 run 合并为片段。
    返回 [(fragment_text, RRGGBB), ...]。
    说明：python-docx 的 paragraph.runs 通常不含 w:hyperlink 内的 run，
    若蓝色仅在超链显示文本上，可能采不到（需再扩展遍历 _p）。
    """
    out: list[tuple[str, str]] = []
    buf = ""
    buf_hex: str | None = None

    for run in paragraph.runs:
        hex6: str | None = None
        rPr = run._element.rPr
        if rPr is not None:
            cel = rPr.find(qn("w:color"))
            if cel is not None:
                raw = _color_element_hex(cel)
                if raw and _is_blue_like_hex(raw):
                    hex6 = raw
        t = run.text or ""
        if hex6 is not None and t:
            if buf_hex is None or buf_hex != hex6:
                if buf and buf_hex:
                    out.append((buf, buf_hex))
                buf = t
                buf_hex = hex6
            else:
                buf += t
        else:
            if buf and buf_hex:
                out.append((buf, buf_hex))
            buf = ""
            buf_hex = None

    if buf and buf_hex:
        out.append((buf, buf_hex))
    return out


def _copy_run_bold_italic(source_run, target_run) -> None:
    try:
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
    except Exception:
        pass


def apply_hex_color_to_char_range(paragraph, start: int, end: int, hex6: str) -> None:
    """把段落内 [start, end) 字符设为 hex6（RRGGBB）；必要时拆分 run。"""
    if start >= end or len(hex6) != 6:
        return
    rgb = RGBColor(int(hex6[0:2], 16), int(hex6[2:4], 16), int(hex6[4:6], 16))

    for _ in range(2048):
        pos = 0
        changed = False
        for run in list(paragraph.runs):
            text = run.text or ""
            rlen = len(text)
            r0 = pos
            r1 = pos + rlen
            pos = r1

            if r1 <= start or r0 >= end:
                continue
            lo = max(start, r0) - r0
            hi = min(end, r1) - r0
            if lo >= hi:
                continue

            changed = True
            if lo == 0 and hi == rlen:
                run.font.color.rgb = rgb
            else:
                before = text[:lo]
                mid = text[lo:hi]
                after = text[hi:]
                run.text = before
                mid_run = paragraph.add_run(mid)
                copy_base_character_format(run, mid_run)
                _copy_run_bold_italic(run, mid_run)
                mid_run.font.color.rgb = rgb
                run._element.addnext(mid_run._element)
                if after:
                    after_run = paragraph.add_run(after)
                    copy_base_character_format(run, after_run)
                    _copy_run_bold_italic(run, after_run)
                    mid_run._element.addnext(after_run._element)
            break

        if not changed:
            break


def reapply_blue_fragments_if_still_present(paragraph, fragments: list[tuple[str, str]]) -> None:
    """
    若改写后的段落全文仍包含某蓝色原文片段，则对该子串出现处恢复为记录下的蓝色；
    若某片段已不在新文中，则不做任何着色（不误标蓝）。
    先匹配较长片段，减少短串误伤；同一文本只保留首次出现的色号。
    """
    if not fragments:
        return

    ordered = sorted(fragments, key=lambda x: len(x[0]), reverse=True)
    text_to_hex: dict[str, str] = {}
    for frag, hx in ordered:
        if frag and frag not in text_to_hex:
            text_to_hex[frag] = hx

    full = paragraph.text
    spans: list[tuple[int, int, str]] = []
    for frag, hx in text_to_hex.items():
        if frag not in full:
            continue
        i = 0
        while True:
            j = full.find(frag, i)
            if j < 0:
                break
            spans.append((j, j + len(frag), hx))
            i = j + max(1, len(frag))

    spans.sort(key=lambda x: x[0], reverse=True)
    for s, e, hx in spans:
        apply_hex_color_to_char_range(paragraph, s, e, hx)


# ---------------------------------------------------------------------------
# 段落文本替换
# ---------------------------------------------------------------------------


def replace_text_preserve_hyperlinks(paragraph, new_text: str) -> None:
    """
    - 无超链：整段视作单一文本块，写入第一个 run，清空其余（减少多 run 比例切碎）。
    - 有超链：超链 run 不动；其余非超链 run 合并承载 new_text。
    """
    new_text = sanitize_for_docx_xml(new_text)

    runs = list(iter_runs_skip_hyperlink(paragraph))
    if not runs:
        return

    for r in runs:
        reset_run_vert_align_to_baseline(r)

    if not paragraph_has_hyperlink_or_field(paragraph):
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""
        return

    # 含超链：不要把 new_text 按旧 run 长度切碎；整块写入第一个非超链 run。
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def redistribute_text_to_docx_runs(new_text: str, runs: list) -> None:
    """保留：仅在你明确需要多 run 比例时使用（当前默认未调用）。"""
    if not runs:
        return

    new_text = sanitize_for_docx_xml(new_text)
    for r in runs:
        reset_run_vert_align_to_baseline(r)

    old_lengths = [len(r.text or "") for r in runs]
    total_old = sum(old_lengths)

    if total_old == 0:
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""
        return

    consumed = 0
    for i, (r, l) in enumerate(zip(runs, old_lengths)):
        if i == len(runs) - 1:
            part = new_text[consumed:]
        else:
            take = round(len(new_text) * (l / total_old))
            part = new_text[consumed : consumed + take]

        r.text = sanitize_for_docx_xml(part)
        consumed += len(part)


# ---------------------------------------------------------------------------
# 斜体：新 run 复制字号/字体；仅对匹配段设置 italic
# ---------------------------------------------------------------------------

# 只拷贝「字形与字号」相关节点，不拷贝 i/b/u/color，避免带跑样式意图
_RPR_CHAR_TAGS = (
    qn("w:rFonts"),
    qn("w:sz"),
    qn("w:szCs"),
    qn("w:lang"),
    qn("w:kern"),
    qn("w:w"),
    qn("w:spacing"),
    qn("w:position"),
)


def copy_base_character_format(source_run, target_run) -> None:
    """
    将 source_run 的字号/字体等基准字符格式复制到 target_run。
    add_run() 默认不带 sz/rFonts，整段会掉进 Word 默认小字号 —— 即你看到的 in vivo 段全文变小。
    """
    s_rpr = source_run._element.rPr
    if s_rpr is None:
        return
    t_rpr = target_run._element.get_or_add_rPr()
    for child in list(t_rpr):
        if child.tag in _RPR_CHAR_TAGS:
            t_rpr.remove(child)
    for tag in _RPR_CHAR_TAGS:
        src_el = s_rpr.find(tag)
        if src_el is not None:
            t_rpr.append(deepcopy(src_el))


ITALIC_REGEXES = [
    r"\b[pP](?:\s*-\s*value)?\s*(?:<=|>=|<|>|=|≤|≥)\s*(?:0?\.\d+|\.\d+|\d+(?:\.\d+)?)\b",
    r"\bin\s+vivo\b",
    r"\bin\s+vitro\b",
]


def apply_italics_in_paragraph(paragraph, regexes: list[str] | None = None) -> None:
    regexes = regexes or ITALIC_REGEXES
    pattern = re.compile("|".join(f"(?:{r})" for r in regexes), flags=re.IGNORECASE)

    for run in list(iter_runs_skip_hyperlink(paragraph)):
        txt = run.text or ""
        if not txt:
            continue

        matches = list(pattern.finditer(txt))
        if not matches:
            continue

        ref_run = run
        run.text = ""
        prev = run
        last = 0

        def _insert_piece(piece_text: str, italic: bool) -> None:
            nonlocal prev
            if piece_text == "":
                return
            newr = paragraph.add_run(piece_text)
            copy_base_character_format(ref_run, newr)
            prev._element.addnext(newr._element)
            prev = newr
            newr.italic = bool(italic)

        for m in matches:
            _insert_piece(txt[last : m.start()], italic=False)
            _insert_piece(txt[m.start() : m.end()], italic=True)
            last = m.end()

        _insert_piece(txt[last:], italic=False)


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------


def apply_rewrites_to_docx(
    input_docx_path: str,
    para_items: dict[str, dict[str, Any]],
    abbrev_map: list[dict],
) -> bytes:
    doc = Document(input_docx_path)

    seen_by_scope: dict[str, set[str]] = {
        "abstract": set(),
        "body": set(),
    }

    for idx, p in enumerate(doc.paragraphs, start=1):
        pid = f"p_{idx:05d}"
        item = para_items.get(pid)
        if not item:
            continue

        blue_frags = collect_blue_fragments(p)

        new_text = item["rewritten_text"]
        title = item.get("title")

        scope = abbrev_scope_for_title(title)
        seen = seen_by_scope[scope]

        new_text = normalize_with_abbrev_map(new_text, abbrev_map or [], seen)

        replace_text_preserve_hyperlinks(p, new_text)
        strip_duplicate_url_in_hyperlinks_and_fields(p, new_text)
        strip_trailing_duplicate_url_if_repeated_earlier(p)
        reset_body_paragraph_run_bold(p, item)
        reset_paragraph_run_colors_to_default(p)
        apply_italics_in_paragraph(p)
        reapply_blue_fragments_if_still_present(p, blue_frags)
        strip_trailing_duplicate_url_if_repeated_earlier(p)

    out_tmp_path = None
    try:
        fd, out_tmp_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)

        doc.save(out_tmp_path)

        with open(out_tmp_path, "rb") as f:
            return f.read()
    finally:
        if out_tmp_path and os.path.exists(out_tmp_path):
            try:
                os.remove(out_tmp_path)
            except Exception:
                pass

        if input_docx_path and os.path.exists(input_docx_path):
            try:
                os.remove(input_docx_path)
            except Exception:
                pass


def download_to_temp(url: str, suffix: str = ".docx", headers: dict | None = None, timeout: int = 60) -> str:
    headers = headers or {}
    r = requests.get(url, stream=True, headers=headers, timeout=timeout)
    r.raise_for_status()
    fd, path = tempfile.mkstemp(suffix=suffix)
    os.close(fd)
    with open(path, "wb") as f:
        for chunk in r.iter_content(chunk_size=1024 * 1024):
            if chunk:
                f.write(chunk)
    return path
